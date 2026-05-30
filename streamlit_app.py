"""
PBS Support Tool
Tab 1: 📊 Behaviour Recording — upload 30-day recording → extract behaviours, patterns & triggers.
Tab 2: 📄 Generate from PBSP — upload PBSP → Support Reference Card + ABC Recording Form.
Tab 3: 🧠 Strategy Recommender — behaviours → AI strategy recommendations + PDF report.
"""

import json
import os
from datetime import date
import streamlit as st
import anthropic
from io import BytesIO
from docx import Document as DocxDocument
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.colors import HexColor, white
from reportlab.pdfbase.pdfmetrics import stringWidth

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

st.set_page_config(page_title="PBS Support Tool", page_icon="📋", layout="centered")

DEEP_BLUE  = HexColor('#0D4F6E')
TEAL       = HexColor('#1A9B8A')
BLUE2      = HexColor('#1A5276')
GREEN      = HexColor('#27AE60')
RED        = HexColor('#C0392B')
AMBER      = HexColor('#D4700A')
LIGHT_TEAL = HexColor('#E0F6F3')
LIGHT_BLUE = HexColor('#DCEEF8')
LIGHT_GRN  = HexColor('#DFFAE9')
LIGHT_RED  = HexColor('#FADBD8')
LIGHT_AMBR = HexColor('#FDEBD0')
MID_GREY   = HexColor('#CCCCCC')
DARK_TEXT  = HexColor('#1A2B35')
MED_TEXT   = HexColor('#4A7C8E')

W, H = A4
LM, RM = 20, 20
CW = W - LM - RM


# ══════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("  |  ".join(dict.fromkeys(cells)))
    return "\n".join(parts)

def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def extract_text_from_excel(file_bytes: bytes) -> str:
    """Convert Excel (e.g. Microsoft Forms export) to structured text for Claude analysis."""
    if not EXCEL_AVAILABLE:
        return ""
    wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        if ws.max_row is None or ws.max_row == 0:
            continue
        parts.append(f"[Sheet: {sheet_name}]")
        headers = []
        for ri, row in enumerate(ws.iter_rows(values_only=True)):
            cells = [str(c) if c is not None else "" for c in row]
            if not any(c.strip() for c in cells):
                continue
            if ri == 0:
                headers = cells
                parts.append("COLUMNS: " + " | ".join(cells))
            else:
                if headers:
                    row_parts = [
                        f"{headers[i]}: {cells[i]}"
                        for i in range(min(len(headers), len(cells)))
                        if i < len(cells) and cells[i].strip()
                    ]
                    if row_parts:
                        parts.append("ENTRY: " + " || ".join(row_parts))
                else:
                    if any(c.strip() for c in cells):
                        parts.append(" | ".join(cells))
    return "\n".join(parts)


# ── Built-in strategy library ──────────────────────────────────────────────────
_APP_DIR           = os.path.dirname(os.path.abspath(__file__))
_LIBRARY_FILENAMES = ["strategy_library.docx", "strategy_library.pdf", "strategy_library.txt"]

@st.cache_resource
def load_builtin_library():
    """Return (text, filename) for a bundled strategy library, or (None, None) if absent."""
    for fname in _LIBRARY_FILENAMES:
        path = os.path.join(_APP_DIR, fname)
        if os.path.exists(path):
            try:
                with open(path, "rb") as f:
                    fb = f.read()
                if fname.endswith(".docx"):
                    text = extract_text_from_docx(fb)
                elif fname.endswith(".pdf"):
                    text = extract_text_from_pdf(fb)
                else:
                    text = fb.decode("utf-8", errors="ignore")
                return text, fname
            except Exception:
                pass
    return None, None


# ══════════════════════════════════════════════════════════════════════════════
# CLAUDE — BEHAVIOUR RECORDING ANALYSIS (Tab 1)
# ══════════════════════════════════════════════════════════════════════════════

BEHAVIOUR_RECORDING_PROMPT = """
You are an experienced Positive Behaviour Support (PBS) Practitioner.
The data below comes from a 30-day behaviour recording tool completed by family members or caregivers.
The data may be exported from Microsoft Forms (as Excel), a PDF form, or a Word document.

Analyse this data and extract clinically meaningful information.
Be flexible about column names and data structure — the recording format will vary.
Look for: dates, times, behaviour descriptions, what happened before (antecedents/triggers),
what happened after (consequences / carer responses), severity, location, and any carer notes.

Return ONLY valid JSON — no commentary, no markdown fences.

JSON structure:
{
  "client_name": "client name if visible in the data, otherwise null",
  "recording_period": "date range covered e.g. '1 May – 30 May 2025', or null if not determinable",
  "total_incidents": <integer — total number of behaviour incidents recorded>,
  "behaviours": [
    {
      "label": "Short clinical label (e.g. 'Physical aggression', 'Self-injurious behaviour', 'Property destruction')",
      "carer_description": "How carers described this behaviour in their own words",
      "frequency": <integer — number of recorded incidents of this behaviour>,
      "descriptors": ["Up to 5 observable descriptions from the data"]
    }
  ],
  "patterns": {
    "time_of_day": ["Patterns in timing — e.g. 'Most incidents between 3–5pm (afternoon transition)'"],
    "day_of_week": ["Day-of-week patterns — e.g. 'Higher frequency on Mondays and after weekends'"],
    "settings": ["Locations or contexts where behaviour tends to occur"],
    "other": ["Any other notable patterns — e.g. escalation over the month, clustering of incidents"]
  },
  "triggers": ["Specific antecedents and triggers recorded by carers — be as specific as the data allows"],
  "carer_responses": ["What carers actually did in response — specific strategies or actions they used"],
  "carer_concerns": "2–3 sentence summary of what the carers appear most concerned about",
  "notes_for_practitioner": ["Clinical observations the PBS Practitioner should follow up — gaps, patterns needing functional assessment, anything that stands out"]
}

Rules:
- Assign clinically appropriate behaviour labels (not just carer wording)
- total_incidents = sum of all individual behaviour frequencies
- Be specific about triggers — avoid generic 'unknown'
- Note in notes_for_practitioner if data is sparse, inconsistent, or if key information is missing
- Do not fabricate data that isn't in the recording

RECORDING DATA:
"""

def extract_behaviour_recording(data_text: str, api_key: str) -> dict:
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model="claude-opus-4-5", max_tokens=2000,
        messages=[{"role": "user", "content": BEHAVIOUR_RECORDING_PROMPT + data_text[:40000]}]
    )
    raw = msg.content[0].text.strip()
    if raw.startswith("```"): raw = "\n".join(raw.split("\n")[1:])
    if raw.endswith("```"):   raw = "\n".join(raw.split("\n")[:-1])
    return json.loads(raw)


def recording_to_sr_format(data: dict) -> tuple:
    """Convert extract_behaviour_recording output → (client_info dict, behaviours list) for Tab 3."""
    triggers_str = "; ".join(data.get("triggers", []))
    client_info = {
        "name":      data.get("client_name") or "Client",
        "age":       "not specified",
        "diagnosis": "not specified",
        "comms":     "refer to behaviour recording",
        "other":     data.get("carer_concerns") or "none provided",
    }
    behaviours = [
        {
            "name":        b.get("label", "Behaviour"),
            "description": (b.get("carer_description") or "") +
                           (f" — {b.get('frequency')} incidents recorded" if b.get("frequency") else ""),
            "triggers":    triggers_str,
        }
        for b in data.get("behaviours", [])
    ]
    return client_info, behaviours


# ══════════════════════════════════════════════════════════════════════════════
# BEHAVIOUR RECORDING ANALYSIS REPORT PDF (Tab 1)
# ══════════════════════════════════════════════════════════════════════════════

def generate_recording_report(data: dict) -> BytesIO:
    buf = BytesIO()
    c   = rl_canvas.Canvas(buf, pagesize=A4)

    BOT, FTR_H, HDR_H = 12, 38, 68
    MINI_HDR_H         = 28
    MIN_Y              = BOT + FTR_H + 20
    RH, SH, GAP        = 17, 20, 8
    page_num           = [1]

    client_name = data.get("client_name") or "Client"
    period      = data.get("recording_period") or "30-day period"

    def draw_footer():
        drect(c, 0, BOT, W, FTR_H, fill=DEEP_BLUE)
        drect(c, 0, BOT, 8, FTR_H, fill=TEAL)
        ft = BOT + FTR_H
        dlbl(c, LM+6, ft-14, "30-Day Behaviour Recording Analysis", 9, bold=True, color=white)
        dlbl(c, W-LM, ft-14, f"Page {page_num[0]}", 9, color=HexColor('#8FC8DC'), align='right')
        dlbl(c, W-LM, ft-28, "CONFIDENTIAL — handle in line with your privacy policy",
             8, italic=True, color=HexColor('#5B9FC0'), align='right')

    def new_page():
        draw_footer(); c.showPage(); page_num[0] += 1
        drect(c, 0, H-MINI_HDR_H, W, MINI_HDR_H, fill=DEEP_BLUE)
        drect(c, 0, H-MINI_HDR_H, 8, MINI_HDR_H, fill=TEAL)
        dlbl(c, LM+6, H-MINI_HDR_H+9,
             f"Behaviour Recording Analysis — {client_name}", 10, bold=True, color=white)
        return H - MINI_HDR_H - GAP

    def ensure(y, needed):
        return new_page() if y - needed < MIN_Y else y

    # ── Header ──
    y = H
    drect(c, 0, y-HDR_H, W, HDR_H, fill=DEEP_BLUE)
    drect(c, 0, y-HDR_H, 8, HDR_H, fill=TEAL)
    dlbl(c, LM+6, y-22, "30-Day Behaviour Recording Analysis", 16, bold=True, color=white)
    dlbl(c, LM+6, y-42, client_name, 13, color=HexColor('#C8E6F5'))
    dlbl(c, LM+6, y-58, f"Recording period: {period}",
         9, italic=True, color=HexColor('#5B9FC0'))
    dlbl(c, W-LM, y-42, f"Total incidents: {data.get('total_incidents', '—')}",
         13, bold=True, color=HexColor('#C8E6F5'), align='right')
    dlbl(c, W-LM, y-58, f"Generated: {date.today().strftime('%d %B %Y')}",
         9, italic=True, color=HexColor('#5B9FC0'), align='right')
    y -= HDR_H + GAP

    # ── Behaviours identified ──
    behaviours = data.get("behaviours", [])
    if behaviours:
        total = data.get("total_incidents") or sum((b.get("frequency") or 0) for b in behaviours) or 1
        y = ensure(y, SH + 40)
        y = sec(c, y, "Behaviours identified by carers", TEAL)
        for b in behaviours:
            freq  = b.get("frequency") or 0
            pct   = f"{round(freq / total * 100)}%" if total else ""
            label = b.get("label", "Behaviour")
            desc  = b.get("carer_description", "")
            descs = b.get("descriptors", [])

            needed = 22 + (len(wrap_text(f"Described as: {desc}", "Helvetica-Oblique", 9, CW-16)) * 13 + 4 if desc else 0) + len(descs) * 14 + 10
            y = ensure(y, needed)

            # Behaviour header bar
            drect(c, LM, y-22, CW, 22, fill=LIGHT_TEAL)
            drect(c, LM, y-22, 5,  22, fill=TEAL)
            dlbl(c, LM+13, y-14, label, 11, bold=True, color=DARK_TEXT)
            dlbl(c, W-RM,  y-14, f"{freq} incidents  ({pct})",
                 10, bold=True, color=TEAL, align='right')
            y -= 22

            if desc:
                desc_lines = wrap_text(f"Carers described: {desc}", "Helvetica-Oblique", 9, CW-16)
                desc_h = len(desc_lines) * 13 + 4
                drect(c, LM, y-desc_h, CW, desc_h, fill=HexColor('#F5FFFE'))
                for li, ln in enumerate(desc_lines):
                    dlbl(c, LM+10, y-11-li*13, ln, 9, italic=True, color=MED_TEXT)
                y -= desc_h

            if descs:
                desc_item_h = len(descs) * 14 + 4
                drect(c, LM, y-desc_item_h, CW, desc_item_h, fill=LIGHT_TEAL)
                for di, d_ in enumerate(descs):
                    drect(c, LM+10, y-10-di*14, 6, 6, fill=TEAL)
                    dlbl(c,  LM+20, y- 5-di*14, d_, 9, color=DARK_TEXT)
                y -= desc_item_h
            y -= 6

    # ── Patterns ──
    patterns  = data.get("patterns", {})
    pat_items = []
    for lbl, items in [("Time of day", patterns.get("time_of_day", [])),
                        ("Day of week",  patterns.get("day_of_week", [])),
                        ("Settings",     patterns.get("settings",    [])),
                        ("Other",        patterns.get("other",       []))]:
        for it in (items or []):
            pat_items.append((lbl, it))

    if pat_items:
        pat_h = len(pat_items) * RH
        y = ensure(y, SH + pat_h)
        y = sec(c, y, "Patterns observed over 30 days", BLUE2)
        drect(c, LM, y-pat_h, CW, pat_h, fill=LIGHT_BLUE)
        for i, (lbl, txt) in enumerate(pat_items):
            dlbl(c, LM+10, y-5-i*RH, f"{lbl}:", 9, bold=True, color=BLUE2)
            dlbl(c, LM+82, y-5-i*RH, txt,        9, color=DARK_TEXT)
        y -= pat_h + GAP

    # ── Triggers ──
    triggers = data.get("triggers", [])
    if triggers:
        trig_h = len(triggers) * RH
        y = ensure(y, SH + trig_h)
        y = sec(c, y, "Triggers identified by carers", AMBER)
        drect(c, LM, y-trig_h, CW, trig_h, fill=LIGHT_AMBR)
        for i, trig in enumerate(triggers):
            drect(c, LM+10, y-9-i*RH, 7, 7, fill=AMBER)
            dlbl(c,  LM+21, y-4-i*RH, trig, 9, color=DARK_TEXT)
        y -= trig_h + GAP

    # ── Carer responses ──
    responses = data.get("carer_responses", [])
    if responses:
        resp_h = len(responses) * RH
        y = ensure(y, SH + resp_h)
        y = sec(c, y, "Strategies carers have already been trying", GREEN)
        drect(c, LM, y-resp_h, CW, resp_h, fill=LIGHT_GRN)
        for i, resp in enumerate(responses):
            drect(c, LM+10, y-9-i*RH, 7, 7, fill=GREEN)
            dlbl(c,  LM+21, y-4-i*RH, resp, 9, color=DARK_TEXT)
        y -= resp_h + GAP

    # ── Carer concerns ──
    concerns = data.get("carer_concerns", "")
    if concerns:
        con_lines = wrap_text(concerns, "Helvetica-Oblique", 9.5, CW-24)
        con_h = len(con_lines) * 15 + 10
        y = ensure(y, SH + con_h)
        y = sec(c, y, "Carer concerns — summary", DEEP_BLUE)
        drect(c, LM, y-con_h, CW, con_h, fill=LIGHT_BLUE)
        for li, ln in enumerate(con_lines):
            dlbl(c, LM+10, y-13-li*15, ln, 9.5, italic=True, color=DARK_TEXT)
        y -= con_h + GAP

    # ── Notes for practitioner ──
    notes = data.get("notes_for_practitioner", [])
    if notes:
        notes_h = len(notes) * RH
        y = ensure(y, SH + notes_h)
        y = sec(c, y, "Notes for PBS Practitioner — follow up", RED)
        drect(c, LM, y-notes_h, CW, notes_h, fill=LIGHT_RED)
        for i, note in enumerate(notes):
            drect(c, LM+10, y-9-i*RH, 7, 7, fill=RED)
            dlbl(c,  LM+21, y-4-i*RH, note, 9, bold=True, color=DARK_TEXT)
        y -= notes_h + GAP

    draw_footer()
    c.save(); buf.seek(0); return buf


# ══════════════════════════════════════════════════════════════════════════════
# CLAUDE — PBSP EXTRACTION (Tab 2)
# ══════════════════════════════════════════════════════════════════════════════

EXTRACTION_PROMPT = """
You are a Positive Behaviour Support (PBS) assistant.
Read the PBSP text below and extract key information into a JSON object.

Rules:
- Keep each bullet point under 90 characters — written for support workers to scan quickly
- Use plain language (no clinical jargon where possible)
- Focus on what is most operationally useful during a shift
- If a field cannot be found, use an empty string or empty list

Return ONLY valid JSON — no commentary, no markdown fences.

JSON structure:
{
  "name": "full name",
  "preferred": "preferred/short name",
  "pronouns": "pronouns",
  "age_info": "Age XX  |  [Diagnosis]  |  [Organisation / Location]",
  "about": ["up to 5 key points a support worker must know about this person"],
  "warning_signs": ["up to 5 observable early warning signs that behaviour is building"],
  "triggers": ["up to 6 known setting events and immediate triggers"],
  "proactive": ["up to 5 proactive strategies — things to DO to prevent behaviour"],
  "reactive": ["up to 5 reactive strategies — what to DO when behaviour occurs"],
  "do_not": ["up to 4 things NOT to do — known escalators"],
  "behaviours": [
    {
      "label": "short behaviour name (e.g. Verbal Outbursts)",
      "descriptors": ["up to 5 observable descriptors of this behaviour"]
    }
  ],
  "setting_events_checklist": ["up to 6 setting events as short checkbox labels"],
  "antecedents_checklist": ["up to 8 common antecedents as short checkbox labels"],
  "staff_responses_checklist": ["up to 7 common staff responses as short checkbox labels"],
  "review_date": "review year or date",
  "practitioner": "name and title",
  "contact": "email or contact"
}

PBSP TEXT:
"""

def extract_client_data(pbsp_text: str, api_key: str) -> dict:
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model="claude-opus-4-5", max_tokens=2000,
        messages=[{"role": "user", "content": EXTRACTION_PROMPT + pbsp_text[:40000]}]
    )
    raw = msg.content[0].text.strip()
    if raw.startswith("```"): raw = "\n".join(raw.split("\n")[1:])
    if raw.endswith("```"):   raw = "\n".join(raw.split("\n")[:-1])
    return json.loads(raw)


# ══════════════════════════════════════════════════════════════════════════════
# CLAUDE — STRATEGY RECOMMENDER (Tab 3)
# ══════════════════════════════════════════════════════════════════════════════

STRATEGY_PROMPT = """\
You are an experienced Positive Behaviour Support (PBS) Practitioner.
Based on the client profile and behaviours described, recommend practical PBS strategies for support workers.

{library_instruction}

Return ONLY valid JSON — no commentary, no markdown fences.

JSON structure:
{{
  "client_summary": "2-3 sentence summary of this person's likely support needs and overall PBS approach",
  "general_strategies": ["Up to 5 general/environmental strategies that apply across all behaviours"],
  "behaviours": [
    {{
      "behaviour": "behaviour name from input",
      "likely_function": "Primary function (Escape/Avoidance | Access | Sensory | Attention) — one sentence rationale",
      "proactive": ["Up to 5 proactive/antecedent strategies to prevent this behaviour"],
      "reactive": ["Up to 5 reactive de-escalation strategies for when this behaviour occurs"],
      "teach_instead": ["Up to 3 replacement skills or alternative communication strategies to build"],
      "avoid": ["Up to 3 specific things NOT to do — common mistakes that escalate this behaviour"]
    }}
  ]
}}

Rules:
- Keep each item under 90 characters — written for support workers to act on quickly
- Be specific to the triggers and context described — avoid generic advice
- Base strategy recommendations on the likely function of each behaviour
- Use plain language, not clinical jargon
- Replacement skills must serve the same function as the behaviour (functionally equivalent)

CLIENT PROFILE:
{profile}

BEHAVIOURS OF CONCERN:
{behaviours}
{library_section}"""

def recommend_strategies(client_info: dict, behaviours: list, api_key: str,
                          library_text: str = None) -> dict:
    profile_text = (
        f"Name: {client_info['name']}\n"
        f"Age: {client_info['age']}\n"
        f"Diagnosis/Condition: {client_info['diagnosis']}\n"
        f"Communication level: {client_info['comms']}\n"
        f"Additional context: {client_info['other']}"
    )
    behaviours_text = ""
    for i, b in enumerate(behaviours, 1):
        behaviours_text += (
            f"\nBehaviour {i}: {b['name']}\n"
            f"  What it looks like: {b['description']}\n"
            f"  Known triggers / when it occurs: {b['triggers']}\n"
        )
    if library_text:
        lib_instruction = (
            "IMPORTANT: Your primary task is to SELECT strategies FROM THE STRATEGY LIBRARY "
            "provided at the end of this prompt. Quote or closely paraphrase library strategies. "
            "Where the library has no relevant strategy for a specific need, you may suggest an "
            "evidence-based alternative and append '(not in library)' to that item."
        )
        lib_section = f"\nSTRATEGY LIBRARY:\n{library_text[:30000]}"
    else:
        lib_instruction = (
            "Generate evidence-based PBS strategies based on the likely function of each behaviour."
        )
        lib_section = ""
    prompt = STRATEGY_PROMPT.format(
        profile=profile_text, behaviours=behaviours_text,
        library_instruction=lib_instruction, library_section=lib_section,
    )
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model="claude-opus-4-5", max_tokens=3000,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = msg.content[0].text.strip()
    if raw.startswith("```"): raw = "\n".join(raw.split("\n")[1:])
    if raw.endswith("```"):   raw = "\n".join(raw.split("\n")[:-1])
    return json.loads(raw)


def pbsp_to_sr_format(data: dict) -> tuple:
    """Convert extract_client_data output → (client_info dict, behaviours list) for Tab 3."""
    age_info = data.get("age_info", "")
    parts    = [p.strip() for p in age_info.split("|")]
    age      = parts[0].replace("Age", "").strip() if parts else "not specified"
    diagnosis = parts[1] if len(parts) > 1 else "not specified"
    client_info = {
        "name":      data.get("name") or data.get("preferred") or "Unknown",
        "age":       age or "not specified",
        "diagnosis": diagnosis or "not specified",
        "comms":     "refer to PBSP",
        "other":     "; ".join(data.get("about", [])) or "none provided",
    }
    triggers_str = "; ".join(data.get("triggers", []))
    behaviours = [
        {
            "name":        b.get("label", "Behaviour"),
            "description": "; ".join(b.get("descriptors", [])),
            "triggers":    triggers_str,
        }
        for b in data.get("behaviours", [])
    ]
    return client_info, behaviours


# ── Free-text behaviour description → strategies ──────────────────────────────

FREETEXT_STRATEGY_PROMPT = """\
You are an experienced Positive Behaviour Support (PBS) Practitioner.
A practitioner has described a client and their behaviours in plain language below.

Your tasks:
1. Read the description and identify each distinct behaviour of concern
2. Assign each a clear clinical label (e.g. "Physical aggression", "Self-injurious behaviour",
   "Property destruction", "Verbal aggression", "Elopement")
3. Infer the likely triggers and context from what is described
4. Determine the likely function of each behaviour
5. Recommend practical PBS strategies

{library_instruction}

Return ONLY valid JSON — no commentary, no markdown fences.

JSON structure:
{{
  "client_summary": "2-3 sentence clinical summary based on the description",
  "general_strategies": ["Up to 5 general/environmental strategies across all behaviours"],
  "behaviours": [
    {{
      "behaviour": "Clinical label you have assigned",
      "likely_function": "Primary function (Escape/Avoidance | Access | Sensory | Attention) — one sentence rationale drawn from the description",
      "proactive": ["Up to 5 proactive strategies to prevent this behaviour"],
      "reactive": ["Up to 5 reactive de-escalation strategies when this behaviour occurs"],
      "teach_instead": ["Up to 3 replacement skills or communication strategies to build"],
      "avoid": ["Up to 3 things NOT to do — common mistakes that escalate this behaviour"]
    }}
  ]
}}

Rules:
- Keep each item under 90 characters — written for support workers to act on quickly
- Use the description to infer triggers and context — be specific, not generic
- Assign clinically appropriate behaviour labels
- Use plain language in strategy recommendations
- Replacement skills must serve the same function as the behaviour

CLIENT PROFILE:
{profile}

PRACTITIONER'S DESCRIPTION:
{freetext}
{library_section}"""


def recommend_from_freetext(client_info: dict, freetext: str, api_key: str,
                              library_text: str = None) -> dict:
    profile_text = (
        f"Name: {client_info['name']}\n"
        f"Age: {client_info['age']}\n"
        f"Diagnosis/Condition: {client_info['diagnosis']}\n"
        f"Communication level: {client_info['comms']}\n"
        f"Additional context: {client_info['other']}"
    )
    if library_text:
        lib_instruction = (
            "IMPORTANT: Select strategies FROM THE STRATEGY LIBRARY provided at the end of this "
            "prompt. Quote or closely paraphrase library strategies. Where the library has no "
            "relevant strategy, suggest an evidence-based alternative and note '(not in library)'."
        )
        lib_section = f"\nSTRATEGY LIBRARY:\n{library_text[:30000]}"
    else:
        lib_instruction = "Generate evidence-based PBS strategies based on function and context."
        lib_section = ""
    prompt = FREETEXT_STRATEGY_PROMPT.format(
        profile=profile_text, freetext=freetext,
        library_instruction=lib_instruction, library_section=lib_section,
    )
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model="claude-opus-4-5", max_tokens=3000,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = msg.content[0].text.strip()
    if raw.startswith("```"): raw = "\n".join(raw.split("\n")[1:])
    if raw.endswith("```"):   raw = "\n".join(raw.split("\n")[:-1])
    return json.loads(raw)


# ══════════════════════════════════════════════════════════════════════════════
# PDF HELPERS (shared)
# ══════════════════════════════════════════════════════════════════════════════

def drect(c, x, y, w, h, fill=None, stroke=None, lw=0.75):
    c.saveState()
    if fill:   c.setFillColor(fill)
    if stroke: c.setStrokeColor(stroke); c.setLineWidth(lw)
    c.rect(x, y, w, h, fill=1 if fill else 0, stroke=1 if stroke else 0)
    c.restoreState()

def dlbl(c, x, y, s, size, bold=False, italic=False, color=DARK_TEXT, align='left'):
    c.saveState()
    c.setFillColor(color)
    face = ('Helvetica-BoldOblique' if bold and italic else
            'Helvetica-Bold'        if bold           else
            'Helvetica-Oblique'     if italic         else 'Helvetica')
    c.setFont(face, size)
    if align == 'center': c.drawCentredString(x, y, s)
    elif align == 'right': c.drawRightString(x, y, s)
    else: c.drawString(x, y, s)
    c.restoreState()

def sec(c, y, title, bg, ht=20):
    drect(c, LM, y - ht, CW, ht, fill=bg)
    dlbl(c, LM + 8, y - ht + 6, title, 11, bold=True, color=white)
    return y - ht

def checkbox(c, x, y, size=8):
    c.saveState(); c.setStrokeColor(MID_GREY); c.setLineWidth(0.75)
    c.rect(x, y, size, size, fill=0, stroke=1); c.restoreState()

def write_line(c, x, y, w):
    c.saveState(); c.setStrokeColor(MID_GREY); c.setLineWidth(0.5)
    c.line(x, y, x + w, y); c.restoreState()

def cb_item(c, x, y, text, size=9, color=DARK_TEXT, bold=False):
    checkbox(c, x, y, 8)
    dlbl(c, x + 13, y + 1, text, size, bold=bold, color=color)

def wrap_text(text: str, font_name: str, font_size: float, max_w: float) -> list:
    words = text.split()
    lines, line = [], ""
    for word in words:
        test = (line + " " + word).strip()
        if stringWidth(test, font_name, font_size) <= max_w:
            line = test
        else:
            if line: lines.append(line)
            line = word
    if line: lines.append(line)
    return lines or [""]


# ══════════════════════════════════════════════════════════════════════════════
# SUPPORT REFERENCE CARD PDF (Tab 2)
# ══════════════════════════════════════════════════════════════════════════════

def generate_support_card(d: dict) -> BytesIO:
    buf = BytesIO()
    c   = rl_canvas.Canvas(buf, pagesize=A4)
    RH, SH, GAP = 22, 20, 7

    ABOUT_H = SH + len(d["about"])         * RH
    WARN_H  = SH + len(d["warning_signs"]) * RH
    SIDE_H  = SH + max(len(d["triggers"]), len(d["proactive"])) * RH
    REACT_H = SH + len(d["reactive"])      * RH
    DONT_H  = SH + len(d["do_not"])        * RH
    HDR_H, FTR_H, BOT = 72, 38, 12
    y = H

    drect(c, 0, y-HDR_H, W, HDR_H, fill=DEEP_BLUE)
    drect(c, 0, y-HDR_H, 8, HDR_H, fill=TEAL)
    px, py = W-LM-58, y-HDR_H+6
    drect(c, px, py, 52, 58, fill=BLUE2, stroke=HexColor('#C8E6F5'), lw=1)
    dlbl(c, px+26, py+28, "PHOTO",      7, color=HexColor('#C8E6F5'), align='center')
    dlbl(c, px+26, py+18, "(optional)", 7, italic=True, color=HexColor('#8FC8DC'), align='center')
    dlbl(c, LM+6, y-22, d["name"],           22, bold=True, color=white)
    dlbl(c, LM+6, y-40, d.get("pronouns",""), 11, color=HexColor('#C8E6F5'))
    dlbl(c, LM+6, y-56, d.get("age_info",""), 10, color=HexColor('#8FC8DC'))
    dlbl(c, LM+6, y-70, "Positive Behaviour Support  —  Behaviour Support Reference Card",
         8, italic=True, color=HexColor('#5B9FC0'))
    y -= HDR_H + GAP

    y = sec(c, y, f"About {d['preferred']}", TEAL)
    drect(c, LM, y-(ABOUT_H-SH), CW, ABOUT_H-SH, fill=LIGHT_TEAL)
    for i, item in enumerate(d["about"]):
        drect(c, LM+10, y-12-i*RH, 7, 7, fill=TEAL)
        dlbl(c,  LM+21, y- 7-i*RH, item, 9.5, color=DARK_TEXT)
    y -= (ABOUT_H-SH) + GAP

    y = sec(c, y, "Early warning signs — watch for these", AMBER)
    drect(c, LM, y-(WARN_H-SH), CW, WARN_H-SH, fill=LIGHT_AMBR)
    for i, item in enumerate(d["warning_signs"]):
        drect(c, LM+10, y-12-i*RH, 7, 7, fill=AMBER)
        dlbl(c,  LM+21, y- 7-i*RH, item, 9.5, color=DARK_TEXT)
    y -= (WARN_H-SH) + GAP

    half, cont = (CW-5)/2, SIDE_H-SH
    drect(c, LM,       y-SH, half, SH, fill=BLUE2)
    dlbl(c, LM+8,      y-13, "Known triggers",       11, bold=True, color=white)
    drect(c, LM,       y-SIDE_H, half, cont, fill=LIGHT_BLUE)
    for i, item in enumerate(d["triggers"]):
        drect(c, LM+10, y-SH-10-i*RH, 7, 7, fill=BLUE2)
        dlbl(c,  LM+21, y-SH- 5-i*RH, item, 9.5, color=DARK_TEXT)
    rx = LM+half+5
    drect(c, rx,       y-SH, half, SH, fill=GREEN)
    dlbl(c, rx+8,      y-13, "Proactive strategies",  11, bold=True, color=white)
    drect(c, rx,       y-SIDE_H, half, cont, fill=LIGHT_GRN)
    for i, item in enumerate(d["proactive"]):
        drect(c, rx+10, y-SH-10-i*RH, 7, 7, fill=GREEN)
        dlbl(c,  rx+21, y-SH- 5-i*RH, item, 9.5, color=DARK_TEXT)
    y -= SIDE_H + GAP

    y = sec(c, y, "When behaviour occurs — do this", TEAL)
    drect(c, LM, y-(REACT_H-SH), CW, REACT_H-SH, fill=LIGHT_TEAL)
    for i, item in enumerate(d["reactive"]):
        drect(c, LM+10, y-12-i*RH, 7, 7, fill=TEAL)
        dlbl(c,  LM+21, y- 7-i*RH, item, 9.5, color=DARK_TEXT)
    y -= (REACT_H-SH) + GAP

    y = sec(c, y, f"DO NOT — things that escalate behaviour for {d['preferred']}", RED)
    drect(c, LM, y-(DONT_H-SH), CW, DONT_H-SH, fill=LIGHT_RED)
    for i, item in enumerate(d["do_not"]):
        drect(c, LM+10, y-12-i*RH, 7, 7, fill=RED)
        dlbl(c,  LM+21, y- 7-i*RH, item, 9.5, bold=True, color=RED)
    y -= (DONT_H-SH) + GAP

    drect(c, 0, BOT, W, FTR_H, fill=DEEP_BLUE)
    drect(c, 0, BOT, 8, FTR_H, fill=TEAL)
    ft = BOT + FTR_H
    dlbl(c, LM+6, ft-16,
         f"Plan review: {d.get('review_date','')}  |  {d.get('practitioner','')}", 9, bold=True, color=white)
    dlbl(c, LM+6, ft-29, d.get("contact",""), 9, italic=True, color=HexColor('#C8E6F5'))
    dlbl(c, W-LM, ft-29, "CONFIDENTIAL — handle in line with your privacy policy",
         8, italic=True, color=HexColor('#5B9FC0'), align='right')
    c.save(); buf.seek(0); return buf


# ══════════════════════════════════════════════════════════════════════════════
# ABC RECORDING FORM PDF (Tab 2)
# ══════════════════════════════════════════════════════════════════════════════

def generate_abc_form(d: dict) -> BytesIO:
    buf = BytesIO()
    c   = rl_canvas.Canvas(buf, pagesize=A4)
    y   = H
    half_cw = (CW - 10) / 2

    drect(c, 0, y-58, W, 58, fill=DEEP_BLUE)
    drect(c, 0, y-58, 7, 58, fill=TEAL)
    dlbl(c, LM+4, y-20, "ABC Behaviour Recording Form", 16, bold=True, color=white)
    dlbl(c, LM+4, y-36,
         f"{d['name']}  |  {d.get('pronouns','')}  |  {d.get('age_info','')}", 9, color=HexColor('#C8E6F5'))
    dlbl(c, LM+4, y-52, "Complete as soon as safely possible after any behaviour of concern",
         8, italic=True, color=HexColor('#8FC8DC'))
    dlbl(c, W-LM, y-52, "CONFIDENTIAL", 8, bold=True, color=HexColor('#5B9FC0'), align='right')
    y -= 63

    drect(c, LM, y-22, CW, 22, fill=LIGHT_BLUE)
    dlbl(c, LM+6,   y-14, "Date:",           9, bold=True, color=BLUE2)
    write_line(c, LM+32,  y-15, 68)
    dlbl(c, LM+110, y-14, "Time:",           9, bold=True, color=BLUE2)
    write_line(c, LM+134, y-15, 30)
    dlbl(c, LM+170, y-14, "to",              9, color=MED_TEXT)
    write_line(c, LM+181, y-15, 30)
    dlbl(c, LM+220, y-14, "Support worker:", 9, bold=True, color=BLUE2)
    write_line(c, LM+297, y-15, 90)
    dlbl(c, LM+398, y-14, "Shift:",          9, bold=True, color=BLUE2)
    for si, sh in enumerate(["Day","Aft","Night"]):
        sx = LM+425+si*38; checkbox(c, sx, y-16); dlbl(c, sx+11, y-14, sh, 8.5, color=DARK_TEXT)
    y -= 27

    drect(c, LM, y-20, 20, 20, fill=AMBER)
    dlbl(c, LM+10, y-13, "S", 11, bold=True, color=white, align='center')
    drect(c, LM+20, y-20, CW-20, 20, fill=LIGHT_AMBR)
    dlbl(c, LM+28, y-13, "Setting events today", 11, bold=True, color=AMBER)
    y -= 20
    se = d.get("setting_events_checklist", [])
    se_h = 10 + ((len(se)+1)//2)*18 + 18
    drect(c, LM, y-se_h, CW, se_h, fill=LIGHT_AMBR)
    for i, item in enumerate(se):
        col, row = i%2, i//2
        cb_item(c, LM+6+col*(half_cw+10), y-14-row*18, item)
    dlbl(c, LM+6, y-se_h+6, "Other:", 9, color=MED_TEXT)
    write_line(c, LM+38, y-se_h+5, CW-44)
    y -= se_h + 5

    drect(c, LM, y-20, 20, 20, fill=BLUE2)
    dlbl(c, LM+10, y-13, "A", 11, bold=True, color=white, align='center')
    drect(c, LM+20, y-20, CW-20, 20, fill=LIGHT_BLUE)
    dlbl(c, LM+28, y-13, "Antecedent — what happened immediately BEFORE?", 11, bold=True, color=BLUE2)
    y -= 20
    ant = d.get("antecedents_checklist", [])
    ant_h = 10 + ((len(ant)+1)//2)*18 + 18
    drect(c, LM, y-ant_h, CW, ant_h, fill=LIGHT_BLUE)
    for i, item in enumerate(ant):
        col, row = i%2, i//2
        cb_item(c, LM+6+col*(half_cw+10), y-14-row*18, item)
    dlbl(c, LM+6, y-ant_h+6, "Other / describe:", 9, color=MED_TEXT)
    write_line(c, LM+96, y-ant_h+5, CW-102)
    y -= ant_h + 5

    drect(c, LM, y-20, 20, 20, fill=TEAL)
    dlbl(c, LM+10, y-13, "B", 11, bold=True, color=white, align='center')
    drect(c, LM+20, y-20, CW-20, 20, fill=LIGHT_TEAL)
    dlbl(c, LM+28, y-13, "Behaviour — what did you observe? (facts only, no interpretations)",
         11, bold=True, color=TEAL)
    y -= 20
    behs = d.get("behaviours", [])
    n_beh = len(behs)
    beh_col_w = (CW-8) / max(n_beh, 1)
    max_desc = max((len(b.get("descriptors",[])) for b in behs), default=3)
    beh_h = 18 + max_desc*16 + 24
    drect(c, LM, y-beh_h, CW, beh_h, fill=LIGHT_TEAL)
    beh_colors = [TEAL, BLUE2, GREEN, AMBER]
    for bi, beh in enumerate(behs):
        bx = LM+4+bi*(beh_col_w+4); bw = beh_col_w-4; bc = beh_colors[bi%len(beh_colors)]
        drect(c, bx, y-18, bw, 18, fill=bc)
        dlbl(c, bx+5, y-12, beh.get("label",""), 9, bold=True, color=white)
        for di, desc in enumerate(beh.get("descriptors",[])):
            cb_item(c, bx+2, y-26-di*16, desc, size=8.5)
    ir = y-beh_h+22
    drect(c, LM+4, ir, CW-8, 16, fill=HexColor('#DFF0EE'))
    dlbl(c, LM+8, ir+5, "Intensity:", 9, bold=True, color=TEAL)
    for ii, lvl in enumerate(["Mild","Moderate","Severe"]):
        sx = LM+54+ii*58; checkbox(c, sx, ir+4); dlbl(c, sx+11, ir+5, lvl, 9, color=DARK_TEXT)
    dlbl(c, LM+232, ir+5, "Duration:", 9, bold=True, color=TEAL)
    write_line(c, LM+272, ir+4, 28)
    dlbl(c, LM+305, ir+5, "mins", 9, color=MED_TEXT)
    dlbl(c, LM+345, ir+5, "Location:", 9, bold=True, color=TEAL)
    write_line(c, LM+385, ir+4, CW-203)
    dlbl(c, LM+6, y-beh_h+8, "Describe what you saw (objective language):", 9, italic=True, color=MED_TEXT)
    write_line(c, LM+6, y-beh_h+2, CW-12)
    y -= beh_h + 5

    drect(c, LM, y-20, 20, 20, fill=GREEN)
    dlbl(c, LM+10, y-13, "C", 11, bold=True, color=white, align='center')
    drect(c, LM+20, y-20, CW-20, 20, fill=LIGHT_GRN)
    dlbl(c, LM+28, y-13, "Consequence — what happened after? What did you do?",
         11, bold=True, color=GREEN)
    y -= 20
    resp = d.get("staff_responses_checklist", [])
    resp_h = 10 + ((len(resp)+1)//2)*17 + 52
    drect(c, LM, y-resp_h, CW, resp_h, fill=LIGHT_GRN)
    dlbl(c, LM+6, y-10, "What did you do?", 9, bold=True, color=GREEN)
    for i, item in enumerate(resp):
        col, row = i%2, i//2
        cb_item(c, LM+6+col*(half_cw+10), y-22-row*17, item)
    ry = y-resp_h+48
    dlbl(c, LM+6, ry, "How did they respond?", 9, bold=True, color=GREEN)
    for ri, r in enumerate(["De-escalated < 5 mins","De-escalated < 30 mins",
                             "Continued > 30 mins","Escalated further"]):
        rx2 = LM+6+ri*134; checkbox(c, rx2, ry-14); dlbl(c, rx2+11, ry-13, r, 8.5, color=DARK_TEXT)
    dlbl(c, LM+6, ry-28, "Other:", 9, color=MED_TEXT)
    write_line(c, LM+38, ry-29, CW-44)
    y -= resp_h + 5

    drect(c, LM, y-18, CW, 18, fill=LIGHT_BLUE)
    dlbl(c, LM+6, y-12, "Additional notes:", 10, bold=True, color=BLUE2)
    y -= 18
    drect(c, LM, y-44, CW, 44, fill=HexColor('#FAFAFA'), stroke=MID_GREY, lw=0.5)
    for li in range(3): write_line(c, LM+6, y-16-li*13, CW-12)
    y -= 48

    drect(c, LM, y-28, CW, 28, fill=LIGHT_TEAL)
    for i, ch in enumerate(["Incident report completed","Handed over to next shift",
                             "Staff debrief completed","Behaviour data entered"]):
        ix = LM+6+i*(CW/4); checkbox(c, ix, y-18); dlbl(c, ix+11, y-16, ch, 8.5, color=DARK_TEXT)
    y -= 32

    drect(c, LM, y-24, CW, 24, fill=DEEP_BLUE)
    dlbl(c, LM+6,   y-10, "Signature:",   9, bold=True, color=white)
    write_line(c, LM+55,  y-11, 110)
    dlbl(c, LM+175, y-10, "Print name:",  9, bold=True, color=white)
    write_line(c, LM+227, y-11, 110)
    dlbl(c, LM+347, y-10, "Date / Time:", 9, bold=True, color=white)
    write_line(c, LM+407, y-11, CW-230)
    dlbl(c, W/2, y-20, f"{d.get('practitioner','')}  |  {d.get('contact','')}",
         7, italic=True, color=HexColor('#8FC8DC'), align='center')

    c.save(); buf.seek(0); return buf


# ══════════════════════════════════════════════════════════════════════════════
# STRATEGY REPORT PDF (Tab 3)
# ══════════════════════════════════════════════════════════════════════════════

def generate_strategy_report(result: dict, client_name: str,
                              practitioner: str = "", contact: str = "") -> BytesIO:
    buf = BytesIO()
    c   = rl_canvas.Canvas(buf, pagesize=A4)

    BOT, FTR_H, HDR_H  = 12, 38, 68
    MINI_HDR_H          = 28
    MIN_Y               = BOT + FTR_H + 20
    RH, SH, GAP         = 17, 20, 8
    page_num            = [1]

    def draw_footer():
        drect(c, 0, BOT, W, FTR_H, fill=DEEP_BLUE)
        drect(c, 0, BOT, 8, FTR_H, fill=TEAL)
        ft = BOT + FTR_H
        dlbl(c, LM+6, ft-14, practitioner or "Positive Behaviour Support", 9, bold=True, color=white)
        if contact:
            dlbl(c, LM+6, ft-28, contact, 9, italic=True, color=HexColor('#C8E6F5'))
        dlbl(c, W-LM, ft-14, f"Page {page_num[0]}", 9, color=HexColor('#8FC8DC'), align='right')
        dlbl(c, W-LM, ft-28, "CONFIDENTIAL — handle in line with your privacy policy",
             8, italic=True, color=HexColor('#5B9FC0'), align='right')

    def new_page():
        draw_footer(); c.showPage(); page_num[0] += 1
        drect(c, 0, H-MINI_HDR_H, W, MINI_HDR_H, fill=DEEP_BLUE)
        drect(c, 0, H-MINI_HDR_H, 8, MINI_HDR_H, fill=TEAL)
        dlbl(c, LM+6, H-MINI_HDR_H+9,
             f"Behaviour Strategy Recommendations — {client_name}", 10, bold=True, color=white)
        return H - MINI_HDR_H - GAP

    def ensure(y, needed):
        return new_page() if y - needed < MIN_Y else y

    y = H
    drect(c, 0, y-HDR_H, W, HDR_H, fill=DEEP_BLUE)
    drect(c, 0, y-HDR_H, 8, HDR_H, fill=TEAL)
    dlbl(c, LM+6, y-22, "Behaviour Strategy Recommendations", 17, bold=True, color=white)
    dlbl(c, LM+6, y-42, client_name, 13, color=HexColor('#C8E6F5'))
    dlbl(c, LM+6, y-58, "Positive Behaviour Support — Strategy Guide for Support Workers",
         9, italic=True, color=HexColor('#5B9FC0'))
    dlbl(c, W-LM, y-58, f"Generated: {date.today().strftime('%d %B %Y')}",
         9, italic=True, color=HexColor('#5B9FC0'), align='right')
    y -= HDR_H + GAP

    summary = result.get("client_summary", "")
    if summary:
        lines = wrap_text(summary, "Helvetica-Oblique", 9.5, CW - 24)
        sum_h = len(lines) * 15 + 10
        y = ensure(y, sum_h + SH)
        y = sec(c, y, "Clinical Summary", TEAL)
        drect(c, LM, y-sum_h, CW, sum_h, fill=LIGHT_TEAL)
        for li, ln in enumerate(lines):
            dlbl(c, LM+10, y-13-li*15, ln, 9.5, italic=True, color=DARK_TEXT)
        y -= sum_h + GAP

    gen = result.get("general_strategies", [])
    if gen:
        gen_h = len(gen) * RH
        y = ensure(y, gen_h + SH)
        y = sec(c, y, "General Strategies — apply across all behaviours", BLUE2)
        drect(c, LM, y-gen_h, CW, gen_h, fill=LIGHT_BLUE)
        for i, item in enumerate(gen):
            drect(c, LM+10, y-9-i*RH, 7, 7, fill=BLUE2)
            dlbl(c,  LM+21, y-4-i*RH, item, 9.5, color=DARK_TEXT)
        y -= gen_h + GAP

    for beh in result.get("behaviours", []):
        b_name  = beh.get("behaviour", "Behaviour")
        b_func  = beh.get("likely_function", "")
        pros    = beh.get("proactive", [])
        reacts  = beh.get("reactive", [])
        teaches = beh.get("teach_instead", [])
        avoids  = beh.get("avoid", [])

        y = ensure(y, 26)
        drect(c, LM, y-26, CW, 26, fill=DEEP_BLUE)
        drect(c, LM, y-26,  6, 26, fill=TEAL)
        dlbl(c, LM+14, y-17, f"Behaviour: {b_name}", 12, bold=True, color=white)
        y -= 26

        y = ensure(y, 22)
        drect(c, LM, y-22, CW, 22, fill=LIGHT_AMBR)
        dlbl(c, LM+8,  y-14, "Likely function:", 9.5, bold=True, color=AMBER)
        func_lines = wrap_text(b_func, "Helvetica-Oblique", 9.5, CW - 112)
        dlbl(c, LM+106, y-14, func_lines[0] if func_lines else "", 9.5, italic=True, color=DARK_TEXT)
        y -= 26

        def draw_section(items, title, bg, light_bg, dot_col):
            nonlocal y
            if not items: return
            h = len(items) * RH
            y = ensure(y, h + SH)
            y = sec(c, y, title, bg)
            drect(c, LM, y-h, CW, h, fill=light_bg)
            for i, item in enumerate(items):
                drect(c, LM+10, y-9-i*RH, 7, 7, fill=dot_col)
                dlbl(c,  LM+21, y-4-i*RH, item, 9.5,
                     bold=(dot_col == RED), color=(dot_col if dot_col == RED else DARK_TEXT))
            y -= h + 4

        draw_section(pros,    "Proactive strategies — prevent this behaviour",         GREEN, LIGHT_GRN,  GREEN)
        draw_section(reacts,  "Reactive strategies — when this behaviour occurs",       TEAL,  LIGHT_TEAL, TEAL)
        draw_section(teaches, "Skills to build — teach as a replacement behaviour",     BLUE2, LIGHT_BLUE, BLUE2)
        draw_section(avoids,  "DO NOT — avoid these with this behaviour",               RED,   LIGHT_RED,  RED)
        y -= GAP

    draw_footer()
    c.save(); buf.seek(0); return buf


# ══════════════════════════════════════════════════════════════════════════════
# NDIS BSP GENERATION (Tab 4)
# ══════════════════════════════════════════════════════════════════════════════

def format_data_for_bsp(t0_data: dict = None, t1_data: dict = None) -> str:
    """Combine available tab data into structured text for BSP generation."""
    parts = []
    if t1_data:
        parts.append(f"CLIENT NAME: {t1_data.get('name', 'Unknown')}")
        parts.append(f"PROFILE: {t1_data.get('age_info', '')}")
        if t1_data.get('pronouns'):
            parts.append(f"PRONOUNS: {t1_data['pronouns']}")
        if t1_data.get('about'):
            parts.append("ABOUT THIS PERSON: " + "; ".join(t1_data['about']))
        if t1_data.get('warning_signs'):
            parts.append("WARNING SIGNS: " + "; ".join(t1_data['warning_signs']))
        if t1_data.get('triggers'):
            parts.append("KNOWN TRIGGERS: " + "; ".join(t1_data['triggers']))
        if t1_data.get('proactive'):
            parts.append("CURRENT PROACTIVE STRATEGIES: " + "; ".join(t1_data['proactive']))
        if t1_data.get('reactive'):
            parts.append("CURRENT REACTIVE STRATEGIES: " + "; ".join(t1_data['reactive']))
        if t1_data.get('do_not'):
            parts.append("DO NOT DO: " + "; ".join(t1_data['do_not']))
        for b in t1_data.get('behaviours', []):
            parts.append(f"BEHAVIOUR — {b.get('label','')}: {'; '.join(b.get('descriptors', []))}")
    if t0_data:
        parts.append(f"\n30-DAY RECORDING (period: {t0_data.get('recording_period','unknown')}):")
        parts.append(f"Total incidents: {t0_data.get('total_incidents','unknown')}")
        if t0_data.get('carer_concerns'):
            parts.append(f"CARER CONCERNS: {t0_data['carer_concerns']}")
        patterns = t0_data.get('patterns', {})
        all_p = (patterns.get('time_of_day',[]) + patterns.get('day_of_week',[]) +
                 patterns.get('settings',[]) + patterns.get('other',[]))
        if all_p:
            parts.append("PATTERNS: " + "; ".join(all_p))
        if t0_data.get('triggers'):
            parts.append("CARER-RECORDED TRIGGERS: " + "; ".join(t0_data['triggers']))
        if t0_data.get('carer_responses'):
            parts.append("STRATEGIES TRIED: " + "; ".join(t0_data['carer_responses']))
        for b in t0_data.get('behaviours', []):
            parts.append(f"RECORDED BEHAVIOUR — {b.get('label','')}: {b.get('frequency','?')} incidents. "
                         f"{b.get('carer_description','')}")
        if t0_data.get('notes_for_practitioner'):
            parts.append("CLINICAL NOTES: " + "; ".join(t0_data['notes_for_practitioner']))
    return '\n'.join(parts) if parts else "No client data available."


NDIS_BSP_PROMPT = """\
You are an experienced NDIS Behaviour Support Practitioner writing a compliant Positive Behaviour Support Plan.

Generate a draft {plan_type} Behaviour Support Plan to NDIS Version 3 template standards.

For regulated restrictive practices:
- Justify using least restrictive principles
- Mark authorisation fields as [PRACTITIONER TO COMPLETE]
- Include a fading plan for each RP

Return ONLY valid JSON — no commentary, no markdown fences.

JSON structure:
{{
  "about_client": "2-3 paragraphs — strengths, what matters to the person, daily life, communication style",
  "communication": "Detailed description of communication abilities and supports needed",
  "behaviours_of_concern": [
    {{
      "name": "behaviour label",
      "description": "Observable, measurable description",
      "frequency": "How often, based on available data",
      "impact": "Impact on participation and quality of life"
    }}
  ],
  "fba_summary": {{
    "setting_events": ["setting event 1"],
    "triggers": ["immediate trigger 1"],
    "functions": ["Behaviour X — likely function: Y with rationale"],
    "hypothesis": "Full hypothesis statement linking antecedents, behaviour and consequences"
  }},
  "goals": [
    {{
      "goal": "SMART, person-centred goal statement",
      "measure": "How progress will be measured",
      "timeframe": "Target timeframe"
    }}
  ],
  "proactive_strategies": ["detailed proactive strategy 1"],
  "reactive_strategies": ["detailed reactive/de-escalation strategy 1"],
  "skill_building": ["replacement skill or communication strategy 1"],
  "crisis_strategy": {{
    "warning_signs": ["observable early warning sign 1"],
    "escalation_signs": ["sign of escalation to crisis 1"],
    "immediate_responses": ["immediate staff action 1"],
    "staff_safety": ["staff safety consideration 1"],
    "environment_management": ["environmental adjustment 1"],
    "post_crisis_support": ["post-crisis support action 1"],
    "when_to_call_emergency": "Clear criteria for calling 000 or emergency services",
    "debrief_process": "Staff debrief and documentation process after a crisis"
  }},
  "restrictive_practices": [
    {{
      "type": "RP type from NDIS framework",
      "description": "Exactly how this practice is implemented in practice",
      "behaviour_addressed": "Which behaviour(s) this RP relates to and why",
      "justification": "Clinical justification — what harm it prevents, why necessary",
      "least_restrictive_evidence": "Evidence that less restrictive alternatives were considered",
      "conditions_of_use": "When and how this RP may be used — triggers, duration limits, who implements",
      "monitoring": ["monitoring requirement 1"],
      "risks": ["identified risk 1"],
      "fading_plan": "Specific, time-bound plan to reduce and eliminate this RP",
      "authorisation_note": "This RP requires NDIS authorisation prior to implementation."
    }}
  ],
  "rp_protocol_notes": "General dignity, rights and monitoring notes for RP implementation",
  "implementation": {{
    "all_staff_responsibilities": ["all support staff responsibility 1"],
    "team_leader_responsibilities": ["team leader responsibility 1"],
    "practitioner_responsibilities": ["behaviour support practitioner responsibility 1"],
    "training_requirements": [
      {{
        "topic": "Training topic",
        "who": "Who requires this training",
        "method": "Delivery method",
        "timeframe": "When it must be completed"
      }}
    ],
    "competency_statement": "What staff must demonstrate before implementing the plan",
    "go_live_requirements": "Requirements to be met before plan implementation",
    "review_schedule": "Routine and trigger-based review schedule"
  }}
}}

PLAN TYPE: {plan_type}
REGULATED RESTRICTIVE PRACTICES TO INCLUDE: {rp_list}

CLIENT DATA:
{client_data}

{extra_context}

RULES:
- Person-centred language throughout — focus on strengths, needs and rights
- Strategies must address the functional hypothesis
- RP sections must reference least restrictive principles and NDIS authorisation
- Crisis strategies must be specific and actionable — not generic
- Australian spelling throughout
- Mark anything requiring clinical verification with [PRACTITIONER TO VERIFY]
"""


def generate_ndis_bsp(client_data_text: str, plan_type: str, rp_types: list,
                       extra_context: str, api_key: str) -> dict:
    rp_list = ', '.join(rp_types) if rp_types else 'None'
    extra_section = f"ADDITIONAL CONTEXT:\n{extra_context}" if extra_context.strip() else ""
    prompt = NDIS_BSP_PROMPT.format(
        plan_type=plan_type, rp_list=rp_list,
        client_data=client_data_text, extra_context=extra_section,
    )
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model="claude-opus-4-5", max_tokens=8000,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = msg.content[0].text.strip()
    if raw.startswith("```"): raw = "\n".join(raw.split("\n")[1:])
    if raw.endswith("```"):   raw = "\n".join(raw.split("\n")[:-1])
    # If JSON is truncated, attempt to close it gracefully
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        # Try to salvage by trimming to last complete top-level key
        last_brace = raw.rfind('"}')
        if last_brace > 0:
            raw = raw[:last_brace + 2] + "\n}}"
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            raise ValueError(
                "The BSP response was too long and could not be parsed. "
                "Try selecting fewer restrictive practices, or use Interim instead of Comprehensive."
            )


def create_bsp_docx(bsp: dict, plan_type: str, client_name: str,
                     practitioner: str = "", contact: str = "") -> BytesIO:
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    NAVY   = RGBColor(0x0D, 0x4F, 0x6E)
    TEALC  = RGBColor(0x1A, 0x9B, 0x8A)
    REDC   = RGBColor(0xC0, 0x39, 0x2B)
    AMBC   = RGBColor(0xD4, 0x70, 0x0A)
    GREYC  = RGBColor(0x4A, 0x4A, 0x4A)

    def h1(text):
        p = doc.add_heading(text, level=1)
        for r in p.runs: r.font.color.rgb = NAVY; r.font.size = Pt(14)
        return p

    def h2(text):
        p = doc.add_heading(text, level=2)
        for r in p.runs: r.font.color.rgb = TEALC; r.font.size = Pt(12)
        return p

    def body(text, italic=False):
        p = doc.add_paragraph(text)
        if italic:
            for r in p.runs: r.italic = True
        return p

    def bul(text):
        try:    return doc.add_paragraph(text, style='List Bullet')
        except: return doc.add_paragraph(f"• {text}")

    def placeholder(field):
        p = doc.add_paragraph()
        r = p.add_run(f"[{field}]")
        r.bold = True; r.font.color.rgb = AMBC
        return p

    def warn(text):
        p = doc.add_paragraph()
        r = p.add_run(f"DRAFT — {text}")
        r.italic = True; r.font.size = Pt(9); r.font.color.rgb = REDC
        return p

    def two_col_table(rows_data):
        t = doc.add_table(rows=len(rows_data), cols=2)
        t.style = 'Table Grid'
        for i, (lbl, val) in enumerate(rows_data):
            t.rows[i].cells[0].text = lbl
            t.rows[i].cells[1].text = val if val else "[PRACTITIONER TO COMPLETE]"
            if t.rows[i].cells[0].paragraphs[0].runs:
                t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        doc.add_paragraph("")
        return t

    # ── Cover ──────────────────────────────────────────────────────────────────
    title = doc.add_heading("Positive Behaviour Support Plan", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs: r.font.color.rgb = NAVY; r.font.size = Pt(18)
    sub = doc.add_paragraph(f"{plan_type.title()} Plan")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs: r.bold = True; r.font.size = Pt(14); r.font.color.rgb = TEALC
    doc.add_paragraph("")
    two_col_table([
        ("Participant name",               client_name),
        ("NDIS number",                    "[NDIS NUMBER]"),
        ("Date of birth",                  "[DATE OF BIRTH]"),
        ("Plan start date",                date.today().strftime('%d %B %Y')),
        ("Plan review date",               "[REVIEW DATE]"),
        ("Behaviour Support Practitioner", practitioner or "[BSP PRACTITIONER NAME]"),
    ])
    warn("AI-generated draft. All content must be reviewed by a qualified practitioner before "
         "implementation or NDIS submission. Regulated restrictive practices require NDIS "
         "authorisation before use.")
    doc.add_page_break()

    # ── S1: About the person ───────────────────────────────────────────────────
    h1(f"1. About {client_name}")
    if bsp.get("about_client"):
        body(bsp["about_client"])
    else:
        placeholder("Describe who this person is — strengths, interests, daily life, what matters to them")
    if bsp.get("communication"):
        h2("Communication")
        body(bsp["communication"])

    # ── S2: Behaviours of concern ──────────────────────────────────────────────
    doc.add_paragraph("")
    h1("2. Behaviours of concern")
    for b in bsp.get("behaviours_of_concern", []):
        h2(b.get("name", "Behaviour"))
        two_col_table([
            ("Description",        b.get("description", "")),
            ("Frequency / rate",   b.get("frequency", "")),
            ("Impact",             b.get("impact", "")),
        ])
    if not bsp.get("behaviours_of_concern"):
        placeholder("List each behaviour of concern with description, frequency and impact")

    # ── S3: FBA Summary ────────────────────────────────────────────────────────
    doc.add_paragraph("")
    h1("3. Functional behaviour assessment summary")
    fba = bsp.get("fba_summary", {})
    h2("Setting events")
    for s in fba.get("setting_events", []): bul(s)
    if not fba.get("setting_events"): placeholder("List setting events from assessment")
    h2("Immediate triggers")
    for t in fba.get("triggers", []): bul(t)
    if not fba.get("triggers"): placeholder("List immediate antecedents/triggers")
    h2("Functions of behaviour")
    for f in fba.get("functions", []): bul(f)
    if not fba.get("functions"): placeholder("Identify likely function for each behaviour")
    h2("Summary hypothesis")
    if fba.get("hypothesis"): body(fba["hypothesis"])
    else: placeholder("Write the functional hypothesis statement")

    # ── S4: Goals ──────────────────────────────────────────────────────────────
    doc.add_paragraph("")
    h1("4. Behaviour support plan goals")
    for i, g in enumerate(bsp.get("goals", []), 1):
        h2(f"Goal {i}")
        two_col_table([
            ("Goal statement",           g.get("goal", "")),
            ("How progress is measured", g.get("measure", "")),
            ("Target timeframe",         g.get("timeframe", "")),
        ])
    if not bsp.get("goals"): placeholder("Add SMART behaviour support goals")

    # ── S5: Strategies ─────────────────────────────────────────────────────────
    doc.add_paragraph("")
    h1("5. Behaviour support strategies")
    h2("5.1  Proactive strategies")
    body("The following proactive strategies should be implemented consistently.", italic=True)
    for s in bsp.get("proactive_strategies", []): bul(s)
    if not bsp.get("proactive_strategies"): placeholder("List proactive strategies")
    h2("5.2  Reactive strategies")
    body("The following strategies should be used when behaviour is observed or escalating.", italic=True)
    for s in bsp.get("reactive_strategies", []): bul(s)
    if not bsp.get("reactive_strategies"): placeholder("List reactive and de-escalation strategies")
    h2("5.3  Skill building and replacement behaviours")
    body("The following skills will be taught to meet the person's needs more appropriately.", italic=True)
    for s in bsp.get("skill_building", []): bul(s)
    if not bsp.get("skill_building"): placeholder("List replacement skills and teaching strategies")

    # ── S6: Crisis strategy ────────────────────────────────────────────────────
    doc.add_page_break()
    h1("6. Crisis and emergency response strategy")
    warn("Crisis strategies must be reviewed and approved by the Behaviour Support Practitioner. "
         "All staff must be trained before implementation.")
    crisis = bsp.get("crisis_strategy", {})
    sections_crisis = [
        ("Early warning signs",           crisis.get("warning_signs", [])),
        ("Signs of escalation to crisis",  crisis.get("escalation_signs", [])),
        ("Immediate staff responses",      crisis.get("immediate_responses", [])),
        ("Staff safety considerations",    crisis.get("staff_safety", [])),
        ("Environmental management",       crisis.get("environment_management", [])),
        ("Post-crisis support",            crisis.get("post_crisis_support", [])),
    ]
    for title_text, items in sections_crisis:
        h2(title_text)
        for item in items: bul(item)
        if not items: placeholder(f"Complete — {title_text.lower()}")
    h2("When to call emergency services (000)")
    if crisis.get("when_to_call_emergency"):
        body(crisis["when_to_call_emergency"])
    else:
        placeholder("Define clear criteria for calling 000")
    h2("Debrief and documentation")
    if crisis.get("debrief_process"):
        body(crisis["debrief_process"])
    else:
        placeholder("Staff debrief and incident documentation process")
    h2("Emergency contacts")
    ec = doc.add_table(rows=5, cols=2)
    ec.style = 'Table Grid'
    for i, (role, detail) in enumerate([
        ("Role", "Contact details"),
        ("On-call manager / team leader", "[NAME / PHONE]"),
        ("Behaviour Support Practitioner", practitioner or "[NAME / PHONE]"),
        ("Emergency services", "000"),
        ("NDIS Quality and Safeguards Commission", "1800 035 544"),
    ]):
        ec.rows[i].cells[0].text = role
        ec.rows[i].cells[1].text = detail
        if i == 0 and ec.rows[i].cells[0].paragraphs[0].runs:
            ec.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            ec.rows[i].cells[1].paragraphs[0].runs[0].bold = True
    doc.add_paragraph("")

    # ── S7: Regulated restrictive practices ────────────────────────────────────
    rps = bsp.get("restrictive_practices", [])
    if rps:
        doc.add_page_break()
        h1("7. Regulated restrictive practices")
        p = doc.add_paragraph()
        r = p.add_run(
            "IMPORTANT: Regulated restrictive practices cannot be implemented without NDIS "
            "authorisation. This section is a DRAFT and requires completion, practitioner review "
            "and formal authorisation before any regulated practice is used."
        )
        r.bold = True; r.font.color.rgb = REDC

        # Register table
        h2("7.1  Regulated restrictive practices register")
        reg = doc.add_table(rows=1, cols=5)
        reg.style = 'Table Grid'
        for i, hdr in enumerate(["RP type", "Behaviour addressed", "Auth. type",
                                   "Auth. date / ref", "Review date"]):
            reg.rows[0].cells[i].text = hdr
            if reg.rows[0].cells[i].paragraphs[0].runs:
                reg.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for rp in rps:
            row = reg.add_row().cells
            row[0].text = rp.get("type", "")
            row[1].text = rp.get("behaviour_addressed", "")
            row[2].text = "[AUTH TYPE]"
            row[3].text = "[DATE / REF]"
            row[4].text = "[REVIEW DATE]"
        doc.add_paragraph("")

        # Individual protocols
        h2("7.2  Regulated restrictive practice protocols")
        for i, rp in enumerate(rps, 1):
            h2(f"Protocol {i}: {rp.get('type', 'Restrictive Practice')}")
            two_col_table([
                ("Type of regulated RP",         rp.get("type", "")),
                ("Behaviour(s) addressed",        rp.get("behaviour_addressed", "")),
                ("Description of practice",       rp.get("description", "")),
                ("Clinical justification",        rp.get("justification", "")),
                ("Least restrictive evidence",    rp.get("least_restrictive_evidence", "")),
                ("Conditions of use",             rp.get("conditions_of_use", "")),
            ])
            h2("Monitoring requirements")
            for m in rp.get("monitoring", []): bul(m)
            if not rp.get("monitoring"): placeholder("List monitoring requirements")
            h2("Identified risks")
            for r_item in rp.get("risks", []): bul(r_item)
            h2("Fading plan")
            if rp.get("fading_plan"): body(rp["fading_plan"])
            else: placeholder("Describe the plan to reduce and eliminate this RP")
            h2("Authorisation details")
            two_col_table([
                ("Authorisation type",            "[NDIS authorisation type]"),
                ("Authorising body / person",     "[AUTHORISING BODY]"),
                ("Authorisation date",            "[DATE]"),
                ("Authorisation reference number","[REFERENCE NUMBER]"),
            ])
        if bsp.get("rp_protocol_notes"):
            h2("General implementation notes")
            body(bsp["rp_protocol_notes"])

    # ── S8: Implementation and training ───────────────────────────────────────
    doc.add_page_break()
    h1("8. Implementation plan and training")
    impl = bsp.get("implementation", {})

    h2("8.1  Roles and responsibilities")
    for section_label, key in [
        ("All support staff",               "all_staff_responsibilities"),
        ("Team leader / coordinator",        "team_leader_responsibilities"),
        ("Behaviour Support Practitioner",   "practitioner_responsibilities"),
    ]:
        h2(section_label)
        for item in impl.get(key, []): bul(item)
        if not impl.get(key): placeholder(f"{section_label} responsibilities")

    h2("8.2  Training requirements")
    training = impl.get("training_requirements", [])
    if training:
        tr = doc.add_table(rows=1, cols=4)
        tr.style = 'Table Grid'
        for i, hdr in enumerate(["Training topic", "Who requires it",
                                   "Delivery method", "Must be completed"]):
            tr.rows[0].cells[i].text = hdr
            if tr.rows[0].cells[i].paragraphs[0].runs:
                tr.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for t_item in training:
            row = tr.add_row().cells
            row[0].text = t_item.get("topic", "")
            row[1].text = t_item.get("who", "")
            row[2].text = t_item.get("method", "")
            row[3].text = t_item.get("timeframe", "")
        doc.add_paragraph("")
    else:
        placeholder("List training requirements")

    h2("8.3  Competency requirements")
    if impl.get("competency_statement"): body(impl["competency_statement"])
    else: placeholder("Define what staff must demonstrate before implementing this plan")

    h2("8.4  Plan go-live requirements")
    if impl.get("go_live_requirements"): body(impl["go_live_requirements"])
    else: placeholder("Requirements to be met before the plan is implemented")

    h2("8.5  Review schedule")
    if impl.get("review_schedule"): body(impl["review_schedule"])
    else: placeholder("Describe routine and trigger-based review schedule")

    # ── S9: Sign-off ──────────────────────────────────────────────────────────
    doc.add_paragraph("")
    h1("9. Plan approval and sign-off")
    so = doc.add_table(rows=5, cols=3)
    so.style = 'Table Grid'
    for i, (role, name, dt) in enumerate([
        ("Role",                         "Name and signature", "Date"),
        ("Behaviour Support Practitioner", practitioner or "", ""),
        ("Supervisor / Senior Practitioner", "", ""),
        ("Participant or nominee",         "", ""),
        ("Provider representative",        "", ""),
    ]):
        so.rows[i].cells[0].text = role
        so.rows[i].cells[1].text = name
        so.rows[i].cells[2].text = dt
        if i == 0:
            for ci in range(3):
                if so.rows[i].cells[ci].paragraphs[0].runs:
                    so.rows[i].cells[ci].paragraphs[0].runs[0].bold = True
    doc.add_paragraph("")
    foot = doc.add_paragraph()
    fr = foot.add_run(
        f"DRAFT document generated {date.today().strftime('%d %B %Y')} to assist "
        f"the Behaviour Support Practitioner. Requires clinical review before submission.  "
        f"{contact or ''}"
    )
    fr.font.size = Pt(8); fr.italic = True; fr.font.color.rgb = GREYC

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════════
# FUNCTIONAL BEHAVIOUR ASSESSMENT (Tab 5)
# ══════════════════════════════════════════════════════════════════════════════

FBA_PROMPT = """\
You are an experienced NDIS Behaviour Support Practitioner interpreting standardised FBA assessment results.

The text below has been extracted from completed assessment tool score reports.

For each tool present:
1. Extract all scores precisely as reported
2. Write an interpretive narrative in the clinical voice used in an NDIS Comprehensive PBSP
3. Link findings to behaviour patterns where the data supports it

Narratives must be:
- Written in third person about the participant
- 3-5 paragraphs per tool (matching the depth of a professional NDIS PBSP)
- Specific to the scores — not generic descriptions of the tool
- Clinically accurate and person-centred

Return ONLY valid JSON — no commentary, no markdown fences.

JSON structure:
{{
  "client_name": "participant name if visible in the data, otherwise null",
  "tools_identified": ["list of tools found"],
  "direct_methods": ["list of direct assessment methods used — observations, session notes, etc."],
  "indirect_methods": ["list of indirect/standardised tools used"],
  "results": {{
    "edaq": {{
      "present": true,
      "completed_by": "name and date if visible",
      "total_score": "X/74",
      "threshold_note": "threshold for relevant age range",
      "narrative": "3-4 paragraph interpretive narrative specific to these scores"
    }},
    "qabf": {{
      "present": true,
      "behaviours": [
        {{
          "behaviour_name": "name of behaviour or presentation assessed",
          "completed_by": "name and date if visible",
          "scores": {{
            "attention": "X/15",
            "escape": "X/15",
            "non_social": "X/15",
            "physical": "X/15",
            "tangible": "X/15"
          }},
          "narrative": "2-3 paragraph interpretive narrative for this behaviour's scores"
        }}
      ]
    }},
    "sensory_profile": {{
      "present": true,
      "completed_by": "name and date if visible",
      "processing": {{
        "seeking": "Much less than others / Less than others / Just like others / More than others / Much more than others",
        "avoiding": "...",
        "sensitivity": "...",
        "registration": "..."
      }},
      "sensory": {{
        "auditory": "...",
        "visual": "...",
        "touch": "...",
        "movement": "...",
        "body_position": "...",
        "oral": "..."
      }},
      "behaviour_scores": {{
        "conduct": "...",
        "social_emotional": "...",
        "attentional": "..."
      }},
      "narrative": "4-5 paragraph interpretive narrative specific to this person's sensory profile"
    }},
    "fast": {{
      "present": true,
      "completed_by": "name and date if visible",
      "scores": {{
        "social_connection": "X/4",
        "social_autonomy": "X/4",
        "automatic_sensory": "X/4",
        "automatic_interoceptive": "X/4"
      }},
      "narrative": "3-4 paragraph interpretive narrative specific to these scores"
    }},
    "tmq": {{
      "present": true,
      "completed_by": "name and date if visible",
      "factors": {{
        "special_interests": "X/5",
        "rumination_anxiety": "X/5",
        "need_for_routines": "X/5",
        "environmental_impact": "X/5",
        "losing_track": "X/5",
        "decision_making": "X/5",
        "anxiety_reducing": "X/5",
        "social_interactions": "X/5",
        "overall_average": "X/5"
      }},
      "narrative": "3-4 paragraph interpretive narrative specific to these scores"
    }},
    "abas": {{
      "present": false,
      "completed_by": null,
      "domain_scores": {{}},
      "narrative": null
    }},
    "abc_data": {{
      "present": false,
      "summary": null,
      "narrative": null
    }}
  }},
  "function_of_presentation": {{
    "behaviour_description": "2-3 paragraphs describing the behaviour pattern, its observable features, and escalation pathway",
    "relevant_evidence": "2-3 paragraphs synthesising findings across all tools — how they collectively explain the behaviour",
    "summary_statement": "1-2 paragraphs stating the confirmed primary and secondary functions of behaviour — evidence-based, not hypothesised"
  }}
}}

RULES:
- Extract scores exactly as reported — do not estimate or fill in missing data
- Mark any tool not found in the data as present: false with null values
- Narratives must be specific to the actual scores — never generic
- The summary_statement in function_of_presentation must state confirmed function (evidence-based), not hypothesis
- Use Australian spelling throughout
- If a score or section is unclear from the PDF, note [SCORE UNCLEAR — VERIFY] in that field

ASSESSMENT DATA:
{text}

CLIENT CONTEXT (from other tabs):
{client_context}
"""


def extract_and_interpret_fba(pdf_texts: list, client_context: str, api_key: str) -> dict:
    combined = "\n\n---\n\n".join(pdf_texts)
    prompt = FBA_PROMPT.format(
        text=combined[:50000],
        client_context=client_context or "Not provided",
    )
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model="claude-opus-4-5", max_tokens=8000,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = msg.content[0].text.strip()
    if raw.startswith("```"): raw = "\n".join(raw.split("\n")[1:])
    if raw.endswith("```"):   raw = "\n".join(raw.split("\n")[:-1])
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        last = raw.rfind('"}')
        if last > 0:
            try: return json.loads(raw[:last+2] + "\n}}")
            except Exception: pass
        raise ValueError(
            "Could not parse FBA results. The PDF text may be too long or unclear. "
            "Try uploading fewer files at once."
        )


def create_fba_docx(fba: dict, client_name: str,
                     practitioner: str = "", contact: str = "") -> BytesIO:
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    NAVY  = RGBColor(0x0D, 0x4F, 0x6E)
    TEALC = RGBColor(0x1A, 0x9B, 0x8A)
    REDC  = RGBColor(0xC0, 0x39, 0x2B)
    AMBC  = RGBColor(0xD4, 0x70, 0x0A)
    GREYC = RGBColor(0x4A, 0x4A, 0x4A)

    def h1(text):
        p = doc.add_heading(text, level=1)
        for r in p.runs: r.font.color.rgb = NAVY; r.font.size = Pt(14)
        return p

    def h2(text):
        p = doc.add_heading(text, level=2)
        for r in p.runs: r.font.color.rgb = TEALC; r.font.size = Pt(12)
        return p

    def body(text, italic=False):
        p = doc.add_paragraph(text)
        if italic:
            for r in p.runs: r.italic = True
        return p

    def bul(text):
        try:    return doc.add_paragraph(text, style='List Bullet')
        except: return doc.add_paragraph(f"• {text}")

    def score_table(rows_data, col_headers=None):
        n_cols = len(rows_data[0]) if rows_data else 2
        t = doc.add_table(rows=len(rows_data), cols=n_cols)
        t.style = 'Table Grid'
        for i, row in enumerate(rows_data):
            for j, cell_val in enumerate(row):
                t.rows[i].cells[j].text = str(cell_val) if cell_val else ""
                if i == 0 and t.rows[i].cells[j].paragraphs[0].runs:
                    t.rows[i].cells[j].paragraphs[0].runs[0].bold = True
        doc.add_paragraph("")
        return t

    def narrative_block(text):
        if not text: return
        for para in text.split('\n'):
            para = para.strip()
            if para:
                body(para)

    # ── Cover ──────────────────────────────────────────────────────────────────
    title = doc.add_heading("Section 8: Functional Behaviour Assessment", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs: r.font.color.rgb = NAVY; r.font.size = Pt(16)
    sub = doc.add_paragraph(client_name)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs: r.bold = True; r.font.size = Pt(13); r.font.color.rgb = TEALC
    doc.add_paragraph("")

    p = doc.add_paragraph()
    r = p.add_run("DRAFT — AI-generated interpretation of uploaded assessment data. "
                  "All narratives must be reviewed and verified by the Behaviour Support "
                  "Practitioner before inclusion in the Comprehensive PBSP.")
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = REDC
    doc.add_page_break()

    # ── Overview ───────────────────────────────────────────────────────────────
    h1("Overview")
    body(
        "A Functional Behaviour Assessment (FBA) was completed by the Behaviour Support "
        f"Practitioner, {practitioner or '[PRACTITIONER NAME]'}, with information gathered "
        "through direct observations, stakeholder consultation, standardised assessment "
        "measures, and review of available records."
    )
    doc.add_paragraph("")

    direct_methods = fba.get("direct_methods", [])
    if direct_methods:
        body("Direct assessment methods:", italic=True)
        for m in direct_methods: bul(m)
    doc.add_paragraph("")

    indirect_methods = fba.get("indirect_methods", [])
    if indirect_methods:
        body("The following standardised tools and information sources were utilised:", italic=True)
        for m in indirect_methods: bul(m)
    doc.add_paragraph("")

    # ── Assessment Results ─────────────────────────────────────────────────────
    h1("Assessment Results")
    results = fba.get("results", {})

    # EDA-Q
    edaq = results.get("edaq", {})
    if edaq.get("present"):
        h2("Extreme Demand Avoidance Questionnaire (EDA-Q)")
        body(
            "The EDA-Q consists of 26 questions aimed to measure the presence of "
            "extreme/pathological demand avoidance in a person's behaviour over the most "
            "recent 6 months. The questionnaire is used for the purpose of understanding "
            "only and is not a diagnostic test."
        )
        doc.add_paragraph("")
        score_table([
            [f"{client_name}'s EDA-Q Results"],
            ["Total score", edaq.get("total_score", "[SCORE]")],
            ["Threshold note", edaq.get("threshold_note", "")],
            ["Completed by", edaq.get("completed_by", "[VERIFY]")],
        ])
        narrative_block(edaq.get("narrative", ""))
        doc.add_paragraph("")

    # QABF
    qabf = results.get("qabf", {})
    if qabf.get("present"):
        h2("Questions About Behavioural Function (QABF)")
        body(
            "The QABF is completed by members of a person's support network to assess "
            "the impact of behaviour on a person's environment. The questionnaire consists "
            "of 25 questions and is completed for each individual behaviour being assessed."
        )
        doc.add_paragraph("")
        for beh in qabf.get("behaviours", []):
            h2(f"{client_name}'s Results from the QABF")
            if beh.get("behaviour_name"):
                body(f"Assessed presentation: {beh['behaviour_name']}", italic=True)
            scores = beh.get("scores", {})
            score_table([
                ["Category", "Score"],
                ["Attention",   scores.get("attention",   "[VERIFY]")],
                ["Escape",      scores.get("escape",      "[VERIFY]")],
                ["Non-social",  scores.get("non_social",  "[VERIFY]")],
                ["Physical",    scores.get("physical",    "[VERIFY]")],
                ["Tangible",    scores.get("tangible",    "[VERIFY]")],
                ["Completed by", beh.get("completed_by", "[VERIFY]"), ""],
            ])
            narrative_block(beh.get("narrative", ""))
            doc.add_paragraph("")

    # Sensory Profile 2.0
    sp = results.get("sensory_profile", {})
    if sp.get("present"):
        h2("Sensory Profile 2.0")
        body(
            "The Sensory Profile 2.0 is a comprehensive caregiver questionnaire which "
            "assesses the way a person interacts with and experiences sensory input across "
            "different environments."
        )
        doc.add_paragraph("")
        h2(f"{client_name}'s Results from the Sensory Profile 2.0")
        if sp.get("completed_by"):
            body(f"Completed by {sp['completed_by']}.", italic=True)
        proc = sp.get("processing", {})
        if proc:
            score_table([
                ["Sensory Profile 2.0 Score Summary – Processing",
                 "Much less than others", "Less than others", "Just like others",
                 "More than others", "Much more than others"],
                _sp_row("Seeking/Seeker",      proc.get("seeking",      "")),
                _sp_row("Avoiding/Avoider",    proc.get("avoiding",     "")),
                _sp_row("Sensitivity/Sensor",  proc.get("sensitivity",  "")),
                _sp_row("Registration/Bystander", proc.get("registration", "")),
            ])
        sens = sp.get("sensory", {})
        if sens:
            score_table([
                ["Sensory Profile 2.0 Score Summary – Sensory",
                 "Much less than others", "Less than others", "Just like others",
                 "More than others", "Much more than others"],
                _sp_row("Auditory Processing",       sens.get("auditory",      "")),
                _sp_row("Visual Processing",         sens.get("visual",        "")),
                _sp_row("Touch Processing",          sens.get("touch",         "")),
                _sp_row("Movement Processing",       sens.get("movement",      "")),
                _sp_row("Body Position Processing",  sens.get("body_position", "")),
                _sp_row("Oral Sensory Processing",   sens.get("oral",          "")),
            ])
        beh_scores = sp.get("behaviour_scores", {})
        if beh_scores:
            score_table([
                ["Sensory Profile 2.0 Score Summary – Behaviour",
                 "Much less than others", "Less than others", "Just like others",
                 "More than others", "Much more than others"],
                _sp_row("Conduct",                 beh_scores.get("conduct",         "")),
                _sp_row("Social Emotional Responses", beh_scores.get("social_emotional", "")),
                _sp_row("Attentional Responses",   beh_scores.get("attentional",     "")),
            ])
        narrative_block(sp.get("narrative", ""))
        doc.add_paragraph("")

    # FAST
    fast = results.get("fast", {})
    if fast.get("present"):
        h2("Functional Analysis Screening Tool (FAST)")
        body(
            "The FAST is a 16-item questionnaire that can be completed by anyone who "
            "regularly interacts with the participant. The scoring suggests suspected "
            "functions of behaviour. The greater the score (highest being 4), the greater "
            "the impact of the corresponding category on behaviour."
        )
        doc.add_paragraph("")
        h2(f"{client_name}'s Results from the FAST")
        scores = fast.get("scores", {})
        score_table([
            ["FAST Category", "Impact on Behaviour"],
            ["Social (need for connection and/or intellectual stimulation)",
             scores.get("social_connection",      "[VERIFY]")],
            ["Social (need for autonomy and overwhelmed by groups and demands)",
             scores.get("social_autonomy",         "[VERIFY]")],
            ["Automatic (need for sensory stimulation)",
             scores.get("automatic_sensory",       "[VERIFY]")],
            ["Automatic (complications in interoceptive registration)",
             scores.get("automatic_interoceptive", "[VERIFY]")],
            ["Completed by", fast.get("completed_by", "[VERIFY]")],
        ])
        narrative_block(fast.get("narrative", ""))
        doc.add_paragraph("")

    # TMQ
    tmq = results.get("tmq", {})
    if tmq.get("present"):
        h2("The Monotropism Questionnaire (TMQ)")
        body(
            "The Monotropism Questionnaire is a questionnaire developed by autistic "
            "researchers, containing 47 statements. The results demonstrate the degree "
            "to which the person experiences and is impacted by Monotropism. The higher "
            "the score, the higher the likely experience and impact."
        )
        doc.add_paragraph("")
        h2(f"{client_name}'s Results from the TMQ")
        factors = tmq.get("factors", {})
        score_table([
            ["TMQ Factor", "Presentation of Factor"],
            ["Special Interests",
             factors.get("special_interests",   "[VERIFY]")],
            ["Rumination and anxiety",
             factors.get("rumination_anxiety",  "[VERIFY]")],
            ["Need for routines",
             factors.get("need_for_routines",   "[VERIFY]")],
            ["Environmental impact on the attention tunnel",
             factors.get("environmental_impact","[VERIFY]")],
            ["Losing track of other factors when focusing on special interests",
             factors.get("losing_track",        "[VERIFY]")],
            ["Struggle with decision-making",
             factors.get("decision_making",     "[VERIFY]")],
            ["Anxiety-reducing effect of special interests",
             factors.get("anxiety_reducing",    "[VERIFY]")],
            ["Managing social interactions",
             factors.get("social_interactions", "[VERIFY]")],
            ["Overall Average",
             factors.get("overall_average",     "[VERIFY]")],
            ["Completed by", tmq.get("completed_by", "[VERIFY]")],
        ])
        narrative_block(tmq.get("narrative", ""))
        doc.add_paragraph("")

    # ABAS (optional)
    abas = results.get("abas", {})
    if abas.get("present"):
        h2("Adaptive Behaviour Assessment System (ABAS)")
        domain_scores = abas.get("domain_scores", {})
        if domain_scores:
            rows = [["Domain", "Score"]]
            for k, v in domain_scores.items():
                rows.append([k, str(v)])
            score_table(rows)
        if abas.get("completed_by"):
            body(f"Completed by {abas['completed_by']}.", italic=True)
        narrative_block(abas.get("narrative", ""))
        doc.add_paragraph("")

    # ABC data (optional)
    abc = results.get("abc_data", {})
    if abc.get("present"):
        h2("ABC Recording Data")
        if abc.get("summary"):
            body(abc["summary"], italic=True)
        narrative_block(abc.get("narrative", ""))
        doc.add_paragraph("")

    # ── Function of Presentation ───────────────────────────────────────────────
    doc.add_page_break()
    h1("Function of Presentation")
    fop = fba.get("function_of_presentation", {})

    if fop.get("behaviour_description"):
        body("Behaviour Description:", italic=False)
        narrative_block(fop["behaviour_description"])
        doc.add_paragraph("")

    if fop.get("relevant_evidence"):
        body("Relevant Evidence:", italic=False)
        narrative_block(fop["relevant_evidence"])
        doc.add_paragraph("")

    if fop.get("summary_statement"):
        body("Summary Statement:", italic=False)
        narrative_block(fop["summary_statement"])
        doc.add_paragraph("")

    # Footer
    doc.add_paragraph("")
    foot = doc.add_paragraph()
    fr = foot.add_run(
        f"FBA Section generated {date.today().strftime('%d %B %Y')} — "
        f"DRAFT, requires practitioner review before inclusion in Comprehensive PBSP.  "
        f"{contact or ''}"
    )
    fr.font.size = Pt(8); fr.italic = True; fr.font.color.rgb = GREYC

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _sp_row(label: str, rating: str) -> list:
    """Build a Sensory Profile table row with an X in the correct column."""
    cols = [
        "Much less than others", "Less than others", "Just like others",
        "More than others", "Much more than others"
    ]
    row = [label, "", "", "", "", ""]
    for i, col in enumerate(cols):
        if rating and col.lower() in rating.lower():
            row[i + 1] = "X"
            break
    return row


# ══════════════════════════════════════════════════════════════════════════════
# STREAMLIT UI
# ══════════════════════════════════════════════════════════════════════════════

st.title("📋 PBS Support Tool")

# API key — sidebar (visible from all tabs)
try:
    api_key = st.secrets.get("ANTHROPIC_API_KEY", "")
except Exception:
    api_key = ""

if not api_key:
    with st.sidebar:
        st.markdown("### 🔑 API Key")
        api_key = st.text_input(
            "Anthropic API key",
            type="password",
            help="Get a key at console.anthropic.com — a few cents per document",
            key="api_key_input",
        )
        st.caption("Your key is never stored.")

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📊 Behaviour Recording",
    "📄 Generate from PBSP",
    "🧠 Strategy Recommender",
    "📋 NDIS BSP Draft",
    "🔍 Functional Behaviour Assessment",
])


# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — BEHAVIOUR RECORDING ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
with tab1:
    st.markdown(
        "Upload the completed **30-day behaviour recording** from families or caregivers. "
        "The tool extracts the behaviours of concern, identifies patterns and triggers, "
        "and summarises what strategies carers have already been trying."
    )
    st.info(
        "📊 **Using Microsoft Forms?** Open your form → Responses → "
        "click the Excel icon to download responses. Upload that file here."
    )
    st.divider()

    accept_types = ["xlsx", "pdf", "docx"] if EXCEL_AVAILABLE else ["pdf", "docx"]
    if not EXCEL_AVAILABLE:
        st.warning("Excel (.xlsx) support requires openpyxl — install it to enable Microsoft Forms exports.")

    rec_file = st.file_uploader(
        "Upload 30-day behaviour recording",
        type=accept_types,
        key="rec_upload",
        help="Excel export from Microsoft Forms (.xlsx), Word (.docx), or PDF"
    )
    rec_btn = st.button(
        "Analyse recording", type="primary",
        disabled=not (rec_file and api_key), key="rec_analyse"
    )

    if rec_btn and rec_file and api_key:
        with st.spinner("Reading recording…"):
            fb   = rec_file.read()
            name = rec_file.name.lower()
            if name.endswith(".xlsx"):
                rec_text = extract_text_from_excel(fb)
            elif name.endswith(".docx"):
                rec_text = extract_text_from_docx(fb)
            else:
                rec_text = extract_text_from_pdf(fb)

            if not rec_text.strip():
                st.error("Could not extract text from this file. Try saving as PDF or Word and re-uploading.")
                st.stop()

        with st.spinner("Analysing behaviour data…"):
            try:
                rec_data = extract_behaviour_recording(rec_text, api_key)
                st.session_state["t0_data"] = rec_data
            except Exception as e:
                st.error(f"Analysis failed: {e}")
                st.stop()

        n_beh = len(rec_data.get("behaviours", []))
        st.success(
            f"✅  Recording analysed — "
            f"**{rec_data.get('total_incidents', '?')} incidents** across "
            f"**{n_beh} behaviour{'s' if n_beh != 1 else ''}**"
        )

    if "t0_data" in st.session_state:
        rec_data = st.session_state["t0_data"]
        st.divider()

        # ── Summary metrics ──
        total = rec_data.get("total_incidents") or "—"
        n_beh = len(rec_data.get("behaviours", []))
        period = rec_data.get("recording_period") or "30 days"
        col1, col2, col3 = st.columns(3)
        with col1: st.metric("Total incidents", total)
        with col2: st.metric("Behaviours identified", n_beh)
        with col3: st.metric("Period", period)

        # ── Behaviours ──
        behs = rec_data.get("behaviours", [])
        if behs:
            st.markdown("##### Behaviours identified")
            t = rec_data.get("total_incidents") or sum((b.get("frequency") or 0) for b in behs) or 1
            for b in behs:
                freq = b.get("frequency") or 0
                pct  = round(freq / t * 100) if t else 0
                with st.expander(f"**{b.get('label', 'Behaviour')}** — {freq} incidents ({pct}%)"):
                    if b.get("carer_description"):
                        st.markdown(f"*Carers described: {b['carer_description']}*")
                    for d_ in b.get("descriptors", []):
                        st.markdown(f"- {d_}")

        # ── Patterns ──
        patterns = rec_data.get("patterns", {})
        all_patterns = (
            patterns.get("time_of_day", []) +
            patterns.get("day_of_week", []) +
            patterns.get("settings",    []) +
            patterns.get("other",       [])
        )
        if all_patterns:
            st.markdown("##### Patterns observed")
            for p in all_patterns:
                st.markdown(f"- {p}")

        # ── Triggers ──
        triggers = rec_data.get("triggers", [])
        if triggers:
            st.markdown("##### Triggers noted by carers")
            for t_ in triggers:
                st.markdown(f"- {t_}")

        # ── Carer responses ──
        responses = rec_data.get("carer_responses", [])
        if responses:
            st.markdown("##### Strategies carers have been using")
            for r in responses:
                st.markdown(f"- {r}")

        # ── Concerns & notes ──
        if rec_data.get("carer_concerns"):
            st.info(f"**Carer concerns:** {rec_data['carer_concerns']}")
        notes = rec_data.get("notes_for_practitioner", [])
        if notes:
            with st.expander("⚠️ Notes for PBS Practitioner"):
                for n in notes:
                    st.markdown(f"- {n}")

        st.divider()

        # ── Download PDF report ──
        with st.spinner("Building analysis report…"):
            try:
                rep_buf = generate_recording_report(rec_data)
            except Exception as e:
                st.error(f"Report error: {e}"); st.stop()

        safe = (rec_data.get("client_name") or "Client").replace(" ", "_")
        st.download_button(
            "📥 Download analysis report (PDF)", rep_buf,
            f"{safe}_Behaviour_Recording_Analysis.pdf",
            "application/pdf", use_container_width=True,
        )
        st.caption(
            "Take this analysis to your next meeting with the family to confirm and refine. "
            "Switch to the **Strategy Recommender** tab to generate strategies based on this recording."
        )

        if st.button("🗑 Clear recording data", key="rec_clear"):
            del st.session_state["t0_data"]
            st.rerun()

    elif not rec_file:
        st.caption("Supported: Excel (.xlsx from Microsoft Forms), Word (.docx), PDF")

    st.divider()
    st.caption(
        "Recording data is processed in memory only and never stored permanently. "
        "Handle all client information in line with your organisation's privacy policy."
    )


# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — GENERATE FROM PBSP
# ══════════════════════════════════════════════════════════════════════════════
with tab2:
    st.markdown(
        "Upload a client's **Positive Behaviour Support Plan** (PDF or Word) and this tool "
        "will generate a **support reference card** and an **ABC recording form** — "
        "both pre-populated with the client's specific information."
    )
    st.divider()

    uploaded = st.file_uploader("Upload PBSP (PDF or Word .docx)", type=["pdf","docx"])
    gen_btn  = st.button("Generate documents", type="primary",
                         disabled=not (uploaded and api_key), key="gen_tab2")

    if gen_btn and uploaded and api_key:
        with st.spinner("Reading plan…"):
            fb = uploaded.read()
            pbsp_text = extract_text_from_docx(fb) if uploaded.name.lower().endswith(".docx") \
                        else extract_text_from_pdf(fb)

        with st.spinner("Extracting client information…"):
            try:   data = extract_client_data(pbsp_text, api_key)
            except Exception as e: st.error(f"Extraction failed: {e}"); st.stop()

        with st.spinner("Generating PDFs…"):
            try:
                card_buf = generate_support_card(data)
                abc_buf  = generate_abc_form(data)
            except Exception as e: st.error(f"PDF error: {e}"); st.stop()

        st.session_state["t1_data"] = data   # share with Tab 3
        st.success(f"✅  Documents generated for **{data.get('name','client')}**")
        st.divider()
        safe = data.get("name","client").replace(" ","_")
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("#### 📄 Support Reference Card")
            st.markdown("One-page quick reference — print and laminate.")
            st.download_button("Download support card (PDF)", card_buf,
                               f"{safe}_Support_Card.pdf", "application/pdf",
                               use_container_width=True)
        with c2:
            st.markdown("#### 📋 ABC Recording Form")
            st.markdown("Pre-populated with this client's behaviours.")
            st.download_button("Download ABC form (PDF)", abc_buf,
                               f"{safe}_ABC_Form.pdf", "application/pdf",
                               use_container_width=True)
        with st.expander("Review extracted information"):
            st.json(data)

    st.divider()
    st.caption("Always review generated documents before distributing to staff. "
               "Handle all client documents in line with your organisation's privacy policy.")


# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 — STRATEGY RECOMMENDER
# ══════════════════════════════════════════════════════════════════════════════
with tab3:
    st.markdown(
        "Upload a client's PBSP and your **strategy library** — the tool will match the most "
        "appropriate strategies from your library to this client's behaviours and generate a "
        "**printable strategy report**."
    )
    st.divider()

    # ── STEP 1: Client source ─────────────────────────────────────────────────
    st.markdown("#### Step 1 — Client information")

    source_opts = ["📄 Upload a PBSP (auto-extract)", "✏️ Enter manually"]
    if "t1_data" in st.session_state:
        source_opts.insert(0,
            f"♻️ Use PBSP from Tab 2 ({st.session_state['t1_data'].get('name','')})")
    if "t0_data" in st.session_state:
        rec_name = st.session_state["t0_data"].get("client_name") or "client"
        source_opts.insert(0, f"📊 Use behaviour recording from Tab 1 ({rec_name})")

    src = st.radio("Where is the client information coming from?",
                   source_opts, key="sr_src", horizontal=True)

    client_info    = None
    valid_behs     = []
    freetext_value = ""

    src = src or source_opts[0]

    if src.startswith("📊"):
        # ── use behaviour recording from Tab 1 ──
        t0 = st.session_state["t0_data"]
        client_info, valid_behs = recording_to_sr_format(t0)
        rec_name = t0.get("client_name") or "Client"
        period   = t0.get("recording_period") or "30-day recording"
        st.success(f"✅  Using behaviour recording for **{rec_name}** ({period})")
        st.info(
            f"**{t0.get('total_incidents', '?')} incidents** recorded across "
            f"**{len(valid_behs)} behaviour(s)**. "
            "Add diagnosis and communication details in the report footer if known."
        )
        with st.expander("View behaviours from recording"):
            for b in t0.get("behaviours", []):
                st.markdown(
                    f"- **{b.get('label', '')}** ({b.get('frequency', '?')} incidents): "
                    f"{b.get('carer_description', '')}"
                )

    elif src.startswith("♻️"):
        # ── reuse Tab 2 PBSP extraction ──
        t1 = st.session_state["t1_data"]
        client_info, valid_behs = pbsp_to_sr_format(t1)
        st.success(f"✅  Using: **{t1.get('name','')}** — {t1.get('age_info','')}")
        with st.expander("View extracted profile and behaviours"):
            cols = st.columns(2)
            with cols[0]:
                st.markdown(f"**Name:** {t1.get('name','')}")
                st.markdown(f"**Profile:** {t1.get('age_info','')}")
                st.markdown("**About:**")
                for a in t1.get("about",[]): st.markdown(f"- {a}")
            with cols[1]:
                st.markdown("**Behaviours of concern:**")
                for b in t1.get("behaviours",[]): st.markdown(
                    f"- **{b.get('label','')}:** " + ", ".join(b.get("descriptors",[])))
                st.markdown("**Triggers:**")
                for t_ in t1.get("triggers",[]): st.markdown(f"- {t_}")

    elif src.startswith("📄"):
        # ── upload PBSP for Tab 3 ──
        sr_pbsp_file = st.file_uploader("Upload client's PBSP (PDF or Word)",
                                         type=["pdf","docx"], key="sr_pbsp_upload")
        if sr_pbsp_file:
            if st.button("Extract client information", type="secondary", key="sr_pbsp_extract"):
                if not api_key:
                    st.error("API key required — enter it in the sidebar.")
                else:
                    with st.spinner("Extracting from PBSP…"):
                        fb = sr_pbsp_file.read()
                        txt = extract_text_from_docx(fb) if sr_pbsp_file.name.lower().endswith(".docx") \
                              else extract_text_from_pdf(fb)
                        try:
                            extracted = extract_client_data(txt, api_key)
                            st.session_state["sr_pbsp_data"] = extracted
                            st.success(f"✅  Extracted: {extracted.get('name','')}")
                        except Exception as e:
                            st.error(f"Extraction failed: {e}")

        if "sr_pbsp_data" in st.session_state:
            pd_ = st.session_state["sr_pbsp_data"]
            client_info, valid_behs = pbsp_to_sr_format(pd_)
            st.info(f"Using extracted data for **{pd_.get('name','')}** — "
                    f"{len(valid_behs)} behaviour(s) found.")
            with st.expander("View extracted behaviours"):
                for b in pd_.get("behaviours",[]): st.markdown(
                    f"- **{b.get('label','')}:** " + ", ".join(b.get("descriptors",[])))
            if st.button("🗑 Clear extracted data", key="sr_pbsp_clear"):
                del st.session_state["sr_pbsp_data"]; st.rerun()

    else:
        # ── manual entry ──
        ca, cb = st.columns(2)
        with ca:
            sr_name = st.text_input("Client name *", key="sr_name",
                                     placeholder="e.g. Alex Thompson")
            sr_age  = st.text_input("Age", key="sr_age", placeholder="e.g. 24")
        with cb:
            sr_diag  = st.text_input("Diagnosis / condition", key="sr_diag",
                                      placeholder="e.g. Autism Spectrum Disorder, FASD, ABI")
            sr_comms = st.selectbox("Communication level", key="sr_comms",
                                     options=["Verbal","Limited verbal","Non-verbal","Uses AAC device"])
        sr_other = st.text_area("Other relevant context (optional)", key="sr_other", height=75,
                                 placeholder="e.g. Routines are very important, history of trauma…")

        st.markdown("**Behaviours of concern**")
        entry_style = st.radio(
            "How would you like to describe the behaviours?",
            ["📝 Describe in your own words", "📋 Structured entry"],
            key="sr_entry_style", horizontal=True,
        )

        if (entry_style or "").startswith("📝"):
            sr_freetext = st.text_area(
                "Describe what you've observed — write naturally",
                key="sr_freetext",
                height=160,
                placeholder=(
                    "Write however feels natural — describe what you see, when it happens, "
                    "and what seems to set it off. The tool will identify the behaviours, "
                    "assign clinical labels, and develop strategies.\n\n"
                    "e.g. 'Alex hits out at staff when demands are placed, especially during "
                    "transitions. He also bangs his head on surfaces when he's frustrated or "
                    "overwhelmed. Sometimes he shouts and swears when things don't go his way.'"
                ),
            )
            freetext_value = (sr_freetext or "").strip()
            valid_behs = [{"_freetext": True}] if freetext_value else []
            if freetext_value:
                client_info = {
                    "name":      (sr_name or "").strip(),
                    "age":       (sr_age  or "").strip() or "not specified",
                    "diagnosis": (sr_diag or "").strip() or "not specified",
                    "comms":     sr_comms or "Verbal",
                    "other":     (sr_other or "").strip() or "none provided",
                }

        else:
            freetext_value = ""
            if "sr_n" not in st.session_state:
                st.session_state.sr_n = 1

            beh_raw = []
            for i in range(st.session_state.sr_n):
                with st.expander(f"Behaviour {i+1}", expanded=True):
                    bn = st.text_input("Behaviour name / label", key=f"sr_bn_{i}",
                                        placeholder="e.g. Physical aggression, Self-injurious behaviour")
                    bd = st.text_area("Describe what it looks like", key=f"sr_bd_{i}", height=70,
                                       placeholder="e.g. Hitting out, throwing objects — lasts 2–5 minutes")
                    bt = st.text_area("Known triggers / when it tends to occur", key=f"sr_bt_{i}", height=70,
                                       placeholder="e.g. When demands are placed, during transitions")
                    beh_raw.append({"name": bn or "", "description": bd or "", "triggers": bt or ""})

            caddbtn, crmbtn = st.columns(2)
            with caddbtn:
                if st.session_state.sr_n < 5 and st.button("＋ Add another behaviour", key="sr_add"):
                    st.session_state.sr_n += 1; st.rerun()
            with crmbtn:
                if st.session_state.sr_n > 1 and st.button("− Remove last", key="sr_rm"):
                    st.session_state.sr_n -= 1; st.rerun()

            valid_behs = [b for b in beh_raw if b["name"].strip()]
            if valid_behs:
                client_info = {
                    "name":      (sr_name or "").strip(),
                    "age":       (sr_age  or "").strip() or "not specified",
                    "diagnosis": (sr_diag or "").strip() or "not specified",
                    "comms":     sr_comms or "Verbal",
                    "other":     (sr_other or "").strip() or "none provided",
                }

    st.divider()

    # ── STEP 2: Strategy library ──────────────────────────────────────────────
    st.markdown("#### Step 2 — Your strategy library")

    builtin_lib_text, builtin_lib_name = load_builtin_library()

    sr_lib_file = st.file_uploader(
        "Upload a different library for this session (optional)",
        type=["pdf","docx"], key="sr_lib_upload",
        help="Overrides the built-in library for this session only",
    ) if builtin_lib_text else st.file_uploader(
        "Upload strategy library (PDF or Word)", type=["pdf","docx"], key="sr_lib_upload",
        help="Strategies, clinical guidelines, PBS manuals — any document with your approved strategies",
    )

    if sr_lib_file:
        fb = sr_lib_file.read()
        uploaded_text = extract_text_from_docx(fb) if sr_lib_file.name.lower().endswith(".docx") \
                        else extract_text_from_pdf(fb)
        st.session_state["sr_lib_text"]     = uploaded_text
        st.session_state["sr_lib_name"]     = sr_lib_file.name
        st.session_state["sr_lib_uploaded"] = True

    if "sr_lib_text" in st.session_state:
        effective_lib_text = st.session_state["sr_lib_text"]
        effective_lib_name = st.session_state["sr_lib_name"]
        is_uploaded_override = st.session_state.get("sr_lib_uploaded", False)
        word_count = len(effective_lib_text.split())
        if is_uploaded_override and builtin_lib_text:
            st.success(
                f"✅  Using uploaded library: **{effective_lib_name}** ({word_count:,} words)"
                f" — overriding built-in library for this session."
            )
            if st.button("↩ Revert to built-in library", key="sr_lib_revert"):
                del st.session_state["sr_lib_text"]
                del st.session_state["sr_lib_name"]
                del st.session_state["sr_lib_uploaded"]
                st.rerun()
        elif is_uploaded_override:
            st.success(f"✅  Library loaded: **{effective_lib_name}** ({word_count:,} words)")
            if st.button("🗑 Clear library", key="sr_lib_clear"):
                del st.session_state["sr_lib_text"]
                del st.session_state["sr_lib_name"]
                del st.session_state["sr_lib_uploaded"]
                st.rerun()
        else:
            pass
    elif builtin_lib_text:
        effective_lib_text = builtin_lib_text
        effective_lib_name = builtin_lib_name
        word_count = len(builtin_lib_text.split())
        st.success(
            f"✅  **Built-in strategy library:** {builtin_lib_name} ({word_count:,} words) "
            f"— always available, no upload needed."
        )
    else:
        effective_lib_text = None
        effective_lib_name = None
        st.markdown(
            "> 💡 **No strategy library loaded.** Claude will generate evidence-based strategies. "
            "To use your own library permanently, add a file named `strategy_library.docx` "
            "(or `.pdf`) to your GitHub repository — it will load automatically for everyone."
        )

    st.divider()

    # ── STEP 3: Report footer ─────────────────────────────────────────────────
    st.markdown("#### Step 3 — Report footer (optional)")
    cp, cc = st.columns(2)
    with cp: sr_prac    = st.text_input("Your name / title", key="sr_prac",
                                         placeholder="e.g. Janine Hogg — PBS Practitioner")
    with cc: sr_contact = st.text_input("Contact email", key="sr_contact",
                                         placeholder="e.g. janine@org.com.au")

    st.divider()

    # ── GENERATE ──────────────────────────────────────────────────────────────
    can_go  = bool(client_info and valid_behs and api_key)
    rec_btn = st.button("Generate strategy recommendations", type="primary",
                         disabled=not can_go, key="sr_gen")

    if not api_key:
        st.info("Enter your Anthropic API key in the sidebar to enable this tool.")
    elif not client_info or not valid_behs:
        st.info("Complete Step 1 to provide client information before generating.")

    if rec_btn and can_go:
        library_text  = effective_lib_text
        lib_label     = effective_lib_name
        client_label  = client_info.get("name", "client")
        is_freetext   = bool(valid_behs and valid_behs[0].get("_freetext"))

        if is_freetext:
            spinner_msg = (
                f"Reading your description and matching strategies for {client_label}…"
                if library_text else
                f"Reading your description and developing strategies for {client_label}…"
            )
        else:
            spinner_msg = (
                f"Matching strategies from your library for {client_label}…"
                if library_text else
                f"Generating evidence-based strategies for {client_label}…"
            )

        with st.spinner(spinner_msg):
            try:
                if is_freetext:
                    result = recommend_from_freetext(
                        client_info, freetext_value, api_key, library_text)
                else:
                    result = recommend_strategies(client_info, valid_behs, api_key, library_text)
            except Exception as e:
                st.error(f"Could not generate recommendations: {e}"); st.stop()

        source_note = f" (matched from **{lib_label}**)" if lib_label else ""
        st.success(f"✅  Recommendations generated for **{client_label}**{source_note}")
        st.divider()

        if result.get("client_summary"):
            st.info(f"**Clinical summary:** {result['client_summary']}")

        gen_s = result.get("general_strategies", [])
        if gen_s:
            st.markdown("##### General strategies — apply across all behaviours")
            for s in gen_s: st.markdown(f"- {s}")

        for beh in result.get("behaviours", []):
            st.divider()
            st.markdown(f"### {beh.get('behaviour','Behaviour')}")
            if beh.get("likely_function"):
                st.markdown(
                    f"<div style='background:#FDEBD0;padding:8px 12px;border-radius:6px;"
                    f"border-left:4px solid #D4700A;margin-bottom:10px'>"
                    f"<strong style='color:#D4700A'>Likely function:</strong> "
                    f"<span style='color:#1A2B35'>{beh['likely_function']}</span></div>",
                    unsafe_allow_html=True)
            cl, cr = st.columns(2)
            with cl:
                if beh.get("proactive"):
                    st.markdown("<div style='background:#DFFAE9;padding:6px 10px;border-radius:4px;"
                                "margin-bottom:6px'><strong style='color:#27AE60'>Proactive strategies"
                                "</strong></div>", unsafe_allow_html=True)
                    for s in beh["proactive"]: st.markdown(f"- {s}")
                if beh.get("teach_instead"):
                    st.markdown("<div style='background:#DCEEF8;padding:6px 10px;border-radius:4px;"
                                "margin:10px 0 6px'><strong style='color:#1A5276'>Skills to build"
                                "</strong></div>", unsafe_allow_html=True)
                    for s in beh["teach_instead"]: st.markdown(f"- {s}")
            with cr:
                if beh.get("reactive"):
                    st.markdown("<div style='background:#E0F6F3;padding:6px 10px;border-radius:4px;"
                                "margin-bottom:6px'><strong style='color:#1A9B8A'>Reactive strategies"
                                "</strong></div>", unsafe_allow_html=True)
                    for s in beh["reactive"]: st.markdown(f"- {s}")
                if beh.get("avoid"):
                    st.markdown("<div style='background:#FADBD8;padding:6px 10px;border-radius:4px;"
                                "margin:10px 0 6px'><strong style='color:#C0392B'>DO NOT"
                                "</strong></div>", unsafe_allow_html=True)
                    for s in beh["avoid"]: st.markdown(f"- {s}")

        st.divider()
        with st.spinner("Building PDF report…"):
            try:
                report_buf = generate_strategy_report(
                    result, client_label,
                    (sr_prac or "").strip(), (sr_contact or "").strip())
            except Exception as e: st.error(f"PDF error: {e}"); st.stop()

        st.download_button(
            "📥 Download strategy report (PDF)", report_buf,
            f"{client_label.replace(' ','_')}_Strategy_Report.pdf",
            "application/pdf", use_container_width=True)

    st.divider()
    st.caption(
        "Recommendations are AI-generated. Always review before implementing. "
        "When a strategy library is uploaded, Claude selects from that document — "
        "verify that selected strategies are appropriate for this individual. "
        "This tool complements but does not replace a formal Behaviour Support Plan."
    )


# ══════════════════════════════════════════════════════════════════════════════
# TAB 4 — NDIS BSP DRAFT
# ══════════════════════════════════════════════════════════════════════════════
with tab4:
    st.markdown(
        "Generate a draft **NDIS Version 3 Behaviour Support Plan** using data already extracted "
        "in the other tabs. The output is a formatted **Word document** structured to the NDIS "
        "template — covering crisis strategy, regulated restrictive practices, RP protocols, and "
        "implementation and training."
    )
    st.warning(
        "**Clinical review required.** All AI-generated content must be reviewed and approved by a "
        "qualified Behaviour Support Practitioner before implementation or NDIS submission. "
        "Regulated restrictive practices require NDIS authorisation — this tool generates a draft only."
    )
    st.divider()

    # ── Data from other tabs ───────────────────────────────────────────────────
    st.markdown("#### Client data from other tabs")
    t0 = st.session_state.get("t0_data")
    t1 = st.session_state.get("t1_data")

    if not t0 and not t1:
        st.error(
            "No client data available. Complete at least one of the following first:\n\n"
            "- **Tab 1** — Upload a 30-day behaviour recording\n"
            "- **Tab 2** — Upload a PBSP to extract client information"
        )
    else:
        col1, col2 = st.columns(2)
        with col1:
            if t1:
                st.success(f"✅ **PBSP data** — {t1.get('name','')} ({t1.get('age_info','')})")
            else:
                st.info("No PBSP data (Tab 2 not yet used)")
        with col2:
            if t0:
                st.success(f"✅ **Behaviour recording** — {t0.get('total_incidents','?')} incidents, "
                           f"{len(t0.get('behaviours',[]))} behaviour(s)")
            else:
                st.info("No recording data (Tab 1 not yet used)")

        bsp_client_name = (t1.get("name") if t1 else None) or \
                          (t0.get("client_name") if t0 else None) or "Client"

        st.divider()

        # ── Plan type ──────────────────────────────────────────────────────────
        st.markdown("#### Plan type")
        plan_type = st.radio(
            "Select NDIS BSP type",
            ["Comprehensive", "Interim"],
            horizontal=True, key="bsp_plan_type",
            help="Comprehensive — full multi-element plan. Interim — shorter, for new participants "
                 "or urgent situations."
        )

        st.divider()

        # ── Restrictive practices ──────────────────────────────────────────────
        st.markdown("#### Regulated restrictive practices")
        st.caption("Select all that are currently in use or being proposed for this person.")
        rp_options = [
            "Physical restraint",
            "Mechanical restraint",
            "Chemical restraint (medication as behaviour support strategy)",
            "Seclusion",
            "Environmental restraint",
        ]
        selected_rps = []
        c1, c2 = st.columns(2)
        for i, rp in enumerate(rp_options):
            with (c1 if i % 2 == 0 else c2):
                if st.checkbox(rp, key=f"rp_{i}"):
                    selected_rps.append(rp)
        if not selected_rps:
            st.info("No regulated restrictive practices selected — the RP section will be omitted.")

        st.divider()

        # ── Additional context ─────────────────────────────────────────────────
        st.markdown("#### Additional context (optional)")
        bsp_extra = st.text_area(
            "Anything the tool doesn't have from the other tabs",
            key="bsp_extra", height=100,
            placeholder="e.g. Relevant history, recent incidents, current medications prescribed "
                        "for behaviour support, participant/family preferences, living situation..."
        )

        # ── Practitioner details ───────────────────────────────────────────────
        st.markdown("#### Practitioner details")
        bp1, bp2 = st.columns(2)
        with bp1:
            bsp_prac = st.text_input("Your name / title", key="bsp_prac",
                                      placeholder="e.g. Janine Hogg — PBS Practitioner")
        with bp2:
            bsp_contact = st.text_input("Contact email", key="bsp_contact",
                                         placeholder="e.g. janine@org.com.au")

        st.divider()

        # ── Generate ───────────────────────────────────────────────────────────
        bsp_btn = st.button(
            "Generate BSP draft", type="primary",
            disabled=not api_key, key="bsp_gen"
        )
        if not api_key:
            st.info("Enter your Anthropic API key in the sidebar to enable this tool.")

        if bsp_btn and api_key:
            client_data_text = format_data_for_bsp(t0_data=t0, t1_data=t1)
            with st.spinner(f"Drafting {plan_type} BSP for {bsp_client_name}…"):
                try:
                    bsp_result = generate_ndis_bsp(
                        client_data_text, plan_type.lower(),
                        selected_rps, bsp_extra or "", api_key
                    )
                    st.session_state["bsp_data"] = bsp_result
                except Exception as e:
                    st.error(f"Generation failed: {e}"); st.stop()
            st.success(f"✅ {plan_type} BSP drafted for **{bsp_client_name}**")

        if "bsp_data" in st.session_state:
            bsp_result = st.session_state["bsp_data"]
            st.divider()

            # On-screen preview
            if bsp_result.get("about_client"):
                preview = bsp_result["about_client"]
                st.info(f"**About {bsp_client_name}:** {preview[:280]}{'…' if len(preview)>280 else ''}")

            behs = bsp_result.get("behaviours_of_concern", [])
            if behs:
                st.markdown(f"**{len(behs)} behaviour(s) of concern identified**")
                for b in behs:
                    st.markdown(f"- **{b.get('name','')}:** {b.get('description','')}")

            crisis = bsp_result.get("crisis_strategy", {})
            if crisis.get("immediate_responses"):
                with st.expander("Preview crisis strategy"):
                    st.markdown("**Immediate responses:**")
                    for r in crisis["immediate_responses"]: st.markdown(f"- {r}")
                    if crisis.get("when_to_call_emergency"):
                        st.markdown(f"**When to call 000:** {crisis['when_to_call_emergency']}")

            rps = bsp_result.get("restrictive_practices", [])
            if rps:
                with st.expander(f"Preview regulated restrictive practices — {len(rps)} included"):
                    for rp in rps:
                        st.markdown(f"**{rp.get('type','')}** — {rp.get('behaviour_addressed','')}")
                        st.markdown(f"*{rp.get('justification','')}*")
                        if rp.get("fading_plan"):
                            st.markdown(f"Fading plan: {rp['fading_plan']}")
                        st.markdown("---")

            impl = bsp_result.get("implementation", {})
            training = impl.get("training_requirements", [])
            if training:
                with st.expander(f"Preview training requirements — {len(training)} items"):
                    for t_item in training:
                        st.markdown(f"- **{t_item.get('topic','')}** "
                                    f"({t_item.get('who','')}) — {t_item.get('timeframe','')}")

            st.divider()

            with st.spinner("Building Word document…"):
                try:
                    docx_buf = create_bsp_docx(
                        bsp_result, plan_type, bsp_client_name,
                        (bsp_prac or "").strip(), (bsp_contact or "").strip()
                    )
                except Exception as e:
                    st.error(f"Document error: {e}"); st.stop()

            safe = bsp_client_name.replace(" ", "_")
            st.download_button(
                f"📥 Download {plan_type} BSP draft (Word)", docx_buf,
                f"{safe}_{plan_type}_BSP_Draft.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

            if st.button("🗑 Clear BSP draft", key="bsp_clear"):
                del st.session_state["bsp_data"]; st.rerun()

    st.divider()
    st.caption(
        "This tool generates a draft BSP to assist the Behaviour Support Practitioner. "
        "It does not replace clinical assessment, professional judgment, or the NDIS requirement "
        "for authorisation of regulated restrictive practices. "
        "Always review all content before implementing strategies or submitting to the NDIS."
    )


# ══════════════════════════════════════════════════════════════════════════════
# TAB 5 — FUNCTIONAL BEHAVIOUR ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════════
with tab5:
    st.markdown(
        "Upload your completed assessment score reports and this tool will extract the scores, "
        "write interpretive narratives for each tool, and generate **Section 8: Functional "
        "Behaviour Assessment** as a formatted Word document — ready to drop into your "
        "Comprehensive PBSP."
    )
    st.info(
        "**Workflow:** Complete the FBA tools → upload score reports here → review and download "
        "Section 8 → use this in your Comprehensive BSP (Tab 4). "
        "The Interim BSP does not include an FBA section — it uses hypothesis language only."
    )
    st.divider()

    # ── Client info ────────────────────────────────────────────────────────────
    st.markdown("#### Client information")
    t1_fba = st.session_state.get("t1_data")
    t0_fba = st.session_state.get("t0_data")

    fba_client_name = (t1_fba.get("name") if t1_fba else None) or \
                      (t0_fba.get("client_name") if t0_fba else None) or ""

    fc1, fc2 = st.columns(2)
    with fc1:
        fba_name = st.text_input(
            "Client name", key="fba_name",
            value=fba_client_name,
            placeholder="e.g. Izayah Mills"
        )
    with fc2:
        fba_prac = st.text_input(
            "Practitioner name / title", key="fba_prac",
            placeholder="e.g. Janine Hogg — Behaviour Support Practitioner"
        )
    fba_contact = st.text_input(
        "Contact email", key="fba_contact",
        placeholder="e.g. janine@ndaffirming.com.au"
    )

    # Build client context for prompt
    fba_client_context = format_data_for_bsp(t0_data=t0_fba, t1_data=t1_fba) \
                         if (t0_fba or t1_fba) else ""

    st.divider()

    # ── Tool uploads ───────────────────────────────────────────────────────────
    st.markdown("#### Assessment tool score reports")
    st.caption(
        "Upload the completed score report PDF for each tool. "
        "You can upload multiple files at once — Claude will identify which tool each belongs to."
    )

    fba_files = st.file_uploader(
        "Upload assessment score reports (PDF)",
        type=["pdf"],
        accept_multiple_files=True,
        key="fba_uploads",
        help="EDA-Q, QABF, Sensory Profile 2.0, FAST, TMQ — upload one or all together"
    )

    # Show which tools are expected
    st.markdown("**Tools typically included:**")
    tool_cols = st.columns(3)
    tools_list = [
        ("EDA-Q", "Extreme Demand Avoidance Questionnaire"),
        ("QABF", "Questions About Behavioural Function"),
        ("Sensory Profile 2.0", "Sensory Profile 2.0"),
        ("FAST", "Functional Analysis Screening Tool"),
        ("TMQ", "Monotropism Questionnaire"),
        ("ABAS", "Adaptive Behaviour Assessment System (optional)"),
    ]
    for i, (short, full) in enumerate(tools_list):
        with tool_cols[i % 3]:
            st.caption(f"**{short}** — {full}")

    st.divider()

    # ── Direct assessment methods ──────────────────────────────────────────────
    st.markdown("#### Direct assessment methods")
    st.caption("Describe the direct observations and methods used (will appear in the Overview section).")
    fba_direct = st.text_area(
        "Direct assessment methods",
        key="fba_direct",
        height=80,
        placeholder=(
            "e.g. Behaviour Support Practitioner observational assessment conducted across "
            "in-home and telehealth sessions; review of practitioner session notes and "
            "tracking data recorded by support workers..."
        )
    )

    st.divider()

    # ── Practitioner declaration ───────────────────────────────────────────────
    fba_btn = st.button(
        "Extract scores and generate Section 8",
        type="primary",
        disabled=not (fba_files and api_key),
        key="fba_gen"
    )

    if not api_key:
        st.info("Enter your Anthropic API key in the sidebar to enable this tool.")
    elif not fba_files:
        st.info("Upload at least one assessment score report PDF to continue.")

    if fba_btn and fba_files and api_key:
        # Extract text from all uploaded PDFs
        pdf_texts = []
        with st.spinner("Reading uploaded score reports…"):
            for f in fba_files:
                try:
                    text = extract_text_from_pdf(f.read())
                    if text.strip():
                        pdf_texts.append(f"[FILE: {f.name}]\n{text}")
                    else:
                        st.warning(f"{f.name} — could not extract text. "
                                   "Try re-saving as a text-based PDF.")
                except Exception as e:
                    st.warning(f"{f.name} — read error: {e}")

        if not pdf_texts:
            st.error("No readable text found in the uploaded files.")
            st.stop()

        # Inject direct methods into context if provided
        ctx = fba_client_context
        if fba_direct and fba_direct.strip():
            ctx = f"DIRECT ASSESSMENT METHODS USED:\n{fba_direct}\n\n{ctx}"

        with st.spinner(f"Extracting scores and writing FBA narratives for "
                        f"{fba_name or 'client'}… (this takes 30–60 seconds)"):
            try:
                fba_result = extract_and_interpret_fba(pdf_texts, ctx, api_key)
                # Inject client name if not found in PDFs
                if not fba_result.get("client_name") and fba_name:
                    fba_result["client_name"] = fba_name
                st.session_state["fba_data"] = fba_result
            except Exception as e:
                st.error(f"FBA generation failed: {e}")
                st.stop()

        tools_found = fba_result.get("tools_identified", [])
        st.success(
            f"✅ FBA Section 8 generated for **{fba_name or 'client'}** — "
            f"{len(tools_found)} tool(s) identified: {', '.join(tools_found)}"
        )

    if "fba_data" in st.session_state:
        fba_result = st.session_state["fba_data"]
        display_name = fba_name or fba_result.get("client_name") or "Client"
        st.divider()

        # On-screen preview
        results = fba_result.get("results", {})

        # Score summaries
        score_preview = []
        if results.get("edaq", {}).get("present"):
            score_preview.append(f"**EDA-Q:** {results['edaq'].get('total_score','—')}")
        if results.get("qabf", {}).get("present"):
            behs = results["qabf"].get("behaviours", [])
            for b in behs:
                sc = b.get("scores", {})
                score_preview.append(
                    f"**QABF ({b.get('behaviour_name','Behaviour')}):** "
                    f"Attn {sc.get('attention','—')} | "
                    f"Escape {sc.get('escape','—')} | "
                    f"Non-social {sc.get('non_social','—')} | "
                    f"Tangible {sc.get('tangible','—')}"
                )
        if results.get("fast", {}).get("present"):
            sc = results["fast"].get("scores", {})
            score_preview.append(
                f"**FAST:** Social-connection {sc.get('social_connection','—')} | "
                f"Social-autonomy {sc.get('social_autonomy','—')} | "
                f"Auto-sensory {sc.get('automatic_sensory','—')}"
            )
        if results.get("tmq", {}).get("present"):
            score_preview.append(
                f"**TMQ overall average:** "
                f"{results['tmq'].get('factors',{}).get('overall_average','—')}"
            )

        if score_preview:
            st.markdown("##### Extracted scores")
            for s in score_preview:
                st.markdown(f"- {s}")

        # Function of presentation preview
        fop = fba_result.get("function_of_presentation", {})
        if fop.get("summary_statement"):
            with st.expander("Preview — Function of Presentation summary"):
                st.markdown(fop["summary_statement"])

        st.divider()

        # Build and download Word doc
        with st.spinner("Building Word document…"):
            try:
                fba_docx = create_fba_docx(
                    fba_result,
                    display_name,
                    (fba_prac or "").strip(),
                    (fba_contact or "").strip()
                )
            except Exception as e:
                st.error(f"Document error: {e}")
                st.stop()

        safe = display_name.replace(" ", "_")
        st.download_button(
            "📥 Download Section 8 — FBA (Word)",
            fba_docx,
            f"{safe}_FBA_Section8.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.success(
            "FBA data is now stored — switch to **Tab 4 (NDIS BSP Draft)** and select "
            "Comprehensive to generate a plan that uses these findings."
        )

        if st.button("🗑 Clear FBA data", key="fba_clear"):
            del st.session_state["fba_data"]; st.rerun()

    st.divider()
    st.caption(
        "Assessment score reports are processed in memory only and never stored permanently. "
        "All AI-generated narratives must be reviewed by the Behaviour Support Practitioner "
        "before inclusion in a Comprehensive PBSP. "
        "The FBA tab generates Section 8 only — use Tab 4 to generate the full Comprehensive plan."
    )
