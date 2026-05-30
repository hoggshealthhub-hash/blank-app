"""
PBS Support Tool
Tab 1: Upload a PBSP → Support Reference Card + ABC Recording Form.
Tab 2: Enter client behaviours → AI strategy recommendations + PDF report.
"""

import json
from datetime import date
import streamlit as st
import anthropic
from io import BytesIO
from docx import Document as DocxDocument
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.colors import HexColor, white
from reportlab.pdfbase.pdfmetrics import stringWidth

st.set_page_config(page_title="PBS Support Tool", page_icon="📋", layout="centered")

DEEP_BLUE  = HexColor('#0D4F6E')
TEAL       = HexColor('#1A9B8A')
BLUE2      = HexColor('#1A5276')
GREEN      = HexColor('#27AE60')
RED        = HexColor('#C0392B')
AMBER      = HexColor('#D4700A')
LIGHT_TEAL = HexColor('#E0F6F3')
LIGHT_BLUE = HexColor('#DCEEF8')
LIGHT_GRN  = HexColor('#DFFAE9')
LIGHT_RED  = HexColor('#FADBD8')
LIGHT_AMBR = HexColor('#FDEBD0')
MID_GREY   = HexColor('#CCCCCC')
DARK_TEXT  = HexColor('#1A2B35')
MED_TEXT   = HexColor('#4A7C8E')

W, H = A4
LM, RM = 20, 20
CW = W - LM - RM

# ── Embedded strategy library ──────────────────────────────────────────────────
EMBEDDED_LIBRARY = """Hoggs Health Hub PBS Strategy Library — Neuro-Affirming, Trauma-Informed, Rights-Based, NDIS-Aligned

S01 Functional Behaviour Assessment FBA: Tier 1 Established Evidence. All populations. The systematic multi-method process of understanding why a behaviour occurs. Behaviours: Physical Aggression, Self-Injurious Behaviour, Property Destruction, Elopement, Verbal Aggression, Task Refusal, Emotional Dysregulation, Stereotypy, PICA, Disruptive Behaviour. Functions: All. Step 1: File review including previous assessments, medical history, therapy reports. Step 2: Structured stakeholder interviews using FAI or open clinical interviewing including the person. Step 3: Operationally define behaviour in observable measurable terms. Step 4: Collect ABC data across multiple environments, times, and people. Step 5: Identify antecedent patterns, consequence patterns, and setting events. Step 6: Develop hypothesis: When antecedent/context, person does behaviour, which results in function. Step 7: Share hypothesis with person and family. Step 8: Use FBA to directly drive every strategy in BSP.

S02 Functional Communication Training FCT: Tier 1 Established Evidence. All populations. Teaches a person to communicate the same message as their behaviour using a more effective accessible form. Communication rights framework. Behaviours: Physical Aggression, Self-Injurious Behaviour, Property Destruction, Verbal Aggression, Elopement, Task Refusal, Emotional Dysregulation, Disruptive Behaviour. Functions: Escape/Avoidance, Attention, Access to Tangibles, Automatic/Sensory. Step 1: Confirm function via FBA — replacement communication must serve exact same function. Step 2: Collaborate with Speech Pathologist to identify appropriate communication form. Step 3: Select form that is efficient and as easy as the behaviour. Step 4: Teach in calm low-demand conditions. Step 5: Honour communication immediately and consistently every time. This is non-negotiable. Step 6: Train all support people across all environments. Step 7: Monitor both behaviour and replacement communication as parallel data streams.

S03 Visual Supports and Predictability Systems: Tier 1 Established Evidence. All populations. Schedules, now/next boards, task breakdowns, visual timers, choice boards. Behaviours: Task Refusal, Emotional Dysregulation, Elopement, Disruptive Behaviour, Physical Aggression, Verbal Aggression, Demand Avoidance PDA. Functions: Escape/Avoidance, Automatic/Sensory. Step 1: Assess where unpredictability or transitions drive distress. Step 2: Determine format matched to person's profile. Step 3: Build schedules collaboratively using person's preferred images. Step 4: Make schedules interactive. Step 5: Use visual timers to signal transitions before they occur. Step 6: Introduce a change symbol to build tolerance for variation. Step 7: Ensure cross-setting consistency.

S04 Environmental Modification and Sensory Environment Design: Tier 1 Established Evidence. All populations. Changing physical, social, or sensory environment to reduce triggering conditions. Behaviours: Physical Aggression, Self-Injurious Behaviour, Property Destruction, Emotional Dysregulation, Elopement, Disruptive Behaviour, Stereotypy. Functions: Escape/Avoidance, Automatic/Sensory, Attention. Step 1: Ecological assessment across all environments. Step 2: Identify sensory contributors: noise, lighting, visual clutter, temperature. Step 3: Modify: dim lighting, acoustic panels, noise-reducing headphones, flexible seating. Step 4: Examine social environment and reduce demand density. Step 5: Ensure enriched environment with continuous access to preferred items. Step 6: Create or designate a regulation space. Step 7: Involve person in designing their environment.

S05 Antecedent Mapping and Trigger Modification: Tier 1 Established Evidence. All populations. Systematically identifying events that precede behaviour then modifying those conditions. Behaviours: Physical Aggression, Self-Injurious Behaviour, Property Destruction, Task Refusal, Verbal Aggression, Elopement. Functions: Escape/Avoidance, Attention, Access to Tangibles. Step 1: Map antecedents using ABC data. Step 2: Identify setting events: disrupted sleep, illness, medication changes, skipped meals. Step 3: Develop setting event protocol with reduced demands for high-vulnerability days. Step 4: Create shift-to-shift communication system for setting event status. Step 5: Remove or modify identified antecedents. Step 6: Document all modifications trialled. Step 7: Balance trigger removal with concurrent teaching strategies.

S06 Choice-Making and Autonomy Architecture: Tier 1 Established Evidence. All populations. Embedding meaningful genuine opportunities for agency and control throughout the day. Behaviours: Task Refusal, Physical Aggression, Verbal Aggression, Elopement, Demand Avoidance PDA, Emotional Dysregulation. Functions: Escape/Avoidance, Access to Tangibles. Step 1: Audit person's current day for genuine choices available. Step 2: Embed micro-choices throughout the day. Step 3: Ensure choices are genuine and honoured. Step 4: Involve person in planning their own schedule, goals, and activities. Step 5: For PDA profiles: reframe demands as invitations using indirect language: I wonder if, Would you like to help me with, You might be interested to know. Step 6: Provide choice in how tasks are completed.

S07 Collaborative Problem Solving CPS: Tier 2 Moderate Evidence. All populations. Ross Greene model. Behaviour reflects lagging skills not wilful defiance. Behaviours: Task Refusal, Physical Aggression, Verbal Aggression, Demand Avoidance PDA, Emotional Dysregulation. Functions: Escape/Avoidance, Attention. Step 1: Identify specific unsolved problems during calm states. Step 2: Empathy step: open with genuine curiosity not accusation. Step 3: Define adult concern around impact. Step 4: Invitation to generate mutually workable solutions addressing both concerns. Step 5: Evaluate and trial agreed solutions. Step 6: Follow up and document. Never use during crisis.

S08 Co-Regulation and Relational Safety: Tier 2 Moderate Evidence. All populations. Calm attuned support person uses regulated nervous system to help another regulate. Polyvagal theory and attachment science. Behaviours: Emotional Dysregulation, Physical Aggression, Verbal Aggression, Elopement, Self-Injurious Behaviour. Functions: Attention, Escape/Avoidance, Automatic/Sensory. Step 1: Prioritise relationship-building with non-demand interest-led time. Step 2: Train support people in low arousal attuned presence. Step 3: Model regulation — staff anxiety transmits directly to participant. Step 4: Use serve and return interaction following person's lead. Step 5: During distress prioritise presence and safety over problem-solving. Step 6: Avoid explanations during or immediately after dysregulated episodes. Step 7: When regulated connect warmly before resuming routines.

S09 Low Arousal Approach: Tier 2 Moderate Evidence. All populations. Andrew McDonnell Studio3. Reducing environmental and interpersonal arousal. Behaviours: Physical Aggression, Verbal Aggression, Emotional Dysregulation, Property Destruction, Elopement, Self-Injurious Behaviour. Functions: Escape/Avoidance, Attention, Automatic/Sensory. Step 1: Train all support people: reduce demands, avoid confrontation, prioritise relational safety over compliance. Step 2: When behaviour emerging: slow down, step back, soften voice, reduce eye contact, remove demands. Step 3: Use minimal calm verbal communication: short, clear, warm statements. Step 4: Remove audience. Step 5: Withdraw from confrontational interactions without abandonment. Step 6: After resolution reconnect gently without revisiting incident immediately. Step 7: Debrief with team after significant incidents.

S10 Sensory Diet and Occupational Therapy Regulation Programme: Tier 2 Moderate Evidence. All populations. Individualised scheduled programme of sensory activities designed by OT. Behaviours: Physical Aggression, Self-Injurious Behaviour, Property Destruction, Emotional Dysregulation, Stereotypy, Disruptive Behaviour. Functions: Automatic/Sensory, Escape/Avoidance. Step 1: Obtain comprehensive sensory assessment from OT using SPM-2 or Sensory Profile 2. Step 2: Identify person's sensory profile across all systems. Step 3: Design daily schedule of sensory activities matched to profile. Step 4: Embed activities at natural transition points. Step 5: Train all support people consistently. Step 6: Ensure continuous access to self-directed sensory tools. Step 7: Review quarterly.

S11 Interoception Awareness Training: Tier 2 Moderate Evidence. All populations. Kelly Mahler curriculum. Teaching recognition of internal body signals. Behaviours: Emotional Dysregulation, Self-Injurious Behaviour, Physical Aggression, Verbal Aggression. Functions: Automatic/Sensory, Escape/Avoidance. Step 1: Introduce body awareness during calm fun states. Step 2: Use Kelly Mahler Interoception Curriculum. Step 3: Create personalised body maps. Step 4: Practise naming sensations in low-stakes activities. Step 5: Build personal body signal glossary. Step 6: Teach support people to prompt body check-ins. Step 7: Progress slowly led by person.

S12 Zones of Regulation: Tier 2 Moderate Evidence. All populations. Leah Kuypers 2011. Blue: low arousal. Green: optimal regulated. Yellow: heightened alert. Red: crisis. Behaviours: Emotional Dysregulation, Physical Aggression, Verbal Aggression, Self-Injurious Behaviour. Functions: Escape/Avoidance, Attention, Automatic/Sensory. Step 1: Introduce zones through engaging teaching during calm states. Step 2: Create personalised zones reference chart with person's own examples. Step 3: Teach personalised toolbox for each zone. Step 4: Practise identifying zones in low-stakes moments. Step 5: Introduce concepts of triggers and tools. Step 6: Embed zones check-ins at natural transition points. Step 7: Build portable physical regulation toolkit.

S13 Social Stories and Narrative Supports: Tier 2 Moderate Evidence. All populations. Carol Gray 2010. Short personalised narratives from person's perspective. Behaviours: Task Refusal, Emotional Dysregulation, Physical Aggression, Elopement, Disruptive Behaviour. Functions: Escape/Avoidance, Automatic/Sensory. Step 1: Identify specific situation generating difficulty. Step 2: Write from person's first-person perspective in positive non-judgmental tone. Step 3: Follow Gray's sentence ratio: primarily descriptive, some perspective, some coaching, minimal directive. Step 4: Keep brief and visually supported. Step 5: Introduce 1-2 days before relevant situation. Step 6: Involve person in reviewing and editing. Step 7: Never use as correction or post-incident lecture.

S14 DIR Floortime: Tier 2 Moderate Evidence. All populations. Stanley Greenspan and Serena Wieder. Developmental relationship-based framework prioritising following person's lead. Behaviours: Emotional Dysregulation, Physical Aggression, Self-Injurious Behaviour, Task Refusal, Disruptive Behaviour. Functions: Attention, Escape/Avoidance, Automatic/Sensory. Step 1: Training in DIR framework required. Step 2: Follow person's lead rather than redirecting. Step 3: Build circles of communication around person's interests. Step 4: Match person's affect: be present warm playful and genuinely engaged. Step 5: Respect self-directed sensory and regulatory behaviour. Step 6: Identify functional emotional developmental level. Step 7: Implement across daily routines.

S15 Safe Retreat Space and Regulation Environment: Tier 2 Moderate Evidence. All populations. Designated person-designed area for voluntary self-initiated access to regulate. Never used as consequence. Fundamentally distinct from seclusion which is a regulated restrictive practice. Behaviours: Emotional Dysregulation, Physical Aggression, Self-Injurious Behaviour, Elopement. Functions: Escape/Avoidance, Automatic/Sensory. Step 1: Design space collaboratively with person. Step 2: Include person's regulatory items: weighted items, soft textures, adjustable lighting, music, movement tools. Step 3: Make access completely voluntary and unconditional. Step 4: Establish communication signal for needing space. Step 5: Support people must not follow person in unless immediate safety risk. Step 6: Create simple signal for when person is ready to re-engage. Step 7: Never remove access as consequence.

S16 PACE Model Playfulness Acceptance Curiosity Empathy: Tier 2 Moderate Evidence. All populations. Dan Hughes Dyadic Developmental Psychotherapy. Relational stance. Behaviours: Emotional Dysregulation, Physical Aggression, Verbal Aggression, Task Refusal, Self-Injurious Behaviour. Functions: Attention, Escape/Avoidance. Step 1: Playfulness: bring warmth, lightness, and joy. Step 2: Acceptance: unconditional positive regard for person separate from behaviour. Step 3: Curiosity: non-judgmental curiosity using I wonder what was going on for you. Step 4: Empathy: reflect emotional experience with compassion without fixing. Step 5: Use PACE in all interactions consistently. Step 6: Train all team members. Step 7: Reflect in supervision on own PACE practice.

S17 Demand Reduction and Low Demand Approach PDA-Informed: Tier 2 Moderate Evidence. All populations. Fundamentally restructuring nature, language, and volume of demands. Behaviours: Task Refusal, Physical Aggression, Verbal Aggression, Elopement, Demand Avoidance PDA, Emotional Dysregulation. Functions: Escape/Avoidance. Step 1: Conduct demand audit to identify which demands are truly essential. Step 2: Replace direct instructions with indirect collaborative language: I wonder if, Would you mind helping me, You might be interested to know. Step 3: Frame participation as optional and interest-led wherever possible. Step 4: Allow person to lead sequence, timing, and format of non-essential activities. Step 5: Avoid negotiating during dysregulated states. Step 6: Educate all stakeholders in PDA-informed approaches. Step 7: Build shared written understanding of non-negotiable demands.

S18 DBT-Informed Skills Emotional Regulation and Distress Tolerance: Tier 2 Moderate Evidence. All populations. Marsha Linehan. Emotional regulation, distress tolerance, mindfulness, interpersonal effectiveness. Behaviours: Emotional Dysregulation, Self-Injurious Behaviour, Verbal Aggression, Physical Aggression. Functions: Escape/Avoidance, Automatic/Sensory, Attention. Step 1: Formal DBT delivery by qualified DBT therapist with clinical supervision. Step 2: Mindfulness: brief sensory-based activities. Step 3: Emotional regulation: identify and name emotions, check the facts. Step 4: Distress tolerance TIPP skills: Temperature cold water, Intense exercise, Paced breathing, Paired muscle relaxation. Step 5: STOP skill: Stop, Take a step back, Observe, Proceed mindfully. Step 6: Build personalised distress tolerance toolkit. Step 7: Adapt all skills to person's profile.

S19 Matched Sensory Alternatives Sensory Substitution: Tier 2 Moderate Evidence. All populations. Providing access to items producing same sensory feedback as behaviour. Behaviours: Self-Injurious Behaviour, Stereotypy, PICA, Property Destruction. Functions: Automatic/Sensory. Step 1: Confirm automatic/sensory function via FBA. Step 2: Sensory analysis with OT. Step 3: Identify matched alternatives providing similar sensory input. Step 4: Assess preference. Step 5: Ensure continuous unconditional access. Step 6: Teach person to access alternatives independently. Step 7: Monitor and rotate as needed.

S20 Enriched Environment and Access to Preferred Activities: Tier 1 Established Evidence. All populations. Continuous access to varied preferred motivating items, activities, and social interaction. Behaviours: Self-Injurious Behaviour, Stereotypy, Physical Aggression, Disruptive Behaviour, PICA, Emotional Dysregulation. Functions: Automatic/Sensory, Attention, Access to Tangibles. Step 1: Conduct thorough preference assessment. Step 2: Identify preferred items, activities, sensory experiences, social interactions. Step 3: Ensure high-preference items freely and continuously accessible throughout day. Step 4: Provide variety by rotating preferred items. Step 5: Ensure social environment is enriched. Step 6: Include activities at multiple sensory levels. Step 7: Reassess and update regularly.

S21 Transition Supports and Change Preparation: Tier 2 Moderate Evidence. All populations. Proactive strategies preparing person for changes in activity, location, or routine. Behaviours: Emotional Dysregulation, Physical Aggression, Verbal Aggression, Task Refusal, Elopement. Functions: Escape/Avoidance, Automatic/Sensory. Step 1: Identify challenging transitions through ABC data. Step 2: Build transition preparation with warnings at 10, 5, and 2 minutes before. Step 3: Use visual timer person can independently reference. Step 4: Provide now/next board. Step 5: For new environments provide advance visual preview. Step 6: Ensure transition destination contains something appealing. Step 7: Practise difficult transitions in low-stakes conditions.

S22 Anxiety Management in Neurodivergent Presentations: Tier 2 Moderate Evidence. All populations. Addressing significant co-occurrence of anxiety — estimated 40-50 percent of autistic people. Behaviours: Emotional Dysregulation, Task Refusal, Demand Avoidance PDA, Elopement, Physical Aggression, Verbal Aggression. Functions: Escape/Avoidance, Automatic/Sensory. Step 1: Assess anxiety comprehensively. Step 2: Psychoeducation using appropriate language. Step 3: Somatic tools: paced breathing, cold water on wrists, bilateral stimulation, movement. Step 4: Cognitive approaches adapted for neurodivergent processing: thought testing, worry time, externalising anxiety. Step 5: Environmental reduction of uncertainty using visual supports. Step 6: Gradual exposure with psychologist where anxiety is phobia-specific. Step 7: Build personalised anxiety management plan.

S23 Strengths-Based and Interests-Led Engagement: Tier 2 Moderate Evidence. All populations. Using person's genuine interests as primary vehicle for engagement. Behaviours: Task Refusal, Demand Avoidance PDA, Emotional Dysregulation, Disruptive Behaviour, Physical Aggression. Functions: Escape/Avoidance, Attention. Step 1: Conduct comprehensive strengths and interests assessment. Step 2: Map interests across domains. Step 3: Use interests as entry point for all skill-building. Step 4: Allow extended time with interests as regulation tool. Step 5: Build support relationship through shared engagement. Step 6: For PDA profiles interest-led engagement is often the only effective entry point. Step 7: Document strengths and interests prominently in BSP.

S24 Mindfulness-Based Approaches Adapted for Neurodivergent People: Tier 2 Moderate Evidence. All populations. Movement-based, sensory-grounded, brief practices. Behaviours: Emotional Dysregulation, Self-Injurious Behaviour, Verbal Aggression, Physical Aggression. Functions: Escape/Avoidance, Automatic/Sensory. Step 1: Assess sensory and attentional profile. Step 2: Begin with externally focused sensory mindfulness: notice five things you can see. Step 3: Introduce brief structured practices: 30-second grounding exercises. Step 4: Use movement as mindfulness vehicle. Step 5: Use visual or audio aids. Step 6: Teach during calm positive conditions. Step 7: Build personalised mindfulness micro-practice.

S25 Medical and Physical Health Review: Tier 1 Established Evidence. All populations. Examining whether behaviour is driven by pain, medical conditions, GI disturbance, sleep disorders, hormonal factors, medication side effects. Behaviours: Self-Injurious Behaviour, Physical Aggression, Verbal Aggression, Emotional Dysregulation, Task Refusal, Disruptive Behaviour. Functions: Automatic/Sensory, Escape/Avoidance. Step 1: Include medical review as mandatory component of every BSP. Step 2: Map behaviour patterns against medical cycles. Step 3: Collaborate with GP or specialist and share PBS data. Step 4: Review current medications for side effects and interactions. Step 5: Assess sleep comprehensively. Step 6: Address dental health proactively. Step 7: Reassess behaviour following medical treatment.

S26 Crisis and Safety Planning: Tier 1 Established Evidence. All populations. Clear step-by-step guidance for recognising escalation, responding safely, and post-crisis recovery. Behaviours: Physical Aggression, Self-Injurious Behaviour, Property Destruction, Emotional Dysregulation, Elopement. Functions: All functions. Step 1: Map escalation cycle with specific observable stages for this person. Step 2: For each stage develop specific guidance: what to do, what to say, what to avoid. Step 3: Early stage: reduce demands, increase space, offer preferred items, use low arousal language. Step 4: Escalation stage: remove audience, minimise verbal interaction, ensure environmental safety. Step 5: Crisis stage: ensure safety using minimum intervention necessary. Any physical intervention must be authorised as regulated restrictive practice. Step 6: Recovery stage: allow time, reconnect warmly with food, drink, or comfort. Step 7: Post-incident reflective debrief to update plan.

S27 AAC and Communication Rights Framework: Tier 1 Established Evidence. All populations. Every person has the right to communicate. No-tech, low-tech, and high-tech approaches. Behaviours: Physical Aggression, Self-Injurious Behaviour, Property Destruction, Verbal Aggression, Elopement, Task Refusal, Emotional Dysregulation. Functions: All functions. Step 1: Ensure AAC assessment by qualified Speech Pathologist with AAC expertise. Step 2: Design communication system based on assessment. Step 3: Implement aided language stimulation: support people model AAC throughout day. Step 4: Ensure AAC has robust vocabulary beyond requesting. Step 5: Respond to all communication attempts with warmth. Step 6: Never restrict access to AAC as consequence. Step 7: Ensure AAC training extends to all environments.

S28 Positive Programming and Lifestyle Redesign: Tier 2 Moderate Evidence. All populations. Comprehensively reviewing and enriching person's daily life. Behaviours: Physical Aggression, Self-Injurious Behaviour, Property Destruction, Emotional Dysregulation, Disruptive Behaviour, Stereotypy. Functions: All functions. Step 1: Conduct quality of life assessment. Step 2: Identify quality of life gaps. Step 3: Design enriched person-centred daily schedule collaboratively. Step 4: Build meaningful community access and social connection. Step 5: Ensure person has real opportunities to contribute and be seen as competent. Step 6: Review schedule monthly. Step 7: Document quality of life goals explicitly in BSP.

S29 Active Support: Tier 1 Established Evidence. Intellectual Disability. Support workers engage people in meaningful participation in all activities throughout day. Behaviours: Disruptive Behaviour, Emotional Dysregulation, Self-Injurious Behaviour, Social Withdrawal. Functions: Attention, Access to Tangibles, Automatic/Sensory. Step 1: Train all support workers in Active Support principles. Step 2: Conduct activity analysis for each daily routine. Step 3: Grade support using least to most principle. Step 4: Use little and often principle. Step 5: Match activities to preferences and strengths. Step 6: Implement organisational monitoring. Step 7: Review participation and behaviour data together.

S30 Errorless Learning: Tier 1 Established Evidence. Intellectual Disability, Acquired Brain Injury. Teaching technique minimising errors through high levels of support. Behaviours: Task Refusal, Physical Aggression, Verbal Aggression, Emotional Dysregulation. Functions: Escape/Avoidance. Step 1: Conduct task analysis. Step 2: Begin with prompt level guaranteeing correct performance. Step 3: Provide warm acknowledgement for correct steps. Step 4: Use prompt fading schedule. Step 5: If error occurs immediately provide correct prompt without attention to error. Step 6: Keep sessions brief and positive ending on success. Step 7: Document prompt levels and data.

S31 Supported Decision-Making and Dignity of Risk: Tier 2 Moderate Evidence. Intellectual Disability. UN CRPD Article 12. Behaviours: Task Refusal, Physical Aggression, Verbal Aggression, Elopement. Functions: Escape/Avoidance, Access to Tangibles. Step 1: Identify decisions being made for person that they could make themselves. Step 2: Develop decision support plan. Step 3: Provide information in accessible formats. Step 4: Allow adequate time. Step 5: Document and respect person's expressed preferences. Step 6: Address risk proportionately. Step 7: Involve formal supported decision-making model for significant decisions.

S32 TEACCH-Informed Structured Teaching: Tier 2 Moderate Evidence. Intellectual Disability, Autism/Neurodivergence. Organising environment, time, tasks, and communication to support independence. Behaviours: Task Refusal, Emotional Dysregulation, Disruptive Behaviour, Elopement. Functions: Escape/Avoidance, Automatic/Sensory. Step 1: Physical structure: organise environment so different areas signal different activities. Step 2: Scheduling: ensure individualised visual schedule. Step 3: Work systems: visual system for each activity. Step 4: Task organisation: left-to-right top-to-bottom. Step 5: Visual instructions embedded in task itself. Step 6: Train all support people. Step 7: Gradually increase complexity as independence improves.

S33 ABI-Informed Behaviour Support Framework: Tier 2 Moderate Evidence. Acquired Brain Injury. Ensuring BSP accounts for specific neurological profile. Behaviours: Physical Aggression, Verbal Aggression, Emotional Dysregulation, Task Refusal, Disruptive Behaviour. Functions: Escape/Avoidance, Attention, Automatic/Sensory. Step 1: Review all neuropsychological assessment reports. Step 2: Collaborate with neuropsychologist or rehabilitation physician. Step 3: Map relationship between cognitive deficits and behaviour: frontal lobe impulsivity, limbic emotional dysregulation, reduced insight, fatigue-driven irritability. Step 4: Design BSP to work with neurological profile. Step 5: Educate all support people about neurological basis. Step 6: Set realistic neurologically appropriate goals. Step 7: Review BSP as neurological recovery progresses.

S34 Fatigue Management Post-ABI: Tier 2 Moderate Evidence. Acquired Brain Injury. Addressing significant impact of post-ABI fatigue — affects up to 80 percent of people post-ABI. Behaviours: Physical Aggression, Verbal Aggression, Emotional Dysregulation, Task Refusal, Disruptive Behaviour. Functions: Escape/Avoidance, Automatic/Sensory. Step 1: Conduct fatigue assessment using Mental Fatigue Scale or Fatigue Severity Scale. Step 2: Map fatigue pattern across day and week. Step 3: Restructure schedule: demanding tasks during peak capacity, low-demand activities during high-fatigue periods. Step 4: Build proactive planned rest breaks before fatigue is reached. Step 5: Educate person and supports about fatigue as neurological phenomenon. Step 6: Develop fatigue monitoring system. Step 7: Adjust demands immediately when fatigue identified.

S35 Compensatory Strategy Training and Cognitive Scaffolding ABI: Tier 2 Moderate Evidence. Acquired Brain Injury. Teaching use of external aids to compensate for acquired cognitive deficits. Behaviours: Task Refusal, Verbal Aggression, Emotional Dysregulation, Disruptive Behaviour. Functions: Escape/Avoidance, Attention. Step 1: Obtain neuropsychological assessment. Step 2: Introduce compensatory strategies in collaboration with OT and neuropsychology. Step 3: Memory aids: memory books, diaries, wall calendars, smartphone alarms. Step 4: Attention supports: reduce distractions, use timers. Step 5: Executive function supports: written daily plans, visual checklists, decision-making templates. Step 6: Train support people to prompt strategy use consistently. Step 7: Fade support progressively.

S36 Recovery-Oriented Practice: Tier 2 Moderate Evidence. Mental Health. Framework prioritising person's self-defined recovery journey. Behaviours: Social Withdrawal, Emotional Dysregulation, Self-Harm, Task Refusal. Functions: Escape/Avoidance, Attention. Step 1: Begin with genuine exploration of person's values, hopes, strengths, and vision. Step 2: Co-produce goals meaningful to person. Step 3: Hold hope actively. Step 4: Prioritise self-determination. Step 5: Identify and build on existing strengths and interests. Step 6: Facilitate connection to peer support. Step 7: Review progress in terms of quality of life and meaning.

S37 Acceptance and Commitment Therapy ACT Informed Approaches: Tier 2 Moderate Evidence. Mental Health, Autism/Neurodivergence. Steven Hayes. Psychological flexibility. Core processes: acceptance, defusion, present-moment awareness, self-as-context, values clarification, committed action. Behaviours: Emotional Dysregulation, Self-Harm, Task Refusal, Social Withdrawal. Functions: Escape/Avoidance, Automatic/Sensory. Step 1: Acceptance: making room for difficult thoughts and feelings without struggling. Step 2: Defusion: observe thoughts rather than be controlled by them. Step 3: Present-moment awareness through mindfulness. Step 4: Self-as-context: stable observing sense of self. Step 5: Values clarification: what genuinely matters to person. Step 6: Committed action: small concrete steps toward values. Step 7: Use visual metaphors and adapted exercises.

S38 Motivational Interviewing MI: Tier 1 Established Evidence. Mental Health, Problematic Gaming, Intellectual Disability. William Miller and Stephen Rollnick. Collaborative person-centred counselling. Behaviours: Task Refusal, Problematic Gaming, Substance Use, Social Withdrawal, Self-Harm. Functions: Escape/Avoidance, Access to Tangibles. Step 1: Establish rapport with genuine empathy. Step 2: Roll with resistance: do not argue, reflect and explore. Step 3: Explore ambivalence using decisional balance. Step 4: Elicit change talk from person. Step 5: Support self-efficacy. Step 6: Avoid righting reflex: do not lecture or moralize. Step 7: When person is ready collaboratively plan specific achievable next steps.

S39 Mental Health Safety Planning: Tier 1 Established Evidence. Mental Health. Collaboratively developed written plan identifying warning signs, coping strategies, social supports, and crisis contacts. Behaviours: Self-Harm/Suicidal Ideation, Emotional Dysregulation. Functions: Escape/Avoidance, Automatic/Sensory. Step 1: Develop during calm or stable period. Step 2: Warning signs: personal early warning signs. Step 3: Internal coping strategies. Step 4: Social distractions: named people and situations. Step 5: People to ask for help: specific named individuals. Step 6: Crisis services: Lifeline 13 11 14, Suicide Call Back Service 1300 659 467, Beyond Blue 1300 22 4636. Step 7: Means restriction where safe and acceptable. Step 8: Follow up and review regularly.

S40 Schema-Informed Understanding of Entrenched Patterns: Tier 2 Moderate Evidence. Mental Health. Jeffrey Young. Deeply entrenched patterns from unmet childhood needs. Schemas: abandonment, defectiveness, emotional deprivation, mistrust. Behaviours: Physical Aggression, Verbal Aggression, Self-Harm, Emotional Dysregulation, Task Refusal. Functions: Escape/Avoidance, Attention. Step 1: Schema assessment via qualified schema therapist or Young Schema Questionnaire. Step 2: Identify core schemas and schema modes. Step 3: Share formulation in accessible validating language. Step 4: Train support people in schema-informed responses. Step 5: Use limited reparenting principles: consistent warmth, positive regard, reliability. Step 6: Identify and modify schema triggers in support environment. Step 7: Support formal schema therapy with qualified therapist.

S41 Function-Based Assessment of Gaming and Technology Use: Tier 2 Moderate Evidence. Autism/Neurodivergence, Mental Health, Intellectual Disability, Problematic Gaming. Applying PBS FBA framework to gaming. Distinguishes functional gaming from genuinely problematic use. Behaviours: Problematic Gaming, Emotional Dysregulation, Social Withdrawal, Task Refusal. Functions: Escape/Avoidance, Automatic/Sensory, Attention, Access to Tangibles. Step 1: Comprehensive assessment before forming conclusions. Step 2: Identify function: regulation, connection, mastery, identity, escape from overwhelming environments. Step 3: Assess harm: significant interference with sleep, physical health, school, or relationships. Step 4: If regulatory function: address underlying regulation needs first. Step 5: If social function: explore whether online connections supplement or replace face-to-face. Step 6: If escape function: assess and address what person is escaping from. Step 7: If harm established use function-based understanding to guide intervention.

S42 Family Systems Approach and Technology Boundaries: Tier 2 Moderate Evidence. Problematic Gaming, Mental Health. Addressing whole family ecology around technology. Behaviours: Problematic Gaming, Emotional Dysregulation, Physical Aggression, Verbal Aggression, Task Refusal. Functions: Escape/Avoidance, Access to Tangibles. Step 1: Work with whole family. Step 2: Psychoeducation for parents about gaming in neurodivergent young people. Step 3: Support parents to move from unilateral device removal to collaborative agreements. Step 4: Collaboratively develop Family Technology Agreement including timing, transitions, natural save points, alternative activities. Step 5: Address transition planning as most common flashpoint. Step 6: Build alternative connection based on young person's interests. Step 7: Address parental wellbeing.

S43 Behavioural Activation and Alternative Engagement Building: Tier 1 Established Evidence. Mental Health, Problematic Gaming, Autism/Neurodivergence. Scheduling and increasing engagement in rewarding values-consistent activities. Behaviours: Social Withdrawal, Problematic Gaming, Emotional Dysregulation, Task Refusal. Functions: Escape/Avoidance, Access to Tangibles, Attention. Step 1: Map person's current activity profile. Step 2: Identify functions that gaming or withdrawal is meeting. Step 3: Brainstorm menu of potential alternative activities with person based on their interests. Step 4: Start with low-barrier immediately rewarding activities. Step 5: Schedule activities alongside gaming initially, not replacing. Step 6: Gradually increase range and duration as alternatives become established. Step 7: When alternatives are established gaming naturally becomes less dominant without external restriction.

Population Index: Autism/Neurodivergence: S01-S28, S32, S37, S41, S43. Intellectual Disability: S01-S32, S38, S41. Acquired Brain Injury: S01-S28, S30, S33, S34, S35. Mental Health: S01-S28, S36-S43. Problematic Gaming: S01-S28, S38, S41, S42, S43.

Evidence: Tier 1 Established: S01, S02, S03, S04, S05, S06, S20, S25, S26, S27, S29, S30, S38, S39, S43. Tier 2 Moderate: S07-S19, S21-S24, S28, S31-S37, S40-S42.

Key References: Carr et al 2002. Gray 2010 Social Story Book. Greene 2014 Lost at School. Greenspan and Wieder 2006 Engaging Autism. Hughes 2009 Attachment-Focused Parenting. Kuypers 2011 Zones of Regulation. Linehan 1993 DBT. Mahler 2019 Interoception Curriculum. McDonnell 2010 Low Arousal Approaches. NDIS Commission 2025 Rules for Specialist Behaviour Support Providers. Porges 2011 Polyvagal Theory. Tiger Hanley Bruzek 2008 FCT Review. Webber Richardson Lambrick 2019 PBS Practice Framework Australia. Hayes Wilson Strosahl 2012 ACT. Miller Rollnick 2013 Motivational Interviewing. Young Klosko Weishaar 2003 Schema Therapy. Mansell Beadle-Brown 2012 Active Support. Uphoff et al 2020 Behavioural Activation Cochrane Review. Stanley Brown 2012 Safety Planning Intervention."""


# ══════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("  |  ".join(dict.fromkeys(cells)))
    return "\n".join(parts)

def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


# ══════════════════════════════════════════════════════════════════════════════
# CLAUDE — PBSP EXTRACTION (Tab 1)
# ══════════════════════════════════════════════════════════════════════════════

EXTRACTION_PROMPT = """
You are a Positive Behaviour Support (PBS) assistant.
Read the PBSP text below and extract key information into a JSON object.

Rules:
- Keep each bullet point under 90 characters — written for support workers to scan quickly
- Use plain language (no clinical jargon where possible)
- Focus on what is most operationally useful during a shift
- If a field cannot be found, use an empty string or empty list

Return ONLY valid JSON — no commentary, no markdown fences.

JSON structure:
{
  "name": "full name",
  "preferred": "preferred/short name",
  "pronouns": "pronouns",
  "age_info": "Age XX  |  [Diagnosis]  |  [Organisation / Location]",
  "about": ["up to 5 key points a support worker must know about this person"],
  "warning_signs": ["up to 5 observable early warning signs that behaviour is building"],
  "triggers": ["up to 6 known setting events and immediate triggers"],
  "proactive": ["up to 5 proactive strategies — things to DO to prevent behaviour"],
  "reactive": ["up to 5 reactive strategies — what to DO when behaviour occurs"],
  "do_not": ["up to 4 things NOT to do — known escalators"],
  "behaviours": [
    {
      "label": "short behaviour name (e.g. Verbal Outbursts)",
      "descriptors": ["up to 5 observable descriptors of this behaviour"]
    }
  ],
  "setting_events_checklist": ["up to 6 setting events as short checkbox labels"],
  "antecedents_checklist": ["up to 8 common antecedents as short checkbox labels"],
  "staff_responses_checklist": ["up to 7 common staff responses as short checkbox labels"],
  "review_date": "review year or date",
  "practitioner": "name and title",
  "contact": "email or contact"
}

PBSP TEXT:
"""

def extract_client_data(pbsp_text: str, api_key: str) -> dict:
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model="claude-opus-4-5", max_tokens=2000,
        messages=[{"role": "user", "content": EXTRACTION_PROMPT + pbsp_text[:40000]}]
    )
    raw = msg.content[0].text.strip()
    if raw.startswith("```"): raw = "\n".join(raw.split("\n")[1:])
    if raw.endswith("```"):   raw = "\n".join(raw.split("\n")[:-1])
    return json.loads(raw)


# ══════════════════════════════════════════════════════════════════════════════
# CLAUDE — STRATEGY RECOMMENDER (Tab 2)
# ══════════════════════════════════════════════════════════════════════════════

STRATEGY_PROMPT = """\
You are an experienced Positive Behaviour Support (PBS) Practitioner.
Based on the client profile and behaviours described, recommend practical PBS strategies for support workers.

{library_instruction}

Return ONLY valid JSON — no commentary, no markdown fences.

JSON structure:
{{
  "client_summary": "2-3 sentence summary of this person's likely support needs and overall PBS approach",
  "general_strategies": ["Up to 5 general/environmental strategies that apply across all behaviours"],
  "behaviours": [
    {{
      "behaviour": "behaviour name from input",
      "likely_function": "Primary function (Escape/Avoidance | Access | Sensory | Attention) — one sentence rationale",
      "proactive": ["Up to 5 proactive/antecedent strategies to prevent this behaviour"],
      "reactive": ["Up to 5 reactive de-escalation strategies for when this behaviour occurs"],
      "teach_instead": ["Up to 3 replacement skills or alternative communication strategies to build"],
      "avoid": ["Up to 3 specific things NOT to do — common mistakes that escalate this behaviour"]
    }}
  ]
}}

Rules:
- Keep each item under 90 characters — written for support workers to act on quickly
- Be specific to the triggers and context described — avoid generic advice
- Base strategy recommendations on the likely function of each behaviour
- Use plain language, not clinical jargon
- Replacement skills must serve the same function as the behaviour (functionally equivalent)

CLIENT PROFILE:
{profile}

BEHAVIOURS OF CONCERN:
{behaviours}
{library_section}"""

def recommend_strategies(client_info: dict, behaviours: list, api_key: str,
                          library_text: str = None) -> dict:
    profile_text = (
        f"Name: {client_info['name']}\n"
        f"Age: {client_info['age']}\n"
        f"Diagnosis/Condition: {client_info['diagnosis']}\n"
        f"Communication level: {client_info['comms']}\n"
        f"Additional context: {client_info['other']}"
    )
    behaviours_text = ""
    for i, b in enumerate(behaviours, 1):
        behaviours_text += (
            f"\nBehaviour {i}: {b['name']}\n"
            f"  What it looks like: {b['description']}\n"
            f"  Known triggers / when it occurs: {b['triggers']}\n"
        )
    if library_text:
        lib_instruction = (
            "IMPORTANT: Your primary task is to SELECT strategies FROM THE STRATEGY LIBRARY "
            "provided at the end of this prompt. Quote or closely paraphrase library strategies. "
            "Where the library has no relevant strategy for a specific need, you may suggest an "
            "evidence-based alternative and append '(not in library)' to that item."
        )
        lib_section = f"\nSTRATEGY LIBRARY:\n{library_text[:30000]}"
    else:
        lib_instruction = (
            "Generate evidence-based PBS strategies based on the likely function of each behaviour."
        )
        lib_section = ""
    prompt = STRATEGY_PROMPT.format(
        profile=profile_text, behaviours=behaviours_text,
        library_instruction=lib_instruction, library_section=lib_section,
    )
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model="claude-opus-4-5", max_tokens=3000,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = msg.content[0].text.strip()
    if raw.startswith("```"): raw = "\n".join(raw.split("\n")[1:])
    if raw.endswith("```"):   raw = "\n".join(raw.split("\n")[:-1])
    return json.loads(raw)


def pbsp_to_sr_format(data: dict) -> tuple:
    age_info = data.get("age_info", "")
    parts    = [p.strip() for p in age_info.split("|")]
    age      = parts[0].replace("Age", "").strip() if parts else "not specified"
    diagnosis = parts[1] if len(parts) > 1 else "not specified"
    client_info = {
        "name":      data.get("name") or data.get("preferred") or "Unknown",
        "age":       age or "not specified",
        "diagnosis": diagnosis or "not specified",
        "comms":     "refer to PBSP",
        "other":     "; ".join(data.get("about", [])) or "none provided",
    }
    triggers_str = "; ".join(data.get("triggers", []))
    behaviours = [
        {
            "name":        b.get("label", "Behaviour"),
            "description": "; ".join(b.get("descriptors", [])),
            "triggers":    triggers_str,
        }
        for b in data.get("behaviours", [])
    ]
    return client_info, behaviours


FREETEXT_STRATEGY_PROMPT = """\
You are an experienced Positive Behaviour Support (PBS) Practitioner.
A practitioner has described a client and their behaviours in plain language below.

Your tasks:
1. Read the description and identify each distinct behaviour of concern
2. Assign each a clear clinical label
3. Infer the likely triggers and context from what is described
4. Determine the likely function of each behaviour
5. Recommend practical PBS strategies

{library_instruction}

Return ONLY valid JSON — no commentary, no markdown fences.

JSON structure:
{{
  "client_summary": "2-3 sentence clinical summary based on the description",
  "general_strategies": ["Up to 5 general/environmental strategies across all behaviours"],
  "behaviours": [
    {{
      "behaviour": "Clinical label you have assigned",
      "likely_function": "Primary function (Escape/Avoidance | Access | Sensory | Attention) — one sentence rationale drawn from the description",
      "proactive": ["Up to 5 proactive strategies to prevent this behaviour"],
      "reactive": ["Up to 5 reactive de-escalation strategies when this behaviour occurs"],
      "teach_instead": ["Up to 3 replacement skills or communication strategies to build"],
      "avoid": ["Up to 3 things NOT to do — common mistakes that escalate this behaviour"]
    }}
  ]
}}

Rules:
- Keep each item under 90 characters — written for support workers to act on quickly
- Use the description to infer triggers and context — be specific, not generic
- Assign clinically appropriate behaviour labels
- Use plain language in strategy recommendations
- Replacement skills must serve the same function as the behaviour

CLIENT PROFILE:
{profile}

PRACTITIONER'S DESCRIPTION:
{freetext}
{library_section}"""


def recommend_from_freetext(client_info: dict, freetext: str, api_key: str,
                              library_text: str = None) -> dict:
    profile_text = (
        f"Name: {client_info['name']}\n"
        f"Age: {client_info['age']}\n"
        f"Diagnosis/Condition: {client_info['diagnosis']}\n"
        f"Communication level: {client_info['comms']}\n"
        f"Additional context: {client_info['other']}"
    )
    if library_text:
        lib_instruction = (
            "IMPORTANT: Select strategies FROM THE STRATEGY LIBRARY provided at the end of this "
            "prompt. Quote or closely paraphrase library strategies. Where the library has no "
            "relevant strategy, suggest an evidence-based alternative and note '(not in library)'."
        )
        lib_section = f"\nSTRATEGY LIBRARY:\n{library_text[:30000]}"
    else:
        lib_instruction = "Generate evidence-based PBS strategies based on function and context."
        lib_section = ""
    prompt = FREETEXT_STRATEGY_PROMPT.format(
        profile=profile_text, freetext=freetext,
        library_instruction=lib_instruction, library_section=lib_section,
    )
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model="claude-opus-4-5", max_tokens=3000,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = msg.content[0].text.strip()
    if raw.startswith("```"): raw = "\n".join(raw.split("\n")[1:])
    if raw.endswith("```"):   raw = "\n".join(raw.split("\n")[:-1])
    return json.loads(raw)


# ══════════════════════════════════════════════════════════════════════════════
# PDF HELPERS (shared)
# ══════════════════════════════════════════════════════════════════════════════

def drect(c, x, y, w, h, fill=None, stroke=None, lw=0.75):
    c.saveState()
    if fill:   c.setFillColor(fill)
    if stroke: c.setStrokeColor(stroke); c.setLineWidth(lw)
    c.rect(x, y, w, h, fill=1 if fill else 0, stroke=1 if stroke else 0)
    c.restoreState()

def dlbl(c, x, y, s, size, bold=False, italic=False, color=DARK_TEXT, align='left'):
    c.saveState()
    c.setFillColor(color)
    face = ('Helvetica-BoldOblique' if bold and italic else
            'Helvetica-Bold'        if bold           else
            'Helvetica-Oblique'     if italic         else 'Helvetica')
    c.setFont(face, size)
    if align == 'center': c.drawCentredString(x, y, s)
    elif align == 'right': c.drawRightString(x, y, s)
    else: c.drawString(x, y, s)
    c.restoreState()

def sec(c, y, title, bg, ht=20):
    drect(c, LM, y - ht, CW, ht, fill=bg)
    dlbl(c, LM + 8, y - ht + 6, title, 11, bold=True, color=white)
    return y - ht

def checkbox(c, x, y, size=8):
    c.saveState(); c.setStrokeColor(MID_GREY); c.setLineWidth(0.75)
    c.rect(x, y, size, size, fill=0, stroke=1); c.restoreState()

def write_line(c, x, y, w):
    c.saveState(); c.setStrokeColor(MID_GREY); c.setLineWidth(0.5)
    c.line(x, y, x + w, y); c.restoreState()

def cb_item(c, x, y, text, size=9, color=DARK_TEXT, bold=False):
    checkbox(c, x, y, 8)
    dlbl(c, x + 13, y + 1, text, size, bold=bold, color=color)

def wrap_text(text: str, font_name: str, font_size: float, max_w: float) -> list:
    words = text.split()
    lines, line = [], ""
    for word in words:
        test = (line + " " + word).strip()
        if stringWidth(test, font_name, font_size) <= max_w:
            line = test
        else:
            if line: lines.append(line)
            line = word
    if line: lines.append(line)
    return lines or [""]


# ══════════════════════════════════════════════════════════════════════════════
# SUPPORT REFERENCE CARD
# ══════════════════════════════════════════════════════════════════════════════

def generate_support_card(d: dict) -> BytesIO:
    buf = BytesIO()
    c   = rl_canvas.Canvas(buf, pagesize=A4)
    RH, SH, GAP = 22, 20, 7

    ABOUT_H = SH + len(d["about"])         * RH
    WARN_H  = SH + len(d["warning_signs"]) * RH
    SIDE_H  = SH + max(len(d["triggers"]), len(d["proactive"])) * RH
    REACT_H = SH + len(d["reactive"])      * RH
    DONT_H  = SH + len(d["do_not"])        * RH
    HDR_H, FTR_H, BOT = 72, 38, 12
    y = H

    drect(c, 0, y-HDR_H, W, HDR_H, fill=DEEP_BLUE)
    drect(c, 0, y-HDR_H, 8, HDR_H, fill=TEAL)
    px, py = W-LM-58, y-HDR_H+6
    drect(c, px, py, 52, 58, fill=BLUE2, stroke=HexColor('#C8E6F5'), lw=1)
    dlbl(c, px+26, py+28, "PHOTO",      7, color=HexColor('#C8E6F5'), align='center')
    dlbl(c, px+26, py+18, "(optional)", 7, italic=True, color=HexColor('#8FC8DC'), align='center')
    dlbl(c, LM+6, y-22, d["name"],           22, bold=True, color=white)
    dlbl(c, LM+6, y-40, d.get("pronouns",""), 11, color=HexColor('#C8E6F5'))
    dlbl(c, LM+6, y-56, d.get("age_info",""), 10, color=HexColor('#8FC8DC'))
    dlbl(c, LM+6, y-70, "Positive Behaviour Support  —  Behaviour Support Reference Card",
         8, italic=True, color=HexColor('#5B9FC0'))
    y -= HDR_H + GAP

    y = sec(c, y, f"About {d['preferred']}", TEAL)
    drect(c, LM, y-(ABOUT_H-SH), CW, ABOUT_H-SH, fill=LIGHT_TEAL)
    for i, item in enumerate(d["about"]):
        drect(c, LM+10, y-12-i*RH, 7, 7, fill=TEAL)
        dlbl(c,  LM+21, y- 7-i*RH, item, 9.5, color=DARK_TEXT)
    y -= (ABOUT_H-SH) + GAP

    y = sec(c, y, "Early warning signs — watch for these", AMBER)
    drect(c, LM, y-(WARN_H-SH), CW, WARN_H-SH, fill=LIGHT_AMBR)
    for i, item in enumerate(d["warning_signs"]):
        drect(c, LM+10, y-12-i*RH, 7, 7, fill=AMBER)
        dlbl(c,  LM+21, y- 7-i*RH, item, 9.5, color=DARK_TEXT)
    y -= (WARN_H-SH) + GAP

    half, cont = (CW-5)/2, SIDE_H-SH
    drect(c, LM,       y-SH, half, SH, fill=BLUE2)
    dlbl(c, LM+8,      y-13, "Known triggers",       11, bold=True, color=white)
    drect(c, LM,       y-SIDE_H, half, cont, fill=LIGHT_BLUE)
    for i, item in enumerate(d["triggers"]):
        drect(c, LM+10, y-SH-10-i*RH, 7, 7, fill=BLUE2)
        dlbl(c,  LM+21, y-SH- 5-i*RH, item, 9.5, color=DARK_TEXT)
    rx = LM+half+5
    drect(c, rx,       y-SH, half, SH, fill=GREEN)
    dlbl(c, rx+8,      y-13, "Proactive strategies",  11, bold=True, color=white)
    drect(c, rx,       y-SIDE_H, half, cont, fill=LIGHT_GRN)
    for i, item in enumerate(d["proactive"]):
        drect(c, rx+10, y-SH-10-i*RH, 7, 7, fill=GREEN)
        dlbl(c,  rx+21, y-SH- 5-i*RH, item, 9.5, color=DARK_TEXT)
    y -= SIDE_H + GAP

    y = sec(c, y, "When behaviour occurs — do this", TEAL)
    drect(c, LM, y-(REACT_H-SH), CW, REACT_H-SH, fill=LIGHT_TEAL)
    for i, item in enumerate(d["reactive"]):
        drect(c, LM+10, y-12-i*RH, 7, 7, fill=TEAL)
        dlbl(c,  LM+21, y- 7-i*RH, item, 9.5, color=DARK_TEXT)
    y -= (REACT_H-SH) + GAP

    y = sec(c, y, f"DO NOT — things that escalate behaviour for {d['preferred']}", RED)
    drect(c, LM, y-(DONT_H-SH), CW, DONT_H-SH, fill=LIGHT_RED)
    for i, item in enumerate(d["do_not"]):
        drect(c, LM+10, y-12-i*RH, 7, 7, fill=RED)
        dlbl(c,  LM+21, y- 7-i*RH, item, 9.5, bold=True, color=RED)
    y -= (DONT_H-SH) + GAP

    drect(c, 0, BOT, W, FTR_H, fill=DEEP_BLUE)
    drect(c, 0, BOT, 8, FTR_H, fill=TEAL)
    ft = BOT + FTR_H
    dlbl(c, LM+6, ft-16,
         f"Plan review: {d.get('review_date','')}  |  {d.get('practitioner','')}", 9, bold=True, color=white)
    dlbl(c, LM+6, ft-29, d.get("contact",""), 9, italic=True, color=HexColor('#C8E6F5'))
    dlbl(c, W-LM, ft-29, "CONFIDENTIAL — handle in line with your privacy policy",
         8, italic=True, color=HexColor('#5B9FC0'), align='right')
    c.save(); buf.seek(0); return buf


# ══════════════════════════════════════════════════════════════════════════════
# ABC RECORDING FORM
# ══════════════════════════════════════════════════════════════════════════════

def generate_abc_form(d: dict) -> BytesIO:
    buf = BytesIO()
    c   = rl_canvas.Canvas(buf, pagesize=A4)
    y   = H
    half_cw = (CW - 10) / 2

    drect(c, 0, y-58, W, 58, fill=DEEP_BLUE)
    drect(c, 0, y-58, 7, 58, fill=TEAL)
    dlbl(c, LM+4, y-20, "ABC Behaviour Recording Form", 16, bold=True, color=white)
    dlbl(c, LM+4, y-36,
         f"{d['name']}  |  {d.get('pronouns','')}  |  {d.get('age_info','')}", 9, color=HexColor('#C8E6F5'))
    dlbl(c, LM+4, y-52, "Complete as soon as safely possible after any behaviour of concern",
         8, italic=True, color=HexColor('#8FC8DC'))
    dlbl(c, W-LM, y-52, "CONFIDENTIAL", 8, bold=True, color=HexColor('#5B9FC0'), align='right')
    y -= 63

    drect(c, LM, y-22, CW, 22, fill=LIGHT_BLUE)
    dlbl(c, LM+6,   y-14, "Date:",           9, bold=True, color=BLUE2)
    write_line(c, LM+32,  y-15, 68)
    dlbl(c, LM+110, y-14, "Time:",           9, bold=True, color=BLUE2)
    write_line(c, LM+134, y-15, 30)
    dlbl(c, LM+170, y-14, "to",              9, color=MED_TEXT)
    write_line(c, LM+181, y-15, 30)
    dlbl(c, LM+220, y-14, "Support worker:", 9, bold=True, color=BLUE2)
    write_line(c, LM+297, y-15, 90)
    dlbl(c, LM+398, y-14, "Shift:",          9, bold=True, color=BLUE2)
    for si, sh in enumerate(["Day","Aft","Night"]):
        sx = LM+425+si*38; checkbox(c, sx, y-16); dlbl(c, sx+11, y-14, sh, 8.5, color=DARK_TEXT)
    y -= 27

    drect(c, LM, y-20, 20, 20, fill=AMBER)
    dlbl(c, LM+10, y-13, "S", 11, bold=True, color=white, align='center')
    drect(c, LM+20, y-20, CW-20, 20, fill=LIGHT_AMBR)
    dlbl(c, LM+28, y-13, "Setting events today", 11, bold=True, color=AMBER)
    y -= 20
    se = d.get("setting_events_checklist", [])
    se_h = 10 + ((len(se)+1)//2)*18 + 18
    drect(c, LM, y-se_h, CW, se_h, fill=LIGHT_AMBR)
    for i, item in enumerate(se):
        col, row = i%2, i//2
        cb_item(c, LM+6+col*(half_cw+10), y-14-row*18, item)
    dlbl(c, LM+6, y-se_h+6, "Other:", 9, color=MED_TEXT)
    write_line(c, LM+38, y-se_h+5, CW-44)
    y -= se_h + 5

    drect(c, LM, y-20, 20, 20, fill=BLUE2)
    dlbl(c, LM+10, y-13, "A", 11, bold=True, color=white, align='center')
    drect(c, LM+20, y-20, CW-20, 20, fill=LIGHT_BLUE)
    dlbl(c, LM+28, y-13, "Antecedent — what happened immediately BEFORE?", 11, bold=True, color=BLUE2)
    y -= 20
    ant = d.get("antecedents_checklist", [])
    ant_h = 10 + ((len(ant)+1)//2)*18 + 18
    drect(c, LM, y-ant_h, CW, ant_h, fill=LIGHT_BLUE)
    for i, item in enumerate(ant):
        col, row = i%2, i//2
        cb_item(c, LM+6+col*(half_cw+10), y-14-row*18, item)
    dlbl(c, LM+6, y-ant_h+6, "Other / describe:", 9, color=MED_TEXT)
    write_line(c, LM+96, y-ant_h+5, CW-102)
    y -= ant_h + 5

    drect(c, LM, y-20, 20, 20, fill=TEAL)
    dlbl(c, LM+10, y-13, "B", 11, bold=True, color=white, align='center')
    drect(c, LM+20, y-20, CW-20, 20, fill=LIGHT_TEAL)
    dlbl(c, LM+28, y-13, "Behaviour — what did you observe? (facts only, no interpretations)",
         11, bold=True, color=TEAL)
    y -= 20
    behs = d.get("behaviours", [])
    n_beh = len(behs)
    beh_col_w = (CW-8) / max(n_beh, 1)
    max_desc = max((len(b.get("descriptors",[])) for b in behs), default=3)
    beh_h = 18 + max_desc*16 + 24
    drect(c, LM, y-beh_h, CW, beh_h, fill=LIGHT_TEAL)
    beh_colors = [TEAL, BLUE2, GREEN, AMBER]
    for bi, beh in enumerate(behs):
        bx = LM+4+bi*(beh_col_w+4); bw = beh_col_w-4; bc = beh_colors[bi%len(beh_colors)]
        drect(c, bx, y-18, bw, 18, fill=bc)
        dlbl(c, bx+5, y-12, beh.get("label",""), 9, bold=True, color=white)
        for di, desc in enumerate(beh.get("descriptors",[])):
            cb_item(c, bx+2, y-26-di*16, desc, size=8.5)
    ir = y-beh_h+22
    drect(c, LM+4, ir, CW-8, 16, fill=HexColor('#DFF0EE'))
    dlbl(c, LM+8, ir+5, "Intensity:", 9, bold=True, color=TEAL)
    for ii, lvl in enumerate(["Mild","Moderate","Severe"]):
        sx = LM+54+ii*58; checkbox(c, sx, ir+4); dlbl(c, sx+11, ir+5, lvl, 9, color=DARK_TEXT)
    dlbl(c, LM+232, ir+5, "Duration:", 9, bold=True, color=TEAL)
    write_line(c, LM+272, ir+4, 28)
    dlbl(c, LM+305, ir+5, "mins", 9, color=MED_TEXT)
    dlbl(c, LM+345, ir+5, "Location:", 9, bold=True, color=TEAL)
    write_line(c, LM+385, ir+4, CW-203)
    dlbl(c, LM+6, y-beh_h+8, "Describe what you saw (objective language):", 9, italic=True, color=MED_TEXT)
    write_line(c, LM+6, y-beh_h+2, CW-12)
    y -= beh_h + 5

    drect(c, LM, y-20, 20, 20, fill=GREEN)
    dlbl(c, LM+10, y-13, "C", 11, bold=True, color=white, align='center')
    drect(c, LM+20, y-20, CW-20, 20, fill=LIGHT_GRN)
    dlbl(c, LM+28, y-13, "Consequence — what happened after? What did you do?",
         11, bold=True, color=GREEN)
    y -= 20
    resp = d.get("staff_responses_checklist", [])
    resp_h = 10 + ((len(resp)+1)//2)*17 + 52
    drect(c, LM, y-resp_h, CW, resp_h, fill=LIGHT_GRN)
    dlbl(c, LM+6, y-10, "What did you do?", 9, bold=True, color=GREEN)
    for i, item in enumerate(resp):
        col, row = i%2, i//2
        cb_item(c, LM+6+col*(half_cw+10), y-22-row*17, item)
    ry = y-resp_h+48
    dlbl(c, LM+6, ry, "How did they respond?", 9, bold=True, color=GREEN)
    for ri, r in enumerate(["De-escalated < 5 mins","De-escalated < 30 mins",
                             "Continued > 30 mins","Escalated further"]):
        rx2 = LM+6+ri*134; checkbox(c, rx2, ry-14); dlbl(c, rx2+11, ry-13, r, 8.5, color=DARK_TEXT)
    dlbl(c, LM+6, ry-28, "Other:", 9, color=MED_TEXT)
    write_line(c, LM+38, ry-29, CW-44)
    y -= resp_h + 5

    drect(c, LM, y-18, CW, 18, fill=LIGHT_BLUE)
    dlbl(c, LM+6, y-12, "Additional notes:", 10, bold=True, color=BLUE2)
    y -= 18
    drect(c, LM, y-44, CW, 44, fill=HexColor('#FAFAFA'), stroke=MID_GREY, lw=0.5)
    for li in range(3): write_line(c, LM+6, y-16-li*13, CW-12)
    y -= 48

    drect(c, LM, y-28, CW, 28, fill=LIGHT_TEAL)
    for i, ch in enumerate(["Incident report completed","Handed over to next shift",
                             "Staff debrief completed","Behaviour data entered"]):
        ix = LM+6+i*(CW/4); checkbox(c, ix, y-18); dlbl(c, ix+11, y-16, ch, 8.5, color=DARK_TEXT)
    y -= 32

    drect(c, LM, y-24, CW, 24, fill=DEEP_BLUE)
    dlbl(c, LM+6,   y-10, "Signature:",   9, bold=True, color=white)
    write_line(c, LM+55,  y-11, 110)
    dlbl(c, LM+175, y-10, "Print name:",  9, bold=True, color=white)
    write_line(c, LM+227, y-11, 110)
    dlbl(c, LM+347, y-10, "Date / Time:", 9, bold=True, color=white)
    write_line(c, LM+407, y-11, CW-230)
    dlbl(c, W/2, y-20, f"{d.get('practitioner','')}  |  {d.get('contact','')}",
         7, italic=True, color=HexColor('#8FC8DC'), align='center')

    c.save(); buf.seek(0); return buf


# ══════════════════════════════════════════════════════════════════════════════
# STRATEGY REPORT PDF (Tab 2)
# ══════════════════════════════════════════════════════════════════════════════

def generate_strategy_report(result: dict, client_name: str,
                              practitioner: str = "", contact: str = "") -> BytesIO:
    buf = BytesIO()
    c   = rl_canvas.Canvas(buf, pagesize=A4)

    BOT, FTR_H, HDR_H  = 12, 38, 68
    MINI_HDR_H          = 28
    MIN_Y               = BOT + FTR_H + 20
    RH, SH, GAP         = 17, 20, 8
    page_num            = [1]

    def draw_footer():
        drect(c, 0, BOT, W, FTR_H, fill=DEEP_BLUE)
        drect(c, 0, BOT, 8, FTR_H, fill=TEAL)
        ft = BOT + FTR_H
        dlbl(c, LM+6, ft-14, practitioner or "Positive Behaviour Support", 9, bold=True, color=white)
        if contact:
            dlbl(c, LM+6, ft-28, contact, 9, italic=True, color=HexColor('#C8E6F5'))
        dlbl(c, W-LM, ft-14, f"Page {page_num[0]}", 9, color=HexColor('#8FC8DC'), align='right')
        dlbl(c, W-LM, ft-28, "CONFIDENTIAL — handle in line with your privacy policy",
             8, italic=True, color=HexColor('#5B9FC0'), align='right')

    def new_page():
        draw_footer(); c.showPage(); page_num[0] += 1
        drect(c, 0, H-MINI_HDR_H, W, MINI_HDR_H, fill=DEEP_BLUE)
        drect(c, 0, H-MINI_HDR_H, 8, MINI_HDR_H, fill=TEAL)
        dlbl(c, LM+6, H-MINI_HDR_H+9,
             f"Behaviour Strategy Recommendations — {client_name}", 10, bold=True, color=white)
        return H - MINI_HDR_H - GAP

    def ensure(y, needed):
        return new_page() if y - needed < MIN_Y else y

    y = H
    drect(c, 0, y-HDR_H, W, HDR_H, fill=DEEP_BLUE)
    drect(c, 0, y-HDR_H, 8, HDR_H, fill=TEAL)
    dlbl(c, LM+6, y-22, "Behaviour Strategy Recommendations", 17, bold=True, color=white)
    dlbl(c, LM+6, y-42, client_name, 13, color=HexColor('#C8E6F5'))
    dlbl(c, LM+6, y-58, "Positive Behaviour Support — Strategy Guide for Support Workers",
         9, italic=True, color=HexColor('#5B9FC0'))
    dlbl(c, W-LM, y-58, f"Generated: {date.today().strftime('%d %B %Y')}",
         9, italic=True, color=HexColor('#5B9FC0'), align='right')
    y -= HDR_H + GAP

    summary = result.get("client_summary", "")
    if summary:
        lines = wrap_text(summary, "Helvetica-Oblique", 9.5, CW - 24)
        sum_h = len(lines) * 15 + 10
        y = ensure(y, sum_h + SH)
        y = sec(c, y, "Clinical Summary", TEAL)
        drect(c, LM, y-sum_h, CW, sum_h, fill=LIGHT_TEAL)
        for li, ln in enumerate(lines):
            dlbl(c, LM+10, y-13-li*15, ln, 9.5, italic=True, color=DARK_TEXT)
        y -= sum_h + GAP

    gen = result.get("general_strategies", [])
    if gen:
        gen_h = len(gen) * RH
        y = ensure(y, gen_h + SH)
        y = sec(c, y, "General Strategies — apply across all behaviours", BLUE2)
        drect(c, LM, y-gen_h, CW, gen_h, fill=LIGHT_BLUE)
        for i, item in enumerate(gen):
            drect(c, LM+10, y-9-i*RH, 7, 7, fill=BLUE2)
            dlbl(c,  LM+21, y-4-i*RH, item, 9.5, color=DARK_TEXT)
        y -= gen_h + GAP

    for beh in result.get("behaviours", []):
        b_name  = beh.get("behaviour", "Behaviour")
        b_func  = beh.get("likely_function", "")
        pros    = beh.get("proactive", [])
        reacts  = beh.get("reactive", [])
        teaches = beh.get("teach_instead", [])
        avoids  = beh.get("avoid", [])

        y = ensure(y, 26)
        drect(c, LM, y-26, CW, 26, fill=DEEP_BLUE)
        drect(c, LM, y-26,  6, 26, fill=TEAL)
        dlbl(c, LM+14, y-17, f"Behaviour: {b_name}", 12, bold=True, color=white)
        y -= 26

        y = ensure(y, 22)
        drect(c, LM, y-22, CW, 22, fill=LIGHT_AMBR)
        dlbl(c, LM+8,  y-14, "Likely function:", 9.5, bold=True, color=AMBER)
        func_lines = wrap_text(b_func, "Helvetica-Oblique", 9.5, CW - 112)
        dlbl(c, LM+106, y-14, func_lines[0] if func_lines else "", 9.5, italic=True, color=DARK_TEXT)
        y -= 26

        def draw_section(items, title, bg, light_bg, dot_col):
            nonlocal y
            if not items: return
            h = len(items) * RH
            y = ensure(y, h + SH)
            y = sec(c, y, title, bg)
            drect(c, LM, y-h, CW, h, fill=light_bg)
            for i, item in enumerate(items):
                drect(c, LM+10, y-9-i*RH, 7, 7, fill=dot_col)
                dlbl(c,  LM+21, y-4-i*RH, item, 9.5,
                     bold=(dot_col == RED), color=(dot_col if dot_col == RED else DARK_TEXT))
            y -= h + 4

        draw_section(pros,    "Proactive strategies — prevent this behaviour",         GREEN, LIGHT_GRN,  GREEN)
        draw_section(reacts,  "Reactive strategies — when this behaviour occurs",       TEAL,  LIGHT_TEAL, TEAL)
        draw_section(teaches, "Skills to build — teach as a replacement behaviour",     BLUE2, LIGHT_BLUE, BLUE2)
        draw_section(avoids,  "DO NOT — avoid these with this behaviour",               RED,   LIGHT_RED,  RED)
        y -= GAP

    draw_footer()
    c.save(); buf.seek(0); return buf


# ══════════════════════════════════════════════════════════════════════════════
# STREAMLIT UI
# ══════════════════════════════════════════════════════════════════════════════

st.title("📋 PBS Support Tool")

# API key — sidebar
try:
    api_key = st.secrets.get("ANTHROPIC_API_KEY", "")
except Exception:
    api_key = ""

if not api_key:
    with st.sidebar:
        st.markdown("### 🔑 API Key")
        api_key = st.text_input(
            "Anthropic API key",
            type="password",
            help="Get a key at console.anthropic.com",
            key="api_key_input",
        )
        st.caption("Your key is never stored.")

tab1, tab2 = st.tabs(["📄 Generate from PBSP", "🧠 Strategy Recommender"])


# ── TAB 1 ─────────────────────────────────────────────────────────────────────
with tab1:
    st.markdown(
        "Upload a client's **Positive Behaviour Support Plan** (PDF or Word) and this tool "
        "will generate a **support reference card** and an **ABC recording form** — "
        "both pre-populated with the client's specific information."
    )
    st.divider()

    uploaded = st.file_uploader("Upload PBSP (PDF or Word .docx)", type=["pdf","docx"])
    gen_btn  = st.button("Generate documents", type="primary",
                         disabled=not (uploaded and api_key), key="gen_tab1")

    if gen_btn and uploaded and api_key:
        with st.spinner("Reading plan…"):
            fb = uploaded.read()
            pbsp_text = extract_text_from_docx(fb) if uploaded.name.lower().endswith(".docx") \
                        else extract_text_from_pdf(fb)

        with st.spinner("Extracting client information…"):
            try:   data = extract_client_data(pbsp_text, api_key)
            except Exception as e: st.error(f"Extraction failed: {e}"); st.stop()

        with st.spinner("Generating PDFs…"):
            try:
                card_buf = generate_support_card(data)
                abc_buf  = generate_abc_form(data)
            except Exception as e: st.error(f"PDF error: {e}"); st.stop()

        st.session_state["t1_data"] = data
        st.success(f"✅  Documents generated for **{data.get('name','client')}**")
        st.divider()
        safe = data.get("name","client").replace(" ","_")
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("#### 📄 Support Reference Card")
            st.markdown("One-page quick reference — print and laminate.")
            st.download_button("Download support card (PDF)", card_buf,
                               f"{safe}_Support_Card.pdf", "application/pdf",
                               use_container_width=True)
        with c2:
            st.markdown("#### 📋 ABC Recording Form")
            st.markdown("Pre-populated with this client's behaviours.")
            st.download_button("Download ABC form (PDF)", abc_buf,
                               f"{safe}_ABC_Form.pdf", "application/pdf",
                               use_container_width=True)
        with st.expander("Review extracted information"):
            st.json(data)

    st.divider()
    st.caption("Always review generated documents before distributing to staff. "
               "Handle all client documents in line with your organisation's privacy policy.")


# ── TAB 2 ─────────────────────────────────────────────────────────────────────
with tab2:
    st.markdown(
        "Describe a client's behaviours and this tool will select the most appropriate strategies "
        "from the **Hoggs Health Hub PBS Strategy Library** and generate a **printable strategy report**."
    )
    st.divider()

    # ── STEP 1: Client source ─────────────────────────────────────────────────
    st.markdown("#### Step 1 — Client information")

    source_opts = ["📄 Upload a PBSP (auto-extract)", "✏️ Enter manually"]
    if "t1_data" in st.session_state:
        source_opts.insert(0,
            f"♻️ Use client from Tab 1 ({st.session_state['t1_data'].get('name','')})")
    src = st.radio("Where is the client information coming from?",
                   source_opts, key="sr_src", horizontal=True)

    client_info    = None
    valid_behs     = []
    freetext_value = ""

    src = src or source_opts[0]
    if src.startswith("♻️"):
        t1 = st.session_state["t1_data"]
        client_info, valid_behs = pbsp_to_sr_format(t1)
        st.success(f"✅  Using: **{t1.get('name','')}** — {t1.get('age_info','')}")
        with st.expander("View extracted profile and behaviours"):
            cols = st.columns(2)
            with cols[0]:
                st.markdown(f"**Name:** {t1.get('name','')}")
                st.markdown(f"**Profile:** {t1.get('age_info','')}")
                st.markdown("**About:**")
                for a in t1.get("about",[]): st.markdown(f"- {a}")
            with cols[1]:
                st.markdown("**Behaviours of concern:**")
                for b in t1.get("behaviours",[]): st.markdown(f"- **{b.get('label','')}:** " +
                    ", ".join(b.get("descriptors",[])))
                st.markdown("**Triggers:**")
                for t in t1.get("triggers",[]): st.markdown(f"- {t}")

    elif src.startswith("📄"):
        sr_pbsp_file = st.file_uploader("Upload client's PBSP (PDF or Word)",
                                         type=["pdf","docx"], key="sr_pbsp_upload")
        if sr_pbsp_file:
            if st.button("Extract client information", type="secondary", key="sr_pbsp_extract"):
                if not api_key:
                    st.error("API key required — enter it in the sidebar.")
                else:
                    with st.spinner("Extracting from PBSP…"):
                        fb = sr_pbsp_file.read()
                        txt = extract_text_from_docx(fb) if sr_pbsp_file.name.lower().endswith(".docx") \
                              else extract_text_from_pdf(fb)
                        try:
                            extracted = extract_client_data(txt, api_key)
                            st.session_state["sr_pbsp_data"] = extracted
                            st.success(f"✅  Extracted: {extracted.get('name','')}")
                        except Exception as e:
                            st.error(f"Extraction failed: {e}")

        if "sr_pbsp_data" in st.session_state:
            pd_ = st.session_state["sr_pbsp_data"]
            client_info, valid_behs = pbsp_to_sr_format(pd_)
            st.info(f"Using extracted data for **{pd_.get('name','')}** — "
                    f"{len(valid_behs)} behaviour(s) found.")
            with st.expander("View extracted behaviours"):
                for b in pd_.get("behaviours",[]): st.markdown(
                    f"- **{b.get('label','')}:** " + ", ".join(b.get("descriptors",[])))
            if st.button("🗑 Clear extracted data", key="sr_pbsp_clear"):
                del st.session_state["sr_pbsp_data"]; st.rerun()

    else:
        ca, cb = st.columns(2)
        with ca:
            sr_name = st.text_input("Client name *", key="sr_name",
                                     placeholder="e.g. Alex Thompson")
            sr_age  = st.text_input("Age", key="sr_age", placeholder="e.g. 24")
        with cb:
            sr_diag  = st.text_input("Diagnosis / condition", key="sr_diag",
                                      placeholder="e.g. Autism Spectrum Disorder, ABI")
            sr_comms = st.selectbox("Communication level", key="sr_comms",
                                     options=["Verbal","Limited verbal","Non-verbal","Uses AAC device"])
        sr_other = st.text_area("Other relevant context (optional)", key="sr_other", height=75,
                                 placeholder="e.g. Routines are very important, history of trauma…")

        st.markdown("**Behaviours of concern**")
        entry_style = st.radio(
            "How would you like to describe the behaviours?",
            ["📝 Describe in your own words", "📋 Structured entry"],
            key="sr_entry_style", horizontal=True,
        )

        if (entry_style or "").startswith("📝"):
            sr_freetext = st.text_area(
                "Describe what you've observed — write naturally",
                key="sr_freetext",
                height=160,
                placeholder=(
                    "Write however feels natural — describe what you see, when it happens, "
                    "and what seems to set it off.\n\n"
                    "e.g. 'Alex hits out at staff when demands are placed, especially during "
                    "transitions. He also bangs his head when frustrated or overwhelmed.'"
                ),
            )
            freetext_value = (sr_freetext or "").strip()
            valid_behs = [{"_freetext": True}] if freetext_value else []
            if freetext_value:
                client_info = {
                    "name":      (sr_name or "").strip(),
                    "age":       (sr_age  or "").strip() or "not specified",
                    "diagnosis": (sr_diag or "").strip() or "not specified",
                    "comms":     sr_comms or "Verbal",
                    "other":     (sr_other or "").strip() or "none provided",
                }

        else:
            freetext_value = ""
            if "sr_n" not in st.session_state:
                st.session_state.sr_n = 1

            beh_raw = []
            for i in range(st.session_state.sr_n):
                with st.expander(f"Behaviour {i+1}", expanded=True):
                    bn = st.text_input("Behaviour name / label", key=f"sr_bn_{i}",
                                        placeholder="e.g. Physical aggression, Self-injurious behaviour")
                    bd = st.text_area("Describe what it looks like", key=f"sr_bd_{i}", height=70,
                                       placeholder="e.g. Hitting out, throwing objects")
                    bt = st.text_area("Known triggers / when it tends to occur", key=f"sr_bt_{i}", height=70,
                                       placeholder="e.g. When demands are placed, during transitions")
                    beh_raw.append({"name": bn or "", "description": bd or "", "triggers": bt or ""})

            caddbtn, crmbtn = st.columns(2)
            with caddbtn:
                if st.session_state.sr_n < 5 and st.button("＋ Add another behaviour", key="sr_add"):
                    st.session_state.sr_n += 1; st.rerun()
            with crmbtn:
                if st.session_state.sr_n > 1 and st.button("− Remove last", key="sr_rm"):
                    st.session_state.sr_n -= 1; st.rerun()

            valid_behs = [b for b in beh_raw if b["name"].strip()]
            if valid_behs:
                client_info = {
                    "name":      (sr_name or "").strip(),
                    "age":       (sr_age  or "").strip() or "not specified",
                    "diagnosis": (sr_diag or "").strip() or "not specified",
                    "comms":     sr_comms or "Verbal",
                    "other":     (sr_other or "").strip() or "none provided",
                }

    st.divider()

    # ── STEP 2: Strategy library — always embedded, override optional ──────────
    st.markdown("#### Step 2 — Strategy library")

    # Show built-in library confirmation
    lib_word_count = len(EMBEDDED_LIBRARY.split())
    st.success(
        f"✅  **Built-in library loaded:** Hoggs Health Hub PBS Strategy Library "
        f"({lib_word_count:,} words) — strategies will always be selected from this library."
    )

    # Allow optional override for this session
    with st.expander("Upload a different library for this session (optional)"):
        sr_lib_file = st.file_uploader(
            "Upload alternative strategy library (PDF or Word)",
            type=["pdf","docx"], key="sr_lib_upload",
        )
        if sr_lib_file:
            fb = sr_lib_file.read()
            uploaded_lib_text = extract_text_from_docx(fb) if sr_lib_file.name.lower().endswith(".docx") \
                            else extract_text_from_pdf(fb)
            st.session_state["sr_lib_text"] = uploaded_lib_text
            st.session_state["sr_lib_name"] = sr_lib_file.name
            wc = len(uploaded_lib_text.split())
            st.success(f"✅  Override library loaded: **{sr_lib_file.name}** ({wc:,} words)")

        if "sr_lib_text" in st.session_state:
            if st.button("↩ Revert to built-in library", key="sr_lib_revert"):
                del st.session_state["sr_lib_text"]
                del st.session_state["sr_lib_name"]
                st.rerun()

    # Resolve effective library
    if "sr_lib_text" in st.session_state:
        effective_lib_text = st.session_state["sr_lib_text"]
        effective_lib_name = st.session_state["sr_lib_name"]
    else:
        effective_lib_text = EMBEDDED_LIBRARY
        effective_lib_name = "Hoggs Health Hub PBS Strategy Library"

    st.divider()

    # ── STEP 3: Report footer ─────────────────────────────────────────────────
    st.markdown("#### Step 3 — Report footer (optional)")
    cp, cc = st.columns(2)
    with cp: sr_prac    = st.text_input("Your name / title", key="sr_prac",
                                         placeholder="e.g. Janine Hogg — PBS Practitioner")
    with cc: sr_contact = st.text_input("Contact email", key="sr_contact",
                                         placeholder="e.g. janine@hoggshealthhub.com.au")

    st.divider()

    # ── GENERATE ──────────────────────────────────────────────────────────────
    can_go  = bool(client_info and valid_behs and api_key)
    rec_btn = st.button("Generate strategy recommendations", type="primary",
                         disabled=not can_go, key="sr_gen")

    if not api_key:
        st.info("Enter your Anthropic API key in the sidebar to enable this tool.")
    elif not client_info or not valid_behs:
        st.info("Complete Step 1 to provide client information before generating.")

    if rec_btn and can_go:
        library_text  = effective_lib_text
        lib_label     = effective_lib_name
        client_label  = client_info.get("name", "client")
        is_freetext   = bool(valid_behs and valid_behs[0].get("_freetext"))

        spinner_msg = (
            f"Matching strategies from {lib_label} for {client_label}…"
        )

        with st.spinner(spinner_msg):
            try:
                if is_freetext:
                    result = recommend_from_freetext(
                        client_info, freetext_value, api_key, library_text)
                else:
                    result = recommend_strategies(client_info, valid_behs, api_key, library_text)
            except Exception as e:
                st.error(f"Could not generate recommendations: {e}"); st.stop()

        st.success(f"✅  Recommendations generated for **{client_label}** from **{lib_label}**")
        st.divider()

        if result.get("client_summary"):
            st.info(f"**Clinical summary:** {result['client_summary']}")

        gen_s = result.get("general_strategies", [])
        if gen_s:
            st.markdown("##### General strategies — apply across all behaviours")
            for s in gen_s: st.markdown(f"- {s}")

        for beh in result.get("behaviours", []):
            st.divider()
            st.markdown(f"### {beh.get('behaviour','Behaviour')}")
            if beh.get("likely_function"):
                st.markdown(
                    f"<div style='background:#FDEBD0;padding:8px 12px;border-radius:6px;"
                    f"border-left:4px solid #D4700A;margin-bottom:10px'>"
                    f"<strong style='color:#D4700A'>Likely function:</strong> "
                    f"<span style='color:#1A2B35'>{beh['likely_function']}</span></div>",
                    unsafe_allow_html=True)
            cl, cr = st.columns(2)
            with cl:
                if beh.get("proactive"):
                    st.markdown("<div style='background:#DFFAE9;padding:6px 10px;border-radius:4px;"
                                "margin-bottom:6px'><strong style='color:#27AE60'>Proactive strategies"
                                "</strong></div>", unsafe_allow_html=True)
                    for s in beh["proactive"]: st.markdown(f"- {s}")
                if beh.get("teach_instead"):
                    st.markdown("<div style='background:#DCEEF8;padding:6px 10px;border-radius:4px;"
                                "margin:10px 0 6px'><strong style='color:#1A5276'>Skills to build"
                                "</strong></div>", unsafe_allow_html=True)
                    for s in beh["teach_instead"]: st.markdown(f"- {s}")
            with cr:
                if beh.get("reactive"):
                    st.markdown("<div style='background:#E0F6F3;padding:6px 10px;border-radius:4px;"
                                "margin-bottom:6px'><strong style='color:#1A9B8A'>Reactive strategies"
                                "</strong></div>", unsafe_allow_html=True)
                    for s in beh["reactive"]: st.markdown(f"- {s}")
                if beh.get("avoid"):
                    st.markdown("<div style='background:#FADBD8;padding:6px 10px;border-radius:4px;"
                                "margin:10px 0 6px'><strong style='color:#C0392B'>DO NOT"
                                "</strong></div>", unsafe_allow_html=True)
                    for s in beh["avoid"]: st.markdown(f"- {s}")

        st.divider()
        with st.spinner("Building PDF report…"):
            try:
                report_buf = generate_strategy_report(
                    result, client_label,
                    (sr_prac or "").strip(), (sr_contact or "").strip())
            except Exception as e: st.error(f"PDF error: {e}"); st.stop()

        st.download_button(
            "📥 Download strategy report (PDF)", report_buf,
            f"{client_label.replace(' ','_')}_Strategy_Report.pdf",
            "application/pdf", use_container_width=True)

    st.divider()
    st.caption(
        "Recommendations are AI-generated and selected from the Hoggs Health Hub PBS Strategy Library. "
        "Always review before implementing. "
        "This tool complements but does not replace a formal Behaviour Support Plan."
    )
